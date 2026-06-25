"""
Word文档操作模块 - AI可调用的Word文档编辑工具
使用python-docx库实现.docx文件的读取/编辑
"""
from .基类 import 操作结果, 操作基类


class 读取Word(操作基类):
    名称 = "读取Word"
    描述 = "读取.docx文件内容，返回所有段落和表格的文本，AI可据此了解文档结构"
    参数结构 = {
        "路径": {"类型": "字符串", "必填": True, "说明": "Word文档(.docx)路径"}
    }

    def 执行(self, 参数: dict) -> 操作结果:
        路径 = 参数.get("路径", "")
        if not 路径:
            return 操作结果.失败("路径为空")
        try:
            from docx import Document
            doc = Document(路径)
            parts = []
            for i, p in enumerate(doc.paragraphs):
                if p.text.strip():
                    parts.append(f"[段落{i}] ({p.style.name}): {p.text}")
            for ti, table in enumerate(doc.tables):
                parts.append(f"\n[表格{ti}] {len(table.rows)}行×{len(table.columns)}列:")
                for ri, row in enumerate(table.rows):
                    cells = [cell.text for cell in row.cells]
                    parts.append(f"  行{ri}: {' | '.join(cells)}")
            if not parts:
                return 操作结果.成功("文档为空")
            return 操作结果.成功("\n".join(parts), {"操作类型": "读取Word", "段落数": len(doc.paragraphs), "表格数": len(doc.tables)})
        except Exception as e:
            return 操作结果.失败(f"读取Word失败: {e}")


class 替换Word文本(操作基类):
    名称 = "替换Word文本"
    描述 = "在.docx文件中查找并替换文本，支持单段和多段（跨段落框选）匹配"
    参数结构 = {
        "路径": {"类型": "字符串", "必填": True, "说明": "Word文档(.docx)路径"},
        "旧文本": {"类型": "字符串", "必填": True, "说明": "要查找的文本（可跨多段，用换行分隔）"},
        "新文本": {"类型": "字符串", "必填": True, "说明": "替换为的文本"},
        "全部替换": {"类型": "布尔", "必填": False, "说明": "True=替换全部匹配，False=仅替换第一个，默认True"}
    }

    def 执行(self, 参数: dict) -> 操作结果:
        路径 = 参数.get("路径", "")
        旧文本 = 参数.get("旧文本", "")
        新文本 = 参数.get("新文本", "")
        全部 = 参数.get("全部替换", True)
        if not 路径 or not 旧文本:
            return 操作结果.失败("路径和旧文本不能为空")
        try:
            from docx import Document
            doc = Document(路径)
            替换数 = 0

            # 按行分割，过滤空行（mammoth渲染的HTML选区在段落间有换行）
            旧行 = [line.strip() for line in 旧文本.split('\n') if line.strip()]
            新行 = [line.strip() for line in 新文本.split('\n') if line.strip()] if 新文本 else []

            def 替换段落列表(paras):
                nonlocal 替换数
                i = 0
                while i < len(paras):
                    if len(旧行) <= 1:
                        # 单行匹配
                        旧 = 旧行[0] if 旧行 else 旧文本.strip()
                        if 旧 in paras[i].text:
                            done = False
                            for run in paras[i].runs:
                                if 旧 in run.text:
                                    run.text = run.text.replace(旧, 新文本)
                                    done = True
                                    替换数 += 1
                                    if not 全部:
                                        return True
                            if not done:
                                paras[i].text = paras[i].text.replace(旧, 新文本)
                                替换数 += 1
                                if not 全部:
                                    return True
                    else:
                        # 多行跨段落匹配
                        if 旧行[0] in paras[i].text:
                            matched = True
                            for j in range(1, len(旧行)):
                                if i + j >= len(paras) or 旧行[j] not in paras[i + j].text:
                                    matched = False
                                    break
                            if matched:
                                if len(新行) >= len(旧行):
                                    for j in range(len(旧行)):
                                        paras[i + j].text = paras[i + j].text.replace(旧行[j], 新行[j])
                                else:
                                    paras[i].text = paras[i].text.replace(旧行[0], 新文本)
                                    for j in range(1, len(旧行)):
                                        paras[i + j].text = paras[i + j].text.replace(旧行[j], "")
                                替换数 += 1
                                if not 全部:
                                    return True
                                i += len(旧行)
                                continue
                    i += 1
                return False

            found = 替换段落列表(doc.paragraphs)
            if found:
                doc.save(路径)
                return 操作结果.成功(f"已替换1处", {"操作类型": "替换Word文本", "替换数": 1})
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        found = 替换段落列表(cell.paragraphs)
                        if found:
                            doc.save(路径)
                            return 操作结果.成功(f"已替换1处", {"操作类型": "替换Word文本", "替换数": 1})
            if 替换数 > 0:
                doc.save(路径)
                return 操作结果.成功(f"已替换{替换数}处", {"操作类型": "替换Word文本", "替换数": 替换数})
            return 操作结果.失败(f"未找到文本: {旧文本}")
        except Exception as e:
            return 操作结果.失败(f"替换Word文本失败: {e}")


class 追加Word段落(操作基类):
    名称 = "追加Word段落"
    描述 = "在.docx文件末尾追加一段文字，可指定样式（如Normal/Heading 1/Heading 2等）"
    参数结构 = {
        "路径": {"类型": "字符串", "必填": True, "说明": "Word文档(.docx)路径"},
        "内容": {"类型": "字符串", "必填": True, "说明": "要追加的文本内容"},
        "样式": {"类型": "字符串", "必填": False, "说明": "段落样式，如Normal/Heading 1/Heading 2，默认Normal"}
    }

    def 执行(self, 参数: dict) -> 操作结果:
        路径 = 参数.get("路径", "")
        内容 = 参数.get("内容", "")
        样式 = 参数.get("样式", "Normal")
        if not 路径 or not 内容:
            return 操作结果.失败("路径和内容不能为空")
        try:
            from docx import Document
            doc = Document(路径)
            doc.add_paragraph(内容, style=样式)
            doc.save(路径)
            return 操作结果.成功(f"已追加段落，样式: {样式}", {"操作类型": "追加Word段落"})
        except Exception as e:
            return 操作结果.失败(f"追加Word段落失败: {e}")


class 插入Word段落(操作基类):
    名称 = "插入Word段落"
    描述 = "在.docx文件指定位置插入一段文字（0=最前，N=第N段之后）"
    参数结构 = {
        "路径": {"类型": "字符串", "必填": True, "说明": "Word文档(.docx)路径"},
        "内容": {"类型": "字符串", "必填": True, "说明": "要插入的文本内容"},
        "位置": {"类型": "整数", "必填": True, "说明": "插入位置，0=最前，N=第N段之后插入"},
        "样式": {"类型": "字符串", "必填": False, "说明": "段落样式，默认Normal"}
    }

    def 执行(self, 参数: dict) -> 操作结果:
        路径 = 参数.get("路径", "")
        内容 = 参数.get("内容", "")
        位置 = 参数.get("位置", 0)
        样式 = 参数.get("样式", "Normal")
        if not 路径 or not 内容:
            return 操作结果.失败("路径和内容不能为空")
        try:
            from docx import Document
            doc = Document(路径)
            if 位置 < 0:
                位置 = 0
            if 位置 > len(doc.paragraphs):
                位置 = len(doc.paragraphs)
            if 位置 < len(doc.paragraphs):
                ref = doc.paragraphs[位置]
                new_p = ref.insert_paragraph_before(内容, style=样式)
            else:
                new_p = doc.add_paragraph(内容, style=样式)
            doc.save(路径)
            return 操作结果.成功(f"已在位置{位置}插入段落", {"操作类型": "插入Word段落"})
        except Exception as e:
            return 操作结果.失败(f"插入Word段落失败: {e}")


class 删除Word段落(操作基类):
    名称 = "删除Word段落"
    描述 = "删除.docx文件中指定序号的段落（0=第一段）"
    参数结构 = {
        "路径": {"类型": "字符串", "必填": True, "说明": "Word文档(.docx)路径"},
        "序号": {"类型": "整数", "必填": True, "说明": "要删除的段落序号，0=第一段"}
    }

    def 执行(self, 参数: dict) -> 操作结果:
        路径 = 参数.get("路径", "")
        序号 = 参数.get("序号", -1)
        if not 路径:
            return 操作结果.失败("路径为空")
        if 序号 < 0:
            return 操作结果.失败("序号必须>=0")
        try:
            from docx import Document
            doc = Document(路径)
            if 序号 >= len(doc.paragraphs):
                return 操作结果.失败(f"序号超出范围，文档共{len(doc.paragraphs)}段")
            p = doc.paragraphs[序号]
            p_text = p.text[:50]
            p._element.getparent().remove(p._element)
            doc.save(路径)
            return 操作结果.成功(f"已删除段落{序号}: {p_text}...", {"操作类型": "删除Word段落"})
        except Exception as e:
            return 操作结果.失败(f"删除Word段落失败: {e}")


class 新建Word文档(操作基类):
    名称 = "新建Word文档"
    描述 = "创建一个新的.docx文件，可写入初始内容"
    参数结构 = {
        "路径": {"类型": "字符串", "必填": True, "说明": "新Word文档路径"},
        "标题": {"类型": "字符串", "必填": False, "说明": "文档标题（Heading样式），默认无"},
        "内容": {"类型": "字符串", "必填": False, "说明": "正文内容，多段用\\n分隔"}
    }

    def 执行(self, 参数: dict) -> 操作结果:
        路径 = 参数.get("路径", "")
        标题 = 参数.get("标题", "")
        内容 = 参数.get("内容", "")
        if not 路径:
            return 操作结果.失败("路径为空")
        try:
            from docx import Document
            doc = Document()
            if 标题:
                doc.add_heading(标题, level=1)
            if 内容:
                for line in 内容.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)
            doc.save(路径)
            return 操作结果.成功(f"已创建Word文档: {路径}", {"操作类型": "新建Word文档"})
        except Exception as e:
            return 操作结果.失败(f"创建Word文档失败: {e}")
