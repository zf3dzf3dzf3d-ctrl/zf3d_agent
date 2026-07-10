"""
网页服务 - 内置Web服务+API接口
API路径全部使用英文，避免中文URL编码问题
"""
import json
import os
import sys
import string
import shutil
import subprocess
import struct
import time
import threading
from datetime import datetime
from http.server import ThreadingHTTPServer, BaseHTTPRequestHandler
from urllib.parse import urlparse, parse_qs, unquote
import urllib.request
from pathlib import Path
import zlib
import array


def _tga转png(路径, 最大宽高: int = 0) -> bytes:
    """纯标准库TGA→PNG转码，支持未压缩(类型2)和RLE压缩(类型10)。
    最大宽高>0时缩放到指定尺寸（用于缩略图加速）。"""
    with open(路径, "rb") as f:
        数据 = f.read()
    if len(数据) < 18:
        raise ValueError("文件过小")
    # TGA头解析
    id长度 = 数据[0]
    颜色表类型 = 数据[1]
    图像类型 = 数据[2]
    宽度 = struct.unpack("<H", 数据[12:14])[0]
    高度 = struct.unpack("<H", 数据[14:16])[0]
    像素位数 = 数据[16]
    描述符 = 数据[17]
    偏移 = 18 + id长度
    if 颜色表类型 != 0:
        颜色表长度 = struct.unpack("<H", 数据[5:7])[0]
        颜色表项大小 = 数据[7]
        偏移 += 颜色表长度 * (颜色表项大小 // 8)
    if 像素位数 not in (24, 32):
        raise ValueError(f"不支持的像素位数: {像素位数}")
    字节每像素 = 像素位数 // 8
    像素数 = 宽度 * 高度
    # 解码像素数据
    if 图像类型 == 2:
        原始 = 数据[偏移:偏移 + 像素数 * 字节每像素]
        if len(原始) < 像素数 * 字节每像素:
            raise ValueError("像素数据不完整")
    elif 图像类型 == 10:
        原始 = bytearray()
        i = 偏移
        while len(原始) < 像素数 * 字节每像素 and i < len(数据):
            头 = 数据[i]; i += 1
            if 头 & 0x80:
                count = (头 & 0x7F) + 1
                if i + 字节每像素 > len(数据):
                    break
                像素 = 数据[i:i + 字节每像素]; i += 字节每像素
                原始.extend(像素 * count)
            else:
                count = 头 + 1
                需要 = count * 字节每像素
                if i + 需要 > len(数据):
                    原始.extend(数据[i:])
                    break
                原始.extend(数据[i:i + 需要]); i += 需要
    else:
        raise ValueError(f"不支持的TGA类型: {图像类型}")
    # BGR(A) → RGBA，用array批量位运算
    if 字节每像素 == 4:
        src = array.array("I")
        src.frombytes(原始[:像素数 * 4])
        dst = array.array("I", [0] * 像素数)
        for i in range(像素数):
            p = src[i]
            dst[i] = (p & 0xFF00FF00) | ((p >> 16) & 0xFF) | ((p & 0xFF) << 16)
        rgba = bytearray(dst.tobytes())
    else:
        rgba = bytearray(像素数 * 4)
        for p in range(像素数):
            b = 原始[p * 3]
            g = 原始[p * 3 + 1]
            r = 原始[p * 3 + 2]
            rgba[p * 4] = r
            rgba[p * 4 + 1] = g
            rgba[p * 4 + 2] = b
            rgba[p * 4 + 3] = 255
    # TGA默认原点在左下角，需翻转
    翻转 = (描述符 & 0x20) == 0
    if 翻转:
        行字节 = 宽度 * 4
        for y in range(高度 // 2):
            r1 = y * 行字节
            r2 = (高度 - 1 - y) * 行字节
            rgba[r1:r1 + 行字节], rgba[r2:r2 + 行字节] = rgba[r2:r2 + 行字节], rgba[r1:r1 + 行字节]
    # 构建PNG
    return _编码png(bytes(rgba), 宽度, 高度)


def _编码png(rgba数据: bytes, 宽度: int, 高度: int) -> bytes:
    """纯标准库RGBA→PNG编码"""
    def _chunk(类型: bytes, 数据: bytes) -> bytes:
        c = 类型 + 数据
        return struct.pack(">I", len(数据)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
    # PNG签名
    签名 = b"\x89PNG\r\n\x1a\n"
    # IHDR
    ihdr = struct.pack(">IIBBBBB", 宽度, 高度, 8, 6, 0, 0, 0)  # 8bit, RGBA
    # IDAT: 每行前加filter字节(0)
    行大小 = 宽度 * 4
    原始 = bytearray()
    for y in range(高度):
        原始.append(0)
        原始.extend(rgba数据[y * 行大小:(y + 1) * 行大小])
    idat = zlib.compress(bytes(原始), 9)
    return 签名 + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", idat) + _chunk(b"IEND", b"")


def _提取docx文本(文件路径):
    """从.docx文件提取带格式的文本，返回HTML（保留字体颜色、粗体、斜体等）"""
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(文件路径)
    html_parts = []

    # 样式名→HTML标签映射
    标题映射 = {"Heading 1": "h1", "Heading 2": "h2", "Heading 3": "h3",
               "Heading 4": "h4", "Heading 5": "h5", "Heading 6": "h6",
               "Title": "h1", "Subtitle": "h2"}

    def _run转html(run):
        """将单个run转为带inline style的HTML"""
        text = run.text or ""
        if not text:
            return ""
        # HTML转义
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        # 保留换行
        text = text.replace("\n", "<br>")

        styles = []
        # 字体颜色
        try:
            if run.font.color and run.font.color.rgb:
                styles.append(f"color:#{run.font.color.rgb}")
        except:
            pass
        # 字体大小
        try:
            if run.font.size:
                styles.append(f"font-size:{run.font.size.pt}pt")
        except:
            pass
        # 字体名称
        try:
            if run.font.name:
                styles.append(f"font-family:{run.font.name}")
        except:
            pass
        # 粗体
        if run.font.bold:
            text = f"<strong>{text}</strong>"
        # 斜体
        if run.font.italic:
            text = f"<em>{text}</em>"
        # 下划线
        if run.font.underline:
            text = f"<u>{text}</u>"
        # 删除线
        try:
            if run.font.strike:
                text = f"<s>{text}</s>"
        except:
            pass
        # 包裹inline style
        if styles:
            text = f'<span style="{";".join(styles)}">{text}</span>'
        return text

    def _段落转html(p):
        """将段落转为HTML"""
        # 检查是否为标题
        style_name = p.style.name if p.style else ""
        tag = 标题映射.get(style_name, "p")
        # 合并所有run
        inner = "".join(_run转html(run) for run in p.runs)
        if not inner.strip():
            inner = "&nbsp;"
        # 对齐方式
        align_map = {WD_ALIGN_PARAGRAPH.CENTER: "center", WD_ALIGN_PARAGRAPH.RIGHT: "right",
                     WD_ALIGN_PARAGRAPH.JUSTIFY: "justify"}
        align = ""
        try:
            if p.alignment and p.alignment in align_map:
                align = f' style="text-align:{align_map[p.alignment]}"'
        except:
            pass
        return f"<{tag}{align}>{inner}</{tag}>"

    def _遍历文档体(doc):
        """按文档顺序遍历段落和表格"""
        from docx.oxml.ns import qn
        body = doc.element.body
        for child in body.iterchildren():
            if child.tag == qn('w:p'):
                # 段落
                for p in doc.paragraphs:
                    if p._element is child:
                        html_parts.append(_段落转html(p))
                        break
            elif child.tag == qn('w:tbl'):
                # 表格
                for tbl in doc.tables:
                    if tbl._element is child:
                        html_parts.append(_表格转html(tbl))
                        break

    def _表格转html(tbl):
        """将表格转为HTML"""
        rows_html = []
        for row in tbl.rows:
            cells_html = []
            for cell in row.cells:
                cell_parts = []
                for p in cell.paragraphs:
                    cell_parts.append(_段落转html(p))
                cells_html.append(f"<td>{''.join(cell_parts) or '&nbsp;'}</td>")
            rows_html.append(f"<tr>{''.join(cells_html)}</tr>")
        return f'<table>{"".join(rows_html)}</table>'

    _遍历文档体(doc)
    return "\n".join(html_parts) if html_parts else "<p>（空文档）</p>"


def _提取doc文本(文件路径):
    """从.doc文件提取文本，返回HTML"""
    import olefile
    ole = olefile.OleFileIO(文件路径)
    word_data = ole.openstream('WordDocument').read()
    flags = struct.unpack_from('<H', word_data, 0x000A)[0]
    fComplex = (flags & 0x0004) != 0
    fExtChar = (flags & 0x1000) != 0
    fWhichTblStm = (flags & 0x0200) != 0
    fcMin = struct.unpack_from('<I', word_data, 0x0018)[0]
    nFib = struct.unpack_from('<H', word_data, 0x0002)[0]
    text = ""
    if nFib >= 0x00C1:
        ccpText = struct.unpack_from('<I', word_data, 0x004C)[0]
        if fComplex:
            table_name = '1Table' if fWhichTblStm else '0Table'
            table_data = ole.openstream(table_name).read()
            fcClx = struct.unpack_from('<I', word_data, 0x01A2)[0]
            lcbClx = struct.unpack_from('<I', word_data, 0x01A6)[0]
            clx = table_data[fcClx:fcClx + lcbClx]
            parts = []
            pos = 0
            while pos < len(clx):
                clxt = clx[pos]
                if clxt == 2:
                    pos += 1
                    cb = struct.unpack_from('<I', clx, pos)[0]
                    pos += 4
                    n = (cb - 4) // 12
                    cps = []
                    for i in range(n + 1):
                        cps.append(struct.unpack_from('<I', clx, pos + i * 4)[0])
                    pcd_off = pos + (n + 1) * 4
                    for i in range(n):
                        pcd = clx[pcd_off + i * 8: pcd_off + (i + 1) * 8]
                        fc_raw = struct.unpack_from('<I', pcd, 2)[0]
                        compressed = (fc_raw & 0x40000000) != 0
                        fc = fc_raw & 0x3FFFFFFF
                        cnt = cps[i + 1] - cps[i]
                        if compressed:
                            off = fc // 2
                            parts.append(word_data[off:off + cnt].decode('cp1252', errors='ignore'))
                        else:
                            parts.append(word_data[fc:fc + cnt * 2].decode('utf-16-le', errors='ignore'))
                    pos = pcd_off + n * 8
                elif clxt == 1:
                    pos += 1
                    cb = struct.unpack_from('<H', clx, pos)[0]
                    pos += 2 + cb
                else:
                    break
            text = ''.join(parts)
        else:
            if fExtChar:
                text = word_data[fcMin:fcMin + ccpText * 2].decode('utf-16-le', errors='ignore')
            else:
                text = word_data[fcMin:fcMin + ccpText].decode('cp1252', errors='ignore')
    else:
        fcMac = struct.unpack_from('<I', word_data, 0x001C)[0]
        text = word_data[fcMin:fcMac].decode('cp1252', errors='ignore')
    ole.close()
    text = text.replace('\x07', '\t').replace('\x0B', '').replace('\x0C', '')
    html_parts = []
    for p in text.split('\r'):
        p = p.strip()
        if p:
            p = p.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            html_parts.append(f'<p>{p}</p>')
    return '\n'.join(html_parts) if html_parts else '<p>（空文档或无法读取）</p>'


def _获取英文模型目录():
    """获取纯英文路径用于存放sherpa-onnx模型（C++底层不支持中文路径）
    优先级：Python安装目录 → 系统Temp → C盘根目录
    """
    import tempfile, sys
    # 候选路径，确保纯ASCII
    候选 = [
        Path(sys.executable).parent / "zf3d_voice_model",   # Python安装目录旁边
        Path(tempfile.gettempdir()) / "zf3d_voice_model",    # 系统Temp
        Path("C:/zf3d_voice_model"),                          # C盘根目录（兜底）
    ]
    for 路径 in 候选:
        try:
            路径_str = str(路径)
            # 检测路径是否纯ASCII
            路径_str.encode('ascii')
            路径.mkdir(parents=True, exist_ok=True)
            return 路径
        except (UnicodeEncodeError, OSError, PermissionError):
            continue
    # 所有候选都失败，最后用Temp（即使有中文也比没有强）
    路径 = Path(tempfile.gettempdir()) / "zf3d_voice_model"
    路径.mkdir(parents=True, exist_ok=True)
    return 路径


# Kokoro TTS 说话人列表（v1_1，103个）
_KOKORO说话人 = [
    {"sid": 0, "名称": "艾洛伊", "性别": "女", "语言": "美式英语", "头像": "👩"},
    {"sid": 1, "名称": "奥黛", "性别": "女", "语言": "美式英语", "头像": "👩"},
    {"sid": 2, "名称": "贝拉", "性别": "女", "语言": "美式英语", "头像": "👩‍🦰"},
    {"sid": 3, "名称": "赫特", "性别": "女", "语言": "美式英语", "头像": "💕"},
    {"sid": 4, "名称": "杰西卡", "性别": "女", "语言": "美式英语", "头像": "👩‍💼"},
    {"sid": 5, "名称": "科瑞", "性别": "女", "语言": "美式英语", "头像": "🌱"},
    {"sid": 6, "名称": "妮可", "性别": "女", "语言": "美式英语", "头像": "👩‍🔬"},
    {"sid": 7, "名称": "诺娃", "性别": "女", "语言": "美式英语", "头像": "✨"},
    {"sid": 8, "名称": "丽芙", "性别": "女", "语言": "美式英语", "头像": "🌿"},
    {"sid": 9, "名称": "莎拉", "性别": "女", "语言": "美式英语", "头像": "👩‍🏫"},
    {"sid": 10, "名称": "斯凯", "性别": "女", "语言": "美式英语", "头像": "☁️"},
    {"sid": 11, "名称": "亚当", "性别": "男", "语言": "美式英语", "头像": "👨"},
    {"sid": 12, "名称": "埃科", "性别": "男", "语言": "美式英语", "头像": "🔊"},
    {"sid": 13, "名称": "埃里克", "性别": "男", "语言": "美式英语", "头像": "👨‍💻"},
    {"sid": 14, "名称": "芬里尔", "性别": "男", "语言": "美式英语", "头像": "🐺"},
    {"sid": 15, "名称": "利亚姆", "性别": "男", "语言": "美式英语", "头像": "👦"},
    {"sid": 16, "名称": "迈克尔", "性别": "男", "语言": "美式英语", "头像": "🎸"},
    {"sid": 17, "名称": "奥尼克斯", "性别": "男", "语言": "美式英语", "头像": "💎"},
    {"sid": 18, "名称": "帕克", "性别": "男", "语言": "美式英语", "头像": "🃏"},
    {"sid": 19, "名称": "圣诞老人", "性别": "男", "语言": "美式英语", "头像": "🎅"},
    {"sid": 20, "名称": "爱丽丝", "性别": "女", "语言": "英式英语", "头像": "👸"},
    {"sid": 21, "名称": "艾玛", "性别": "女", "语言": "英式英语", "头像": "👒"},
    {"sid": 22, "名称": "伊莎贝拉", "性别": "女", "语言": "英式英语", "头像": "💃"},
    {"sid": 23, "名称": "莉莉", "性别": "女", "语言": "英式英语", "头像": "🌷"},
    {"sid": 24, "名称": "丹尼尔", "性别": "男", "语言": "英式英语", "头像": "🦁"},
    {"sid": 25, "名称": "费布尔", "性别": "男", "语言": "英式英语", "头像": "📖"},
    {"sid": 26, "名称": "乔治", "性别": "男", "语言": "英式英语", "头像": "🐉"},
    {"sid": 27, "名称": "刘易斯", "性别": "男", "语言": "英式英语", "头像": "🎩"},
    {"sid": 28, "名称": "朵拉", "性别": "女", "语言": "英语", "头像": "🧭"},
    {"sid": 29, "名称": "亚历克斯", "性别": "男", "语言": "英语", "头像": "🧑"},
    {"sid": 30, "名称": "西维斯", "性别": "女", "语言": "法语", "头像": "🇫🇷"},
    {"sid": 31, "名称": "阿尔法", "性别": "女", "语言": "印地语", "头像": "α"},
    {"sid": 32, "名称": "贝塔", "性别": "女", "语言": "印地语", "头像": "β"},
    {"sid": 33, "名称": "欧米伽", "性别": "男", "语言": "印地语", "头像": "Ω"},
    {"sid": 34, "名称": "普西", "性别": "男", "语言": "印地语", "头像": "ψ"},
    {"sid": 35, "名称": "萨拉", "性别": "女", "语言": "意大利语", "头像": "🍝"},
    {"sid": 36, "名称": "尼古拉", "性别": "男", "语言": "意大利语", "头像": "🍕"},
    {"sid": 37, "名称": "阿尔法", "性别": "女", "语言": "日语", "头像": "🌸"},
    {"sid": 38, "名称": "权狐", "性别": "女", "语言": "日语", "头像": "🦊"},
    {"sid": 39, "名称": "鼠", "性别": "女", "语言": "日语", "头像": "🐭"},
    {"sid": 40, "名称": "手袋", "性别": "女", "语言": "日语", "头像": "👜"},
    {"sid": 41, "名称": "蜘蛛", "性别": "男", "语言": "日语", "头像": "🕷️"},
    {"sid": 42, "名称": "朵拉", "性别": "女", "语言": "葡萄牙语", "头像": "🇧🇷"},
    {"sid": 43, "名称": "圣诞老人", "性别": "男", "语言": "葡萄牙语", "头像": "🎁"},
    {"sid": 44, "名称": "小贝", "性别": "女", "语言": "中文", "头像": "👧"},
    {"sid": 45, "名称": "小妮", "性别": "女", "语言": "中文", "头像": "🧒"},
    {"sid": 46, "名称": "晓晓", "性别": "女", "语言": "中文", "头像": "👩"},
    {"sid": 47, "名称": "晓伊", "性别": "女", "语言": "中文", "头像": "💁‍♀️"},
    {"sid": 48, "名称": "云健", "性别": "男", "语言": "中文", "头像": "💪"},
    {"sid": 49, "名称": "云希", "性别": "男", "语言": "中文", "头像": "👨"},
    {"sid": 50, "名称": "云夏", "性别": "男", "语言": "中文", "头像": "👦"},
    {"sid": 51, "名称": "云扬", "性别": "男", "语言": "中文", "头像": "🎙️"},
]
# v1.1新增的50个中文说话人（sid 52-102，来自LongMaoData数据集）
_中文音色名 = [
    "嘉怡","子轩","雨桐","浩然","思琪","俊杰","梦瑶","天宇","欣怡","志强",
    "雅婷","博文","若曦","子墨","佳琪","明轩","诗涵","宇辰","雨欣","健豪",
    "淑芬","伟杰","美玲","家豪","秀英","建国","丽华","国庆","玉兰","建军",
    "晓燕","德明","春梅","学东","红梅","永强","翠兰","建军","秋菊","国华",
    "丽君","文博","静雯","子涵","雨泽","思源","博文","佳豪","若彤","梓萱"
]
_中文头像池 = ["👩","👨","👧","👦","🧑","👱‍♀️","👱‍♂️","👩‍🦰","👨‍🦰","👩‍🦱","👨‍🦱","👩‍🦳","👨‍🦳","👩‍💼","👨‍💼","👩‍🏫","👨‍🏫","👩‍🔬","👨‍🔬","👩‍💻","👨‍💻","👩‍🎨","👨‍🎨","👩‍🔧","👨‍🔧","🧑‍🌾","🧑‍🍳","🧑‍🎤","🧑‍🎤","🧑‍🏭","🧑‍💼","🧑‍🏫","🧑‍🔬","🧑‍💻","🧑‍🎨","🧑‍🔧","👷‍♀️","👷‍♂️","👮‍♀️","👮‍♂️","🕵️‍♀️","🕵️‍♂️","💂‍♀️","💂‍♂️","🤴","👸","👳‍♀️","👳‍♂️","🧕"]
for _i in range(52, 103):
    _idx = _i - 52
    _KOKORO说话人.append({"sid": _i, "名称": _中文音色名[_idx % len(_中文音色名)], "性别": "女" if _idx % 2 == 0 else "男", "语言": "中文", "头像": _中文头像池[_idx % len(_中文头像池)]})


def _获取KokoroTTS引擎():
    """懒加载Kokoro TTS引擎，返回 sherpa_onnx.OfflineTts 实例或 None"""
    if 网页请求处理器._kokoroTTS引擎 is not None:
        return 网页请求处理器._kokoroTTS引擎
    try:
        import sherpa_onnx
    except ImportError:
        return None
    模型目录 = _获取英文模型目录() / "kokoro-tts"
    # 兼容 int8 和非 int8 版本的模型文件名
    模型文件 = 模型目录 / "model.onnx"
    if not 模型文件.exists():
        模型文件 = 模型目录 / "model.int8.onnx"
    if not 模型文件.exists():
        return None
    try:
        tts_config = sherpa_onnx.OfflineTtsConfig(
            model=sherpa_onnx.OfflineTtsModelConfig(
                kokoro=sherpa_onnx.OfflineTtsKokoroModelConfig(
                    model=str(模型文件),
                    voices=str(模型目录 / "voices.bin"),
                    tokens=str(模型目录 / "tokens.txt"),
                    data_dir=str(模型目录 / "espeak-ng-data"),
                    lexicon=f"{模型目录 / 'lexicon-us-en.txt'},{模型目录 / 'lexicon-zh.txt'}",
                ),
                provider="cpu",
                num_threads=2,
            ),
            max_num_sentences=1,
        )
        if not tts_config.validate():
            return None
        网页请求处理器._kokoroTTS引擎 = sherpa_onnx.OfflineTts(tts_config)
        return 网页请求处理器._kokoroTTS引擎
    except Exception as e:
        print(f"  ⚠️ Kokoro TTS引擎加载失败: {e}")
        return None


def _获取当前员工语音配置(员工名):
    """从员工配置或系统配置获取语音设置"""
    默认 = {"引擎": "本地", "说话人ID": 47, "语速": 1.0, "edge音色": "zh-CN-XiaoxiaoNeural"}
    try:
        系统配置 = 网页请求处理器.配置加载器.配置缓存.get("系统配置", {})
        系统语音 = 系统配置.get("语音输出", {})
        if 系统语音:
            默认.update(系统语音)
    except Exception:
        pass
    if not 员工名 or 员工名 == "母体":
        return 默认
    try:
        员工配置 = 网页请求处理器.配置加载器.配置缓存.get("员工配置", {})
        for 员工 in 员工配置.get("员工列表", []):
            if 员工.get("姓名") == 员工名:
                语音 = 员工.get("语音")
                if 语音:
                    默认.update(语音)
                break
    except Exception:
        pass
    return 默认


def _float32转WAV(samples, sample_rate):
    """将 sherpa-onnx 输出的 float32 采样转为 WAV 临时文件，返回文件路径"""
    import wave, tempfile, numpy as np
    samples_int16 = (np.asarray(samples, dtype=np.float32) * 32767).astype(np.int16)
    wav路径 = os.path.join(tempfile.gettempdir(), f'zf3d_kokoro_{threading.get_ident()}.wav')
    with wave.open(wav路径, 'wb') as wf:
        wf.setnchannels(1)
        wf.setsampwidth(2)
        wf.setframerate(sample_rate)
        wf.writeframes(samples_int16.tobytes())
    return wav路径


def _检查KokoroTTS模型存在():
    """检查Kokoro TTS模型是否已下载"""
    模型目录 = _获取英文模型目录() / "kokoro-tts"
    return (模型目录 / "model.onnx").exists() or (模型目录 / "model.int8.onnx").exists()


class 网页请求处理器(BaseHTTPRequestHandler):
    """HTTP请求处理器"""
    界面目录 = None
    文件管理器 = None
    配置加载器 = None
    模型直连器 = None
    模块注册 = None
    操作注册中心 = None
    运行诊断器 = None  # 运行诊断器实例
    当前模型名 = None  # 当前对话使用的模型名
    _启动器实例 = None
    _定时任务调度器 = None
    _tts主界面状态 = {"播放中": False, "代次": 0}  # 主界面语音播报状态
    _tts轮盘状态 = {"播放中": False, "代次": 0}    # 轮盘朗读状态
    _sherpa识别器 = None  # sherpa-onnx 离线语音识别器（懒加载）
    _sherpa流式识别器 = None  # sherpa-onnx 流式语音识别器（懒加载）
    _语音安装状态 = {"步骤": "", "进度": 0, "完成": False, "错误": ""}  # 安装进度
    _kokoroTTS引擎 = None  # sherpa-onnx Kokoro TTS引擎（懒加载）
    _tts安装状态 = {"步骤": "", "进度": 0, "完成": False, "错误": ""}  # TTS模型安装进度
    _最后打开的文件夹 = None  # 前端最后打开的文件夹（录音/录屏默认保存位置）
    _录屏设置 = {"点击效果": False, "点击音效": False, "音效音量": 50, "帧率": 30, "音频模式": "system",
                  "麦克风音量": 1.0, "麦克风静音": False, "系统音量": 1.0, "系统静音": False,
                  "dshow设备名": ""}

    def do_GET(self):
        self._http方法 = "GET"
        try:
            解析结果 = urlparse(self.path)
            路径 = unquote(解析结果.path)
            查询串 = 解析结果.query or ""

            # WebSocket升级检测
            if 路径 == "/api/voice-stream" and self.headers.get("Upgrade", "").lower() == "websocket":
                self._处理WebSocket语音()
                return

            if 路径 == "/" or 路径 == "/index.html":
                self._返回文件(self.界面目录 / "主页.html", "text/html", 查询串)
            elif ".." in 路径:
                # 防止路径穿越攻击
                self.send_response(403)
                self.end_headers()
                self.wfile.write("forbidden".encode("utf-8"))
            elif 路径.endswith(".css"):
                self._返回文件(self.界面目录 / 路径.lstrip("/"), "text/css", 查询串)
            elif 路径.endswith(".js"):
                self._返回文件(self.界面目录 / 路径.lstrip("/"), "application/javascript", 查询串)
            elif 路径.startswith("/monaco/"):
                # Monaco Editor 静态文件（JS/CSS/字体等）
                文件路径 = self.界面目录 / 路径.lstrip("/")
                文件类型 = self._猜测类型(路径)
                if 路径.endswith(".js"):
                    文件类型 = "application/javascript"
                elif 路径.endswith(".css"):
                    文件类型 = "text/css"
                self._返回文件(文件路径, 文件类型, 查询串)
            elif 路径.startswith("/api/"):
                self._处理API_GET(路径, 解析结果)
            else:
                self._返回文件(self.界面目录 / 路径.lstrip("/"), self._猜测类型(路径), 查询串)
        except Exception as e:
            if isinstance(e, (ConnectionAbortedError, ConnectionResetError, BrokenPipeError, OSError)):
                return  # 客户端已断开/连接异常，无需处理
            print(f"  ❌ GET异常: {e}")
            if self.运行诊断器:
                self.运行诊断器.记录错误("网页服务.do_GET", e)
            try:
                self._返回JSON({"错误": f"服务器异常: {str(e)}"}, 500)
            except Exception:
                return  # 响应也失败了，放弃

    def do_POST(self):
        self._http方法 = "POST"
        self._处理POST请求()

    def do_PUT(self):
        self._http方法 = "PUT"
        self._处理POST请求()

    def do_DELETE(self):
        self._http方法 = "DELETE"
        self._处理POST请求()

    def _处理POST请求(self):
        try:
            解析结果 = urlparse(self.path)
            路径 = unquote(解析结果.path)
            if 路径.startswith("/api/"):
                内容长度 = int(self.headers.get("Content-Length", 0))
                原始体 = self.rfile.read(内容长度) if 内容长度 > 0 else b"{}"
                ctype = self.headers.get("Content-Type", "")
                if "multipart/form-data" in ctype:
                    # multipart请求不解析JSON，保留原始字节供handler读取
                    self._multipart_body = 原始体
                    self._处理API_POST(路径, {})
                else:
                    请求体 = 原始体.decode("utf-8") if 原始体 else "{}"
                    try:
                        请求数据 = json.loads(请求体)
                    except json.JSONDecodeError:
                        请求数据 = {}
                    self._处理API_POST(路径, 请求数据)
            else:
                self._返回JSON({"错误": "未知路径"}, 404)
        except Exception as e:
            if isinstance(e, (ConnectionAbortedError, ConnectionResetError, BrokenPipeError, OSError)):
                return  # 客户端已断开/连接异常，无需处理
            print(f"  ❌ POST异常: {e}")
            if self.运行诊断器:
                self.运行诊断器.记录错误("网页服务.do_POST", e)
            try:
                self._返回JSON({"错误": f"服务器异常: {str(e)}"}, 500)
            except Exception:
                return  # 响应也失败了，放弃

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "http://localhost:8765")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type, Authorization")
        self.end_headers()

    def _检查鉴权(self) -> bool:
        """检查API鉴权：启用时非localhost请求需携带正确令牌"""
        配置 = self.配置加载器.配置缓存.get("系统配置", {})
        鉴权配置 = 配置.get("API鉴权", {})
        if not 鉴权配置.get("启用", False):
            return True
        # localhost免鉴权
        客户端地址 = self.client_address[0]
        if 客户端地址 in ("127.0.0.1", "::1", "localhost"):
            return True
        # 检查Bearer token
        令牌 = 鉴权配置.get("令牌", "")
        if not 令牌:
            return True
        auth头 = self.headers.get("Authorization", "")
        if auth头.startswith("Bearer "):
            提供的令牌 = auth头[7:]
            if 提供的令牌 == 令牌:
                return True
        return False

    def _处理API_GET(self, 路径: str, 解析结果):
        if not self._检查鉴权():
            self._返回JSON({"错误": "未授权：缺少或无效的令牌"}, 401)
            return
        if 路径 == "/api/config":
            self._返回JSON(self.配置加载器.配置缓存)
        elif 路径 == "/api/files":
            参数 = parse_qs(解析结果.query)
            目录 = 参数.get("path", ["./"])[0]
            结果 = self.文件管理器.列目录(目录)
            self._返回JSON(结果)
        elif 路径 == "/api/file-tree":
            参数 = parse_qs(解析结果.query)
            目录 = 参数.get("path", ["./"])[0]
            深度 = int(参数.get("depth", ["3"])[0])
            # 记录最后打开的文件夹（轮盘录音/录屏默认保存位置）
            if 目录 and 目录 != "./":
                网页请求处理器._最后打开的文件夹 = 目录
            结果 = self.文件管理器.目录树(目录, 深度)
            self._返回JSON(结果)
        elif 路径 == "/api/folder-size":
            参数 = parse_qs(解析结果.query)
            目录 = 参数.get("path", ["./"])[0]
            校验 = self.文件管理器._校验权限(目录, "读")
            if not 校验["允许"]:
                self._返回JSON({"成功": False, "错误": 校验["原因"]}, 403)
                return
            完整路径 = self.文件管理器._解析路径(目录)
            if not 完整路径.exists() or not 完整路径.is_dir():
                self._返回JSON({"成功": False, "错误": "目录不存在"}, 404)
                return
            总大小 = 0
            文件数 = 0
            文件夹数 = 0
            try:
                for 根, 目录们, 文件们 in os.walk(完整路径):
                    文件夹数 += len(目录们)
                    文件数 += len(文件们)
                    for f in 文件们:
                        try:
                            fp = os.path.join(根, f)
                            if not os.path.islink(fp):
                                总大小 += os.path.getsize(fp)
                        except OSError:
                            continue
            except OSError:
                pass
            self._返回JSON({"成功": True, "大小": 总大小, "文件数": 文件数, "文件夹数": 文件夹数})
        elif 路径 == "/api/image":
            参数 = parse_qs(解析结果.query)
            图片路径 = 参数.get("path", [""])[0]
            校验 = self.文件管理器._校验权限(图片路径, "读")
            if not 校验["允许"]:
                self._返回JSON({"错误": 校验["原因"]}, 403)
                return
            完整路径 = self.文件管理器._解析路径(图片路径)
            if not 完整路径.exists() or not 完整路径.is_file():
                self._返回JSON({"错误": "文件不存在"}, 404)
                return
            后缀 = 完整路径.suffix.lower()
            if 后缀 == ".tga":
                try:
                    png数据 = _tga转png(完整路径)
                    self.send_response(200)
                    self.send_header("Content-Type", "image/png")
                    self.send_header("Content-Length", len(png数据))
                    self.send_header("Cache-Control", "max-age=3600")
                    self.end_headers()
                    self.wfile.write(png数据)
                except Exception as e:
                    self._返回JSON({"错误": f"TGA转码失败: {str(e)}"}, 500)
                return
            类型映射 = {".png":"image/png",".jpg":"image/jpeg",".jpeg":"image/jpeg",".gif":"image/gif",".webp":"image/webp",".bmp":"image/bmp",".svg":"image/svg+xml"}
            类型 = 类型映射.get(后缀, "application/octet-stream")
            try:
                with open(完整路径, "rb") as f:
                    数据 = f.read()
                self.send_response(200)
                self.send_header("Content-Type", 类型)
                self.send_header("Content-Length", len(数据))
                self.send_header("Cache-Control", "max-age=3600")
                self.end_headers()
                self.wfile.write(数据)
            except Exception as e:
                self._返回JSON({"错误": str(e)}, 500)
        elif 路径 == "/api/audio":
            参数 = parse_qs(解析结果.query)
            音频路径 = 参数.get("path", [""])[0]
            校验 = self.文件管理器._校验权限(音频路径, "读")
            if not 校验["允许"]:
                self._返回JSON({"错误": 校验["原因"]}, 403)
                return
            完整路径 = self.文件管理器._解析路径(音频路径)
            if not 完整路径.exists() or not 完整路径.is_file():
                self._返回JSON({"错误": "文件不存在"}, 404)
                return
            后缀 = 完整路径.suffix.lower()
            类型映射 = {".mp3":"audio/mpeg",".wav":"audio/wav",".ogg":"audio/ogg",".m4a":"audio/mp4",".flac":"audio/flac",".aac":"audio/aac",".opus":"audio/opus",".wma":"audio/x-ms-wma"}
            类型 = 类型映射.get(后缀, "application/octet-stream")
            try:
                文件大小 = 完整路径.stat().st_size
                range_header = self.headers.get("Range")
                if range_header:
                    import re as _re
                    m = _re.match(r'bytes=(\d*)-(\d*)', range_header)
                    if m:
                        start = int(m.group(1)) if m.group(1) else 0
                        end = int(m.group(2)) if m.group(2) else 文件大小 - 1
                        end = min(end, 文件大小 - 1)
                        长度 = end - start + 1
                        self.send_response(206)
                        self.send_header("Content-Type", 类型)
                        self.send_header("Content-Length", 长度)
                        self.send_header("Content-Range", f"bytes {start}-{end}/{文件大小}")
                        self.send_header("Accept-Ranges", "bytes")
                        self.send_header("Cache-Control", "max-age=3600")
                        self.end_headers()
                        with open(完整路径, "rb") as f:
                            f.seek(start)
                            剩余 = 长度
                            while 剩余 > 0:
                                块 = f.read(min(65536, 剩余))
                                if not 块:
                                    break
                                self.wfile.write(块)
                                剩余 -= len(块)
                        return
                with open(完整路径, "rb") as f:
                    数据 = f.read()
                self.send_response(200)
                self.send_header("Content-Type", 类型)
                self.send_header("Content-Length", len(数据))
                self.send_header("Accept-Ranges", "bytes")
                self.send_header("Cache-Control", "max-age=3600")
                self.end_headers()
                self.wfile.write(数据)
            except Exception as e:
                self._返回JSON({"错误": str(e)}, 500)
        elif 路径 == "/api/music-proxy":
            """音乐流代理 — 接收BV号，实时获取B站音频流并转发给前端"""
            参数 = parse_qs(解析结果.query)
            bvid = 参数.get("bvid", [""])[0]
            if not bvid:
                self._返回JSON({"错误": "缺少bvid参数"}, 400)
                return
            try:
                import ssl
                ctx = ssl.create_default_context()
                ctx.check_hostname = False
                ctx.verify_mode = ssl.CERT_NONE

                # 1. 访问B站首页获取cookie
                cj = __import__('http.cookiejar', fromlist=['CookieJar']).CookieJar()
                _opener = urllib.request.build_opener(
                    urllib.request.HTTPCookieProcessor(cj),
                    urllib.request.HTTPSHandler(context=ctx)
                )
                _opener.addheaders = [
                    ("User-Agent", "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"),
                    ("Referer", "https://www.bilibili.com"),
                    ("Accept", "application/json, text/plain, */*"),
                ]
                try:
                    _opener.open("https://www.bilibili.com", timeout=5)
                except Exception:
                    pass

                # 2. 获取cid
                info_url = f"https://api.bilibili.com/x/web-interface/view?bvid={bvid}"
                info_resp = _opener.open(info_url, timeout=10)
                info = json.loads(info_resp.read().decode("utf-8"))
                info_resp.close()
                if info.get("code") != 0:
                    self._返回JSON({"错误": f"获取视频信息失败: {info.get('message', '')}"}, 500)
                    return
                cid = info["data"]["cid"]

                # 3. 获取音频流
                play_url = f"https://api.bilibili.com/x/player/playurl?bvid={bvid}&cid={cid}&fnval=16&qn=0"
                play_resp = _opener.open(play_url, timeout=10)
                play_data = json.loads(play_resp.read().decode("utf-8"))
                play_resp.close()

                audio_list = play_data.get("data", {}).get("dash", {}).get("audio", [])
                if not audio_list:
                    self._返回JSON({"错误": "无音频流"}, 500)
                    return

                audio_url = audio_list[0].get("baseUrl") or audio_list[0].get("base_url", "")

                # 4. 代理转发音频流
                audio_req = urllib.request.Request(audio_url, headers={
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
                    "Referer": "https://www.bilibili.com",
                    "Accept": "*/*",
                    "Origin": "https://www.bilibili.com",
                })
                # 转发Range请求
                range_header = self.headers.get("Range")
                if range_header:
                    audio_req.add_header("Range", range_header)

                audio_resp = _opener.open(audio_req, timeout=30)
                原始类型 = audio_resp.headers.get("Content-Type", "")
                # B站返回application/octet-stream，浏览器audio标签不认，强制改成audio/mp4
                if "octet-stream" in 原始类型 or not 原始类型:
                    类型 = "audio/mp4"
                else:
                    类型 = 原始类型
                长度 = audio_resp.headers.get("Content-Length", "")
                content_range = audio_resp.headers.get("Content-Range", "")
                status = audio_resp.status if hasattr(audio_resp, 'status') else 200

                self.send_response(status)
                self.send_header("Content-Type", 类型)
                if 长度:
                    self.send_header("Content-Length", 长度)
                if content_range:
                    self.send_header("Content-Range", content_range)
                self.send_header("Access-Control-Allow-Origin", "*")
                self.send_header("Accept-Ranges", "bytes")
                self.end_headers()

                while True:
                    数据 = audio_resp.read(65536)
                    if not 数据:
                        break
                    try:
                        self.wfile.write(数据)
                    except (BrokenPipeError, ConnectionResetError, ConnectionAbortedError):
                        break
                audio_resp.close()
            except Exception as e:
                self._返回JSON({"错误": str(e)}, 500)
        elif 路径 == "/api/music-download":
            """按需下载B站音频→转MP3→返回本地文件路径"""
            参数 = parse_qs(解析结果.query)
            bvid = 参数.get("bvid", [""])[0]
            歌名 = 参数.get("name", ["未知"])[0]
            if not bvid:
                self._返回JSON({"错误": "缺少bvid参数"}, 400)
                return
            try:
                from 操作.音乐 import _下载并转换, _添加到音乐库, _加载音乐库
                # 先查库，已下载过的不重复
                库 = _加载音乐库()
                for s in 库.get("歌曲列表", []):
                    if s.get("bvid") == bvid:
                        路径 = s.get("路径", "")
                        if 路径 and os.path.exists(路径):
                            self._返回JSON({"成功": True, "文件路径": 路径, "歌名": s.get("歌名", 歌名), "已缓存": True})
                            return
                # 下载
                import re as _re
                干净名 = _re.sub(r'[<>:"/\\|?*]', '', 歌名[:30])
                结果 = _下载并转换(bvid, 干净名)
                if 结果:
                    文件路径, 原始标题, 时长秒 = 结果
                    _添加到音乐库(文件路径, 歌名, "B站", bvid, 时长秒)
                    self._返回JSON({"成功": True, "文件路径": 文件路径, "歌名": 歌名})
                else:
                    self._返回JSON({"错误": "下载失败"}, 500)
            except Exception as e:
                self._返回JSON({"错误": str(e)}, 500)
        elif 路径 == "/api/download-status":
            """查询后台下载进度"""
            try:
                from 操作.多线程下载 import 多线程下载
                with 多线程下载._下载进度锁:
                    进度表 = dict(多线程下载._下载进度表)
                self._返回JSON({"成功": True, "下载列表": 进度表})
            except Exception as e:
                self._返回JSON({"错误": str(e)}, 500)
        elif 路径 == "/api/video":
            参数 = parse_qs(解析结果.query)
            视频路径 = 参数.get("path", [""])[0]
            校验 = self.文件管理器._校验权限(视频路径, "读")
            if not 校验["允许"]:
                self._返回JSON({"错误": 校验["原因"]}, 403)
                return
            完整路径 = self.文件管理器._解析路径(视频路径)
            if not 完整路径.exists() or not 完整路径.is_file():
                self._返回JSON({"错误": "文件不存在"}, 404)
                return
            后缀 = 完整路径.suffix.lower()
            需转码 = 后缀 in [".avi", ".wmv"]
            if 需转码:
                self.send_response(200)
                self.send_header("Content-Type", "video/mp4")
                self.send_header("Cache-Control", "no-cache")
                self.end_headers()
                proc = subprocess.Popen(
                    ["ffmpeg", "-i", str(完整路径), "-c:v", "libx264", "-preset", "ultrafast",
                     "-c:a", "aac", "-b:a", "128k", "-movflags", "frag_keyframe+empty_moov",
                     "-f", "mp4", "-threads", "2", "pipe:1"],
                    stdout=subprocess.PIPE, stderr=subprocess.DEVNULL
                )
                try:
                    while True:
                        块 = proc.stdout.read(65536)
                        if not 块:
                            break
                        self.wfile.write(块)
                except (BrokenPipeError, ConnectionResetError):
                    pass
                finally:
                    proc.stdout.close()
                    proc.kill()
                return
            类型映射 = {".mp4":"video/mp4",".webm":"video/webm",".mkv":"video/x-matroska",".avi":"video/x-msvideo",".wmv":"video/x-ms-wmv",".mov":"video/quicktime",".flv":"video/x-flv",".ts":"video/mp2t"}
            类型 = 类型映射.get(后缀, "video/mp4")
            文件大小 = 完整路径.stat().st_size
            range_header = self.headers.get("Range")
            if range_header:
                import re
                m = re.match(r"bytes=(\d+)-(\d*)", range_header)
                if m:
                    start = int(m.group(1))
                    end = int(m.group(2)) if m.group(2) else 文件大小 - 1
                    end = min(end, 文件大小 - 1)
                    长度 = end - start + 1
                    self.send_response(206)
                    self.send_header("Content-Type", 类型)
                    self.send_header("Content-Length", 长度)
                    self.send_header("Content-Range", f"bytes {start}-{end}/{文件大小}")
                    self.send_header("Accept-Ranges", "bytes")
                    self.send_header("Cache-Control", "max-age=3600")
                    self.end_headers()
                    with open(完整路径, "rb") as f:
                        f.seek(start)
                        剩余 = 长度
                        while 剩余 > 0:
                            块 = f.read(min(65536, 剩余))
                            if not 块:
                                break
                            try:
                                self.wfile.write(块)
                            except (ConnectionAbortedError, ConnectionResetError, BrokenPipeError):
                                return
                            剩余 -= len(块)
                    return
            self.send_response(200)
            self.send_header("Content-Type", 类型)
            self.send_header("Content-Length", 文件大小)
            self.send_header("Accept-Ranges", "bytes")
            self.send_header("Cache-Control", "max-age=3600")
            self.end_headers()
            try:
                with open(完整路径, "rb") as f:
                    while True:
                        块 = f.read(65536)
                        if not 块:
                            break
                        self.wfile.write(块)
            except Exception as e:
                self._返回JSON({"错误": str(e)}, 500)
        elif 路径 == "/api/file-content":
            参数 = parse_qs(解析结果.query)
            文件路径 = 参数.get("path", [""])[0]
            校验 = self.文件管理器._校验权限(文件路径, "读")
            if not 校验["允许"]:
                self._返回JSON({"错误": 校验["原因"]}, 403)
                return
            完整路径 = self.文件管理器._解析路径(文件路径)
            if not 完整路径.exists() or not 完整路径.is_file():
                self._返回JSON({"错误": "文件不存在"}, 404)
                return
            文件大小 = 完整路径.stat().st_size
            self.send_response(200)
            self.send_header("Content-Type", "application/octet-stream")
            self.send_header("Content-Length", 文件大小)
            self.send_header("Cache-Control", "no-cache")
            self.end_headers()
            try:
                with open(完整路径, "rb") as f:
                    while True:
                        块 = f.read(65536)
                        if not 块:
                            break
                        self.wfile.write(块)
            except Exception as e:
                self._返回JSON({"错误": str(e)}, 500)
        elif 路径 == "/api/doc-content":
            参数 = parse_qs(解析结果.query)
            文件路径 = 参数.get("path", [""])[0]
            校验 = self.文件管理器._校验权限(文件路径, "读")
            if not 校验["允许"]:
                self._返回JSON({"错误": 校验["原因"]}, 403)
                return
            完整路径 = self.文件管理器._解析路径(文件路径)
            if not 完整路径.exists() or not 完整路径.is_file():
                self._返回JSON({"错误": "文件不存在"}, 404)
                return
            try:
                html = _提取doc文本(str(完整路径))
                self._返回JSON({"成功": True, "html": html})
            except Exception as e:
                self._返回JSON({"错误": str(e)}, 500)
        elif 路径 == "/api/docx-content":
            参数 = parse_qs(解析结果.query)
            文件路径 = 参数.get("path", [""])[0]
            校验 = self.文件管理器._校验权限(文件路径, "读")
            if not 校验["允许"]:
                self._返回JSON({"错误": 校验["原因"]}, 403)
                return
            完整路径 = self.文件管理器._解析路径(文件路径)
            if not 完整路径.exists() or not 完整路径.is_file():
                self._返回JSON({"错误": "文件不存在"}, 404)
                return
            try:
                html = _提取docx文本(str(完整路径))
                self._返回JSON({"成功": True, "html": html})
            except Exception as e:
                self._返回JSON({"错误": str(e)}, 500)
        elif 路径 == "/api/audit-log":
            self._返回JSON({"日志": self.文件管理器.获取审计日志()})
        elif 路径 == "/api/pending":
            self._返回JSON({"待确认": self.文件管理器.获取待确认()})
        elif 路径 == "/api/modules":
            self._返回JSON({"模块": list(self.模块注册.keys()) if self.模块注册 else []})
        elif 路径 == "/api/dev-reload-sse":
            self._处理开发热重载SSE()
        elif 路径 == "/api/status":
            对话状态 = {}
            if self.模块注册 and "对话" in self.模块注册:
                对话状态 = self.模块注册["对话"].获取状态()
            当前模型 = "默认"
            if self.模型直连器:
                当前模型 = self.模型直连器.当前模型名 or "默认"
            # 从系统配置读取版本号（唯一版本源）
            系统版本 = "未知"
            try:
                if self.配置加载器:
                    系统配置 = self.配置加载器.获取配置("系统配置")
                    系统版本 = 系统配置.get("版本", "未知")
            except Exception:
                pass
            self._返回JSON({
                "状态": "运行中", "版本": 系统版本,
                "对话": 对话状态,
                "当前模型": 当前模型,
                "操作数": len(self.操作注册中心.列出所有操作()) if self.操作注册中心 else 0
            })
        elif 路径 == "/api/tts-status":
            self._返回JSON({
                "正在播放": 网页请求处理器._tts主界面状态["播放中"],
                "轮盘播放": 网页请求处理器._tts轮盘状态["播放中"]
            })
        elif 路径 == "/api/actions":
            if self.操作注册中心:
                self._返回JSON({"操作": self.操作注册中心.获取操作JSON描述()})
            else:
                self._返回JSON({"操作": []})
        elif 路径 == "/api/token-stats":
            """获取Token使用统计"""
            from 模型直连器 import 模型直连器类
            self._返回JSON({"成功": True, "统计": 模型直连器类.获取Token统计()})
        elif 路径 == "/api/cache-stats":
            """获取LLM缓存统计"""
            from 模型直连器 import 模型直连器类
            self._返回JSON({"成功": True, "统计": 模型直连器类.获取缓存统计()})
        elif 路径 == "/api/cache-clear":
            """清空LLM缓存"""
            from 模型直连器 import 模型直连器类
            模型直连器类.清空缓存()
            self._返回JSON({"成功": True, "消息": "缓存已清空"})
        elif 路径 == "/api/conv-search":
            """搜索对话内容（SQLite全文搜索）"""
            参数 = parse_qs(解析结果.query)
            关键词 = 参数.get("q", [""])[0]
            if 关键词:
                try:
                    from 存储引擎 import 获取存储引擎
                    引擎 = 获取存储引擎()
                    结果 = 引擎.搜索对话(关键词)
                    self._返回JSON({"成功": True, "结果": 结果})
                except Exception as e:
                    self._返回JSON({"成功": False, "错误": str(e)})
            else:
                self._返回JSON({"成功": False, "错误": "缺少搜索关键词"})
        elif 路径 == "/api/reasoning-stream":
            """轮询获取推理流（实时显示AI操作过程）"""
            参数 = parse_qs(解析结果.query)
            上次索引 = int(参数.get("index", ["0"])[0])
            if self.模块注册 and "对话" in self.模块注册:
                try:
                    结果 = self.模块注册["对话"].获取推理流(上次索引)
                    self._返回JSON(结果)
                except Exception as e:
                    self._返回JSON({"成功": False, "错误": str(e)})
            else:
                self._返回JSON({"成功": False, "错误": "对话模块未就绪"})
        elif 路径 == "/api/wf-tasks":
            """GET: 获取工作流定时任务列表"""
            调度器 = getattr(网页请求处理器, '_定时任务调度器', None)
            if not 调度器:
                self._返回JSON({"成功": False, "错误": "定时任务调度器未就绪"})
                return
            if 调度器:
                self._返回JSON({"成功": True, "任务列表": 调度器.获取工作流任务列表()})
            else:
                self._返回JSON({"成功": False, "错误": "定时任务调度器未就绪"})
        elif 路径 == "/api/history":
            if self.模块注册 and "对话" in self.模块注册:
                历史 = self.模块注册["对话"].获取历史()
                self._返回JSON({"历史": 历史})
            else:
                self._返回JSON({"历史": []})
        elif 路径 == "/api/tasks":
            if self.模块注册 and "任务" in self.模块注册:
                self._返回JSON(self.模块注册["任务"]._列出任务())
            else:
                self._返回JSON({"任务列表": []})
        elif 路径 == "/api/models":
            # 返回可用模型列表
            模型列表 = []
            if self.模型直连器:
                模型列表 = self.模型直连器.获取模型列表()
            if not 模型列表:
                模型列表.append({"名称": "默认模型", "当前": True})
            self._返回JSON({"模型": 模型列表})
        elif 路径 == "/api/stock-panel":
            """股票盘面：指数+涨幅榜+跌幅榜+市场总览（缓存）"""
            参数 = parse_qs(解析结果.query)
            页码 = int(参数.get("page", ["1"])[0])
            排序字段 = 参数.get("sort", ["f3"])[0]  # 默认涨幅
            排序方向 = 参数.get("order", ["desc"])[0]  # desc降序 asc升序
            from 股票缓存 import 获取股票缓存
            缓存 = 获取股票缓存()
            结果 = 缓存.读取或请求(f"panel_{页码}_{排序字段}_{排序方向}", "panel", lambda: self._获取股票盘面(页码, 排序字段, 排序方向))
            self._返回JSON(结果)
        elif 路径 == "/api/stock-kline":
            """股票K线数据（支持日K/周K/月K，缓存+增量更新）"""
            参数 = parse_qs(解析结果.query)
            代码 = 参数.get("code", [""])[0]
            周期 = 参数.get("period", ["daily"])[0]
            from 股票缓存 import 获取股票缓存
            缓存 = 获取股票缓存()
            结果 = 缓存.读取或请求(f"kline_{代码}_{周期}", "kline", lambda: self._获取股票K线(代码, 周期))
            self._返回JSON(结果)
        elif 路径 == "/api/stock-minute":
            """股票分时数据（缓存）"""
            参数 = parse_qs(解析结果.query)
            代码 = 参数.get("code", [""])[0]
            from 股票缓存 import 获取股票缓存
            缓存 = 获取股票缓存()
            结果 = 缓存.读取或请求(f"minute_{代码}", "minute", lambda: self._获取股票分时(代码))
            self._返回JSON(结果)
        elif 路径 == "/api/stock-search":
            """搜索股票（代码/名称模糊匹配，缓存）"""
            参数 = parse_qs(解析结果.query)
            关键词 = 参数.get("q", [""])[0]
            from 股票缓存 import 获取股票缓存
            缓存 = 获取股票缓存()
            结果 = 缓存.读取或请求(f"search_{关键词}", "search", lambda: self._搜索股票(关键词))
            self._返回JSON(结果)
        elif 路径 == "/api/stock-detail":
            """个股详情（PE/PB/市值/换手率等，缓存）"""
            参数 = parse_qs(解析结果.query)
            代码 = 参数.get("code", [""])[0]
            from 股票缓存 import 获取股票缓存
            缓存 = 获取股票缓存()
            结果 = 缓存.读取或请求(f"detail_{代码}", "detail", lambda: self._获取股票详情(代码))
            self._返回JSON(结果)
        elif 路径 == "/api/stock-sectors":
            """板块行情（行业+概念，缓存）"""
            from 股票缓存 import 获取股票缓存
            缓存 = 获取股票缓存()
            结果 = 缓存.读取或请求("sectors", "sectors", lambda: self._获取板块行情())
            self._返回JSON(结果)
        elif 路径 == "/api/stock-capital-flow":
            """个股资金流向明细（缓存）"""
            参数 = parse_qs(解析结果.query)
            代码 = 参数.get("code", [""])[0]
            from 股票缓存 import 获取股票缓存
            缓存 = 获取股票缓存()
            结果 = 缓存.读取或请求(f"flow_{代码}", "flow", lambda: self._获取资金流向(代码))
            self._返回JSON(结果)
        elif 路径 == "/api/stock-batch":
            """批量查询自选股行情"""
            参数 = parse_qs(解析结果.query)
            代码列表 = 参数.get("codes", [""])[0].split(",")
            代码列表 = [c.strip() for c in 代码列表 if c.strip()]
            self._返回JSON(self._批量查询行情(代码列表))
        elif 路径 == "/api/stock-export":
            """导出K线数据为CSV"""
            参数 = parse_qs(解析结果.query)
            代码 = 参数.get("code", [""])[0]
            周期 = 参数.get("period", ["daily"])[0]
            self._返回CSV(self._导出K线CSV(代码, 周期), f"{代码}_{周期}.csv")
        elif 路径 == "/api/stock-cache-stats":
            """股票缓存统计"""
            from 股票缓存 import 获取股票缓存
            缓存 = 获取股票缓存()
            self._返回JSON({"成功": True, "统计": 缓存.获取缓存统计()})
        elif 路径 == "/api/stock-cache-clear":
            """清空股票缓存"""
            from 股票缓存 import 获取股票缓存
            缓存 = 获取股票缓存()
            缓存.清空缓存()
            self._返回JSON({"成功": True, "消息": "股票缓存已清空"})
        elif 路径 == "/api/stock-bulk-start":
            """启动全量K线+财务数据下载"""
            参数 = parse_qs(解析结果.query)
            周期 = 参数.get("period", ["daily"])[0]
            增量 = 参数.get("incremental", ["1"])[0] != "0"
            含财务 = 参数.get("finance", ["1"])[0] != "0"
            强制刷新 = 参数.get("refresh", ["0"])[0] == "1"
            print(f"[股票下载] 启动请求: 周期={周期} 增量={增量} 财务={含财务} 刷新={强制刷新}")
            from 股票缓存 import 获取下载引擎
            引擎 = 获取下载引擎()
            结果 = 引擎.启动下载(周期, 增量, 含财务, 强制刷新列表=强制刷新)
            print(f"[股票下载] 启动结果: {结果}")
            self._返回JSON(结果)
        elif 路径 == "/api/stock-bulk-progress":
            """查询全量下载进度"""
            from 股票缓存 import 获取下载引擎, 获取股票缓存
            引擎 = 获取下载引擎()
            缓存 = 获取股票缓存()
            进度 = 引擎.获取进度()
            统计 = 缓存.取本地K线统计()
            财务统计 = 缓存.取财务数据统计()
            self._返回JSON({"成功": True, "进度": 进度, "本地统计": 统计, "财务统计": 财务统计})
        elif 路径 == "/api/stock-bulk-stop":
            """停止全量下载"""
            from 股票缓存 import 获取下载引擎
            引擎 = 获取下载引擎()
            引擎.停止()
            self._返回JSON({"成功": True, "消息": "正在停止..."})
        elif 路径 == "/api/stock-finance":
            """查询单只股票本地财务数据"""
            参数 = parse_qs(解析结果.query)
            代码 = 参数.get("code", [""])[0]
            from 股票缓存 import 获取股票缓存
            缓存 = 获取股票缓存()
            数据 = 缓存.取财务数据(代码)
            self._返回JSON({"成功": True, "数据": 数据})
        elif 路径 == "/api/drives":
            # 返回可用磁盘驱动器列表 + 用户文件夹快捷方式
            驱动器列表 = []
            用户目录 = os.path.expanduser("~")
            快捷方式列表 = [
                {"盘符": "桌面", "路径": os.path.join(用户目录, "Desktop"), "标签": "桌面", "图标": "🖥️", "类型": "文件夹"},
                {"盘符": "文档", "路径": os.path.join(用户目录, "Documents"), "标签": "文档", "图标": "📄", "类型": "文件夹"},
                {"盘符": "下载", "路径": os.path.join(用户目录, "Downloads"), "标签": "下载", "图标": "📥", "类型": "文件夹"},
                {"盘符": "图片", "路径": os.path.join(用户目录, "Pictures"), "标签": "图片", "图标": "🖼️", "类型": "文件夹"},
            ]
            for 快捷方式 in 快捷方式列表:
                if os.path.exists(快捷方式["路径"]):
                    驱动器列表.append(快捷方式)
            if sys.platform == "win32":
                for 盘符 in string.ascii_uppercase:
                    驱动器路径 = f"{盘符}:\\"
                    if os.path.exists(驱动器路径):
                        try:
                            使用 = shutil.disk_usage(驱动器路径)
                            总大小GB = round(使用.total / (1024**3), 1)
                            已用GB = round(使用.used / (1024**3), 1)
                            可用GB = round(使用.free / (1024**3), 1)
                            驱动器列表.append({
                                "盘符": f"{盘符}:",
                                "路径": 驱动器路径,
                                "标签": f"本地磁盘 {盘符}",
                                "图标": "💾",
                                "类型": "磁盘",
                                "总大小GB": 总大小GB,
                                "已用GB": 已用GB,
                                "可用GB": 可用GB,
                            })
                        except:
                            驱动器列表.append({"盘符": f"{盘符}:", "路径": 驱动器路径, "标签": f"本地磁盘 {盘符}", "图标": "💾", "类型": "磁盘"})
            else:
                驱动器列表.append({"盘符": "/", "路径": "/", "标签": "根目录", "图标": "💾", "类型": "磁盘"})
            self._返回JSON({"驱动器": 驱动器列表})
        elif 路径 == "/api/folder-dialog":
            # Windows原生文件夹选择对话框
            选中路径 = self._打开文件夹选择对话框()
            self._返回JSON({"路径": 选中路径})
        elif 路径 == "/api/conversations":
            if self.模块注册 and "对话" in self.模块注册:
                self._返回JSON({"对话列表": self.模块注册["对话"].获取对话列表(), "当前ID": self.模块注册["对话"].当前对话ID})
            else:
                self._返回JSON({"对话列表": [], "当前ID": None})
        elif 路径 == "/api/current-plan":
            if self.模块注册 and "对话" in self.模块注册:
                self._返回JSON({"计划": self.模块注册["对话"].当前计划, "工作模式": self.模块注册["对话"].工作模式})
            else:
                self._返回JSON({"计划": None})
        elif 路径 == "/api/checkpoint-info":
            """获取检查点信息"""
            if self.模块注册 and "对话" in self.模块注册:
                self._返回JSON({"有检查点": self.模块注册["对话"].有检查点()})
            else:
                self._返回JSON({"有检查点": False})
        elif 路径 == "/api/health":
            """系统健康自检"""
            启动器 = getattr(self, '_启动器实例', None)
            if 启动器:
                self._返回JSON(启动器.自检())
            else:
                self._返回JSON({"状态": "未知", "错误": "启动器实例未绑定"})
        elif 路径 == "/api/action-stats":
            """操作调用统计"""
            if self.操作注册中心:
                self._返回JSON(self.操作注册中心.获取操作统计())
            else:
                self._返回JSON({"错误": "操作注册中心未就绪"})
        elif 路径 == "/api/engine-diff":
            """对比工作引擎与主引擎文件差异"""
            try:
                结果 = self._引擎差异分析()
                self._返回JSON(结果)
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/engine-backups":
            """列出可用备份"""
            try:
                项目根 = self.配置加载器.项目根目录
                备份目录 = 项目根 / "引擎管理" / "备份"
                备份列表 = []
                if 备份目录.exists():
                    for d in sorted(备份目录.iterdir(), reverse=True):
                        if d.is_dir():
                            文件数 = sum(1 for _ in d.rglob("*") if _.is_file())
                            备份列表.append({"名称": d.name, "文件数": 文件数})
                self._返回JSON({"成功": True, "备份列表": 备份列表})
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/evolution-status":
            """获取进化引擎状态"""
            启动器 = getattr(self, '_启动器实例', None)
            if 启动器 and hasattr(启动器, '进化引擎'):
                self._返回JSON({"成功": True, "状态": 启动器.进化引擎.获取状态()})
            else:
                self._返回JSON({"成功": False, "错误": "进化引擎未启动"})
        elif 路径 == "/api/evolution-records":
            """获取进化历史记录"""
            try:
                参数 = parse_qs(解析结果.query)
                项目根 = self.配置加载器.项目根目录
                记录目录 = 项目根 / "隐私区" / "我的工作引擎" / "进化记录"
                记录列表 = []
                关键词 = 参数.get("关键词", [""])[0]
                if 记录目录.exists():
                    for d in sorted(记录目录.iterdir(), reverse=True):
                        清单文件 = d / "修改清单.json"
                        if not 清单文件.exists():
                            continue
                        try:
                            with open(清单文件, "r", encoding="utf-8") as f:
                                记录 = json.load(f)
                            文件名 = 记录.get("文件", "")
                            if 关键词 and 关键词 not in 文件名:
                                continue
                            修改说明 = ""
                            问题描述 = ""
                            if 记录.get("修改详情"):
                                详情 = 记录["修改详情"][0] if isinstance(记录["修改详情"], list) else 记录["修改详情"]
                                修改说明 = 详情.get("说明", "") if isinstance(详情, dict) else str(详情)
                            if 记录.get("问题列表"):
                                问题项 = 记录["问题列表"][0] if isinstance(记录["问题列表"], list) else 记录["问题列表"]
                                问题描述 = 问题项.get("问题描述", "") if isinstance(问题项, dict) else str(问题项)
                            记录列表.append({
                                "文件": 文件名,
                                "时间": 记录.get("时间", ""),
                                "轮次": 记录.get("轮次", 0),
                                "状态": "审查通过" if 记录.get("审查意见") else "待审查",
                                "风险等级": 记录.get("风险", "低"),
                                "审查意见": 记录.get("审查意见", ""),
                                "修改说明": 修改说明,
                                "问题描述": 问题描述,
                                "原始代码": "",
                                "完整代码": ""
                            })
                        except Exception:
                            continue
                self._返回JSON({"成功": True, "记录": 记录列表})
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/wheel-config":
            """GET：读取快速呼出配置"""
            try:
                配置路径 = self.配置加载器.项目根目录 / "公共区" / "配置" / "快速呼出配置.json"
                with open(配置路径, "r", encoding="utf-8") as f:
                    配置 = json.load(f)
                self._返回JSON({"成功": True, "配置": 配置})
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/voice-status":
            """检查语音输入引擎状态"""
            系统配置 = self.配置加载器.配置缓存.get("系统配置", {})
            语音配置 = 系统配置.get("语音输入", {})
            引擎 = 语音配置.get("引擎", "浏览器")
            模型 = 语音配置.get("本地模型", "paraformer-zh-int8")
            已安装 = False
            try:
                import sherpa_onnx
                已安装 = True
            except ImportError:
                pass
            # 检查模型文件是否存在（纯英文路径，sherpa-onnx不支持中文路径）
            流式目录 = _获取英文模型目录() / "paraformer-streaming"
            流式存在 = (流式目录 / "encoder.int8.onnx").exists() and (流式目录 / "decoder.int8.onnx").exists() and (流式目录 / "tokens.txt").exists()
            self._返回JSON({"成功": True, "已安装": 已安装, "模型存在": 流式存在, "引擎": 引擎, "模型": 模型})
        elif 路径 == "/api/employee-list":
            """获取员工列表"""
            if self.模块注册 and "员工管理" in self.模块注册:
                self._返回JSON(self.模块注册["员工管理"].获取员工列表())
            else:
                self._返回JSON({"成功": False, "错误": "员工管理模块未加载"})
        elif 路径 == "/api/employee-tree":
            """获取员工树形结构"""
            if self.模块注册 and "员工管理" in self.模块注册:
                self._返回JSON(self.模块注册["员工管理"].获取员工树())
            else:
                self._返回JSON({"成功": False, "错误": "员工管理模块未加载"})
        elif 路径 == "/api/employee-notify":
            """获取员工定时提醒"""
            提醒列表 = []
            for 姓名, 信息 in 网页请求处理器._员工提醒队列.items():
                if 信息["待发送"]:
                    提醒列表.append({
                        "姓名": 姓名,
                        "头像": 信息["头像"],
                        "消息": 信息["待发送"],
                        "时间": datetime.now().strftime("%H:%M"),
                        "弹窗时长秒": 信息.get("弹窗时长秒", 30),
                    })
                    信息["待发送"] = None
            self._返回JSON({"成功": True, "数据": 提醒列表})
        elif 路径.startswith("/api/employee-history"):
            """获取员工对话历史"""
            参数 = parse_qs(解析结果.query)
            姓名 = 参数.get("姓名", [""])[0]
            历史 = 网页请求处理器._员工对话历史.get(姓名, [])
            # 内存为空时从记忆文件加载
            if not 历史 and self.模块注册 and "员工管理" in self.模块注册:
                运行时结果 = self.模块注册["员工管理"].获取运行时配置(姓名)
                if 运行时结果.get("成功"):
                    记忆路径 = 运行时结果["数据"].get("记忆路径")
                    if 记忆路径:
                        try:
                            import json as _json
                            from pathlib import Path as _P
                            绝对路径 = self.配置加载器.项目根目录 / 记忆路径.lstrip("./")
                            if 绝对路径.exists():
                                记忆数据 = _json.loads(绝对路径.read_text(encoding="utf-8"))
                                for m in 记忆数据:
                                    if m.get("用户"):
                                        历史.append({"role": "user", "content": m["用户"]})
                                    if m.get("助手"):
                                        历史.append({"role": "assistant", "content": m["助手"]})
                                网页请求处理器._员工对话历史[姓名] = 历史
                        except Exception:
                            pass
            self._返回JSON({"成功": True, "数据": 历史})
        elif 路径 == "/api/employee-current":
            """获取当前活跃员工"""
            if self.模块注册 and "员工管理" in self.模块注册:
                self._返回JSON(self.模块注册["员工管理"].获取当前员工())
            else:
                self._返回JSON({"成功": False, "错误": "员工管理模块未加载"})
        elif 路径.startswith("/api/employee-config?") or 路径 == "/api/employee-config":
            """获取员工运行时配置"""
            if self.模块注册 and "员工管理" in self.模块注册:
                参数 = parse_qs(解析结果.query)
                姓名 = 参数.get("姓名", [""])[0]
                self._返回JSON(self.模块注册["员工管理"].获取运行时配置(姓名))
            else:
                self._返回JSON({"成功": False, "错误": "员工管理模块未加载"})
        elif 路径 == "/api/tts-voices":
            """获取Kokoro TTS可用说话人列表"""
            self._返回JSON({"成功": True, "数据": _KOKORO说话人})
        elif 路径 == "/api/tts-config":
            """获取TTS输出配置"""
            try:
                系统配置 = self.配置加载器.配置缓存.get("系统配置", {})
                语音输出 = 系统配置.get("语音输出", {})
                已安装 = False
                模型存在 = False
                try:
                    import sherpa_onnx
                    已安装 = True
                except ImportError:
                    pass
                模型存在 = _检查KokoroTTS模型存在()
                self._返回JSON({"成功": True, "配置": 语音输出, "已安装": 已安装, "模型存在": 模型存在})
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/tts-install-status":
            """查询TTS模型安装进度"""
            self._返回JSON({"成功": True, "状态": 网页请求处理器._tts安装状态})
        else:
            print(f"  ❌ 未知GET API: {路径}")
            self._返回JSON({"错误": "未知API: " + 路径}, 404)

    # ============ WebSocket 流式语音识别 ============
    def _处理WebSocket语音(self):
        """处理WebSocket升级请求，建立流式语音识别连接"""
        import hashlib, base64 as _b64
        ws_key = self.headers.get("Sec-WebSocket-Key", "")
        if not ws_key:
            self.send_response(400)
            self.end_headers()
            return
        magic = "258EAFA5-E914-47DA-95CA-C5AB0DC85B11"
        accept = _b64.b64encode(hashlib.sha1((ws_key + magic).encode()).digest()).decode()
        self.send_response(101)
        self.send_header("Upgrade", "websocket")
        self.send_header("Connection", "Upgrade")
        self.send_header("Sec-WebSocket-Accept", accept)
        self.end_headers()
        try:
            import sherpa_onnx
        except ImportError:
            self._ws发送JSON({"类型": "错误", "错误": "sherpa-onnx未安装"})
            self._ws关闭(1000)
            return
        if not hasattr(网页请求处理器, '_sherpa流式识别器') or 网页请求处理器._sherpa流式识别器 is None:
            模型目录 = _获取英文模型目录() / "paraformer-streaming"
            if not (模型目录 / "encoder.int8.onnx").exists():
                self._ws发送JSON({"类型": "错误", "错误": "流式模型未安装，请在设置→语音中安装"})
                self._ws关闭(1000)
                return
            try:
                网页请求处理器._sherpa流式识别器 = sherpa_onnx.OnlineRecognizer.from_paraformer(
                    encoder=str(模型目录 / "encoder.int8.onnx"),
                    decoder=str(模型目录 / "decoder.int8.onnx"),
                    tokens=str(模型目录 / "tokens.txt"),
                    num_threads=2,
                    enable_endpoint_detection=True,
                    rule1_min_trailing_silence=2.4,
                    rule2_min_trailing_silence=1.2,
                    rule3_min_utterance_length=20.0,
                )
            except Exception as e:
                self._ws发送JSON({"类型": "错误", "错误": f"流式模型加载失败: {e}"})
                self._ws关闭(1000)
                return
        识别器 = 网页请求处理器._sherpa流式识别器
        流 = 识别器.create_stream()
        已确认文字 = ""
        while True:
            try:
                帧 = self._ws读取帧()
                if 帧 is None:
                    break
                opcode, payload = 帧
                if opcode == 0x8:
                    break
                if opcode == 0x1:
                    消息 = payload.decode("utf-8", errors="replace")
                    if 消息 == "end":
                        识别器.decode_stream(流)
                        文字 = 识别器.get_result(流).strip()
                        if 文字:
                            文字 += "。"  # 结束时加句号
                            已确认文字 += 文字
                            self._ws发送JSON({"类型": "增量", "文字": 文字})
                        self._ws发送JSON({"类型": "最终", "文字": 已确认文字})
                        break
                    continue
                if opcode == 0x2:
                    import numpy as np
                    samples = np.frombuffer(payload, dtype=np.float32)
                    if len(samples) > 0:
                        流.accept_waveform(16000, samples)
                        if 识别器.is_ready(流):
                            识别器.decode_stream(流)
                            部分文字 = 识别器.get_result(流).strip()
                            if 部分文字:
                                self._ws发送JSON({"类型": "部分", "文字": 部分文字})
                        if 识别器.is_endpoint(流):
                            识别器.reset(流)
                            if 部分文字:
                                # 断句时加句号
                                部分文字 += "。"
                                已确认文字 += 部分文字
                                self._ws发送JSON({"类型": "增量", "文字": 部分文字})
            except (ConnectionAbortedError, ConnectionResetError, BrokenPipeError, OSError):
                break
            except Exception as e:
                print(f"  ⚠️ WebSocket语音错误: {e}")
                break
        self._ws关闭(1000)

    def _ws发送JSON(self, 数据: dict):
        payload = json.dumps(数据, ensure_ascii=False).encode("utf-8")
        self._ws发送帧(0x1, payload)

    def _ws发送帧(self, opcode: int, payload: bytes):
        header = bytearray()
        header.append(0x80 | opcode)
        长度 = len(payload)
        if 长度 < 126:
            header.append(长度)
        elif 长度 < 65536:
            header.append(126)
            header.extend(长度.to_bytes(2, "big"))
        else:
            header.append(127)
            header.extend(长度.to_bytes(8, "big"))
        self.wfile.write(header + payload)
        self.wfile.flush()

    def _ws读取帧(self):
        try:
            头 = self.rfile.read(2)
            if len(头) < 2:
                return None
            b1, b2 = 头[0], 头[1]
            opcode = b1 & 0x0F
            masked = (b2 & 0x80) != 0
            长度 = b2 & 0x7F
            if 长度 == 126:
                长度 = int.from_bytes(self.rfile.read(2), "big")
            elif 长度 == 127:
                长度 = int.from_bytes(self.rfile.read(8), "big")
            mask_key = self.rfile.read(4) if masked else b""
            # 循环读取完整payload（rfile.read可能返回不完整）
            payload = b""
            remaining = 长度
            while remaining > 0:
                chunk = self.rfile.read(remaining)
                if not chunk:
                    break
                payload += chunk
                remaining -= len(chunk)
            if masked and payload:
                payload = bytes(payload[i] ^ mask_key[i % 4] for i in range(len(payload)))
            if opcode == 0x8:
                return (0x8, payload)
            return (opcode, payload)
        except (ConnectionAbortedError, ConnectionResetError, BrokenPipeError, OSError):
            return None

    def _ws关闭(self, code: int = 1000):
        try:
            self._ws发送帧(0x8, code.to_bytes(2, "big"))
        except Exception:
            pass

    def _处理API_POST(self, 路径: str, 数据: dict):
        if not self._检查鉴权():
            self._返回JSON({"错误": "未授权：缺少或无效的令牌"}, 401)
            return
        if 路径 == "/api/restart":
            self._返回JSON({"成功": True, "消息": "重启中..."})
            def _延迟重启():
                time.sleep(0.5)
                try:
                    import subprocess, sys, os
                    # 获取项目根目录
                    项目根 = str(网页请求处理器.配置加载器.项目根目录)
                    python = sys.executable
                    启动脚本 = os.path.join(项目根, "公共区", "内核", "启动器.py")
                    # 设置环境变量标记重启，使新进程不再自动打开浏览器
                    env = os.environ.copy()
                    env["_ZF3D_RESTART"] = "1"
                    # 创建detached新进程
                    if sys.platform == 'win32':
                        subprocess.Popen([python, 启动脚本], cwd=项目根, env=env,
                                       creationflags=subprocess.CREATE_NEW_PROCESS_GROUP | subprocess.DETACHED_PROCESS)
                    else:
                        subprocess.Popen([python, 启动脚本], cwd=项目根, env=env,
                                       start_new_session=True)
                except Exception as e:
                    print(f"⚠️ 重启失败: {e}")
                # 立即强制退出当前进程（不走停止流程，避免端口释放延迟）
                import os as _os
                _os._exit(0)
            threading.Thread(target=_延迟重启, daemon=True).start()
        elif 路径 == "/api/chat":
            self._处理对话(数据)
        elif 路径 == "/api/employee-chat":
            self._处理员工对话(数据)
        elif 路径 == "/api/employee-task":
            self._处理员工任务(数据)
        elif 路径 == "/api/employee-workflow":
            self._处理员工工作流(数据)
        elif 路径 == "/api/wf-save":
            """保存节点图到文件"""
            import os as _os
            目录 = self.配置加载器.项目根目录 / "节点图"
            目录.mkdir(parents=True, exist_ok=True)
            文件名 = 数据.get("文件名", "未命名")
            安全名 = "".join(c for c in 文件名 if c not in '/\\:*?"<>|') or "未命名"
            路径文件 = 目录 / (安全名 + ".json")
            覆盖 = 数据.get("覆盖", False)
            if 路径文件.exists() and not 覆盖:
                self._返回JSON({"成功": False, "已存在": True})
                return
            try:
                路径文件.write_text(json.dumps(数据.get("图", {}), ensure_ascii=False, indent=2), encoding="utf-8")
                self._返回JSON({"成功": True, "路径": str(路径文件)})
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/wf-load":
            """从文件载入节点图"""
            import os as _os
            目录 = self.配置加载器.项目根目录 / "节点图"
            if not 目录.exists():
                目录.mkdir(parents=True, exist_ok=True)
            文件名 = 数据.get("文件名", "")
            if 文件名:
                路径文件 = 目录 / (文件名 + ".json")
                try:
                    图 = json.loads(路径文件.read_text(encoding="utf-8"))
                    self._返回JSON({"成功": True, "图": 图})
                except Exception as e:
                    self._返回JSON({"成功": False, "错误": str(e)})
            else:
                列表 = [f.stem for f in 目录.glob("*.json")]
                self._返回JSON({"成功": True, "列表": 列表})
        elif 路径 == "/api/wf-check-images":
            """检查图片节点中的图片是否还存在，返回有效的图片列表"""
            import os as _os_chk
            图片列表 = 数据.get("图片列表", [])
            有效 = []
            for p in 图片列表:
                if p and _os_chk.exists(p):
                    有效.append(p)
            self._返回JSON({"成功": True, "有效图片": 有效})
        elif 路径 == "/api/wf-poll-images":
            """轮询ComfyUI异步生成的图片"""
            新图片 = list(网页请求处理器._comfyui图片队列)
            if 新图片:
                print(f"  [WF-Poll] 返回{len(新图片)}张图片: {新图片}")
            网页请求处理器._comfyui图片队列.clear()
            self._返回JSON({"成功": True, "图片列表": 新图片})
        elif 路径 == "/api/wf-scan-comfyui":
            """扫描ComfyUI工作流目录中所有_api.json文件"""
            import os as _os, pathlib as _pl
            结果 = []
            # 搜索目录列表
            搜索目录列表 = []
            # 1. ComfyUI用户工作流目录
            try:
                from 操作.ComfyUI操作 import _获取工作流目录
                comfyui目录 = _获取工作流目录()
                if _os.path.isdir(comfyui目录):
                    搜索目录列表.append(comfyui目录)
            except Exception:
                pass
            # 2. 项目内置工作流目录
            try:
                项目根 = self.配置加载器.项目根目录
                内置目录 = 项目根 / "公共区" / "工作流"
                if 内置目录.exists():
                    搜索目录列表.append(str(内置目录))
            except Exception:
                pass
            # 扫描所有_api.json
            seen = set()
            for 搜索目录 in 搜索目录列表:
                if not _os.path.exists(搜索目录):
                    continue
                for 根, _, 文件列表 in _os.walk(搜索目录):
                    # 计算相对搜索目录的子文件夹路径
                    相对路径 = _os.path.relpath(根, 搜索目录)
                    if 相对路径 == ".":
                        文件夹 = "其他"
                    else:
                        # 取第一级文件夹名
                        文件夹 = 相对路径.split(_os.sep)[0]
                        # 清理常见前缀
                        文件夹 = 文件夹.replace("01常用", "").replace("02图片", "图片").replace("03其他", "其他").replace("04图片", "图片").replace("05音频", "音频").replace("06其他", "其他").replace("07工程", "工程")
                        if not 文件夹:
                            文件夹 = "其他"
                    for f in 文件列表:
                        if f.endswith("_api.json"):
                            full = _os.path.join(根, f)
                            if full in seen:
                                continue
                            seen.add(full)
                            短名 = f.replace("_api.json", "")
                            # 分类：优先用文件夹名，其次用文件名关键词
                            分类 = 文件夹
                            if 分类 == "其他" or 分类 == "workflows":
                                for kw, cat in [("文生图","图片"),("text_to_image","图片"),("图生视频","视频"),("文生视频","视频"),("i2v","视频"),("t2v","视频"),
                                               ("图片修改","图片"),("image_edit","图片"),("放大","图片"),("upscal","图片"),
                                               ("反推","其他"),("interrogat","其他"),
                                               ("视频","视频"),("audio","音频"),("音乐","音频"),("tts","音频")]:
                                    if kw in 短名.lower():
                                        分类 = cat
                                        break
                            # 只分两类：图片 和 视频，其他归入其他
                            if 分类 not in ("图片", "视频", "音频", "工程", "其他"):
                                分类 = "其他"
                            结果.append({"名称": 短名, "文件名": f, "路径": full, "分类": 分类})
            # 提示信息
            提示 = ""
            if not 结果:
                提示 = "未找到任何_api格式工作流。请在ComfyUI中打开工作流→菜单→保存(API格式)→文件名以_api结尾（如：我的工作流_api.json）→保存到ComfyUI的user/default/workflows/目录"
            self._返回JSON({"成功": True, "工作流列表": 结果, "提示": 提示})
        elif 路径 == "/api/wf-delete":
            """删除节点图文件"""
            目录 = self.配置加载器.项目根目录 / "节点图"
            文件名 = 数据.get("文件名", "")
            安全名 = "".join(c for c in 文件名 if c not in '/\\:*?"<>|') or "未命名"
            路径文件 = 目录 / (安全名 + ".json")
            try:
                路径文件.unlink(missing_ok=True)
                self._返回JSON({"成功": True})
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/wf-logs":
            """查询工作流执行日志"""
            try:
                from 存储引擎 import 获取存储引擎
                存 = 获取存储引擎()
                会话ID = 数据.get("会话ID", "")
                日志 = 存.查询工作流日志(会话ID if 会话ID else None)
                self._返回JSON({"成功": True, "日志": 日志})
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/wf-auto-design":
            """AI自动设计工作流节点图"""
            try:
                需求 = 数据.get("需求", "")
                员工列表 = 数据.get("员工列表", [])
                if not 需求:
                    self._返回JSON({"成功": False, "错误": "缺少需求描述"})
                    return
                if not self.模型直连器:
                    self._返回JSON({"成功": False, "错误": "模型直连器未就绪"})
                    return

                # 构建员工信息
                员工信息 = "\n".join([f"- {e['姓名']}（{e.get('角色','')}）" for e in 员工列表])

                提示词 = f"""你是工作流设计专家。根据用户需求，设计一个节点工作流图。

## 可用员工
{员工信息}

## 设计原则
1. 优先使用已有员工
2. 如果需要新能力，创建自定义员工（取一个简洁的中文名，如"审稿员""润色师"）
3. 每个员工只做一件事，职责单一明确
4. 工作流要有清晰的流水线：上游产出→下游消费
5. 不要创建多余的节点，每个节点都要有实际用途

## 输出格式
只输出JSON，不要任何其他文字。格式：
{{
  "nodes": [
    {{"id": "n1", "type": "employee", "name": "员工名", "config": {{"员工名": "已有员工名", "指令": "具体告诉该员工做什么"}}}},
    {{"id": "n2", "type": "employee", "name": "自定义员工名", "config": {{"员工名": "自定义员工名", "角色": "精确描述该员工的能力和职责", "指令": "具体告诉该员工做什么"}}}}
  ],
  "conns": [
    {{"from": "n1", "to": "n2"}}
  ],
  "frame": {{"text": "分组名称（概括这个工作流的用途）", "color": "#4EC9B0"}}
}}

## 规则
1. 所有节点都是employee类型
2. 已有员工：config只需"员工名"和"指令"
3. 新员工：config需要"员工名"+"角色"（描述能力）+"指令"（具体任务）
4. "指令"必须具体明确，如"请审查文章的语法错误并逐条列出"
5. 连线决定执行顺序：from的输出传给to的输入
6. 一个节点可以连多个下游（并行），也可以多个上游汇入（合并输入）
7. id用n1,n2,n3...递增
8. frame用于把所有节点归入一个分组，text概括工作流用途

## 用户需求
{需求}
"""

                结果 = self.模型直连器.发送消息(
                    [{"role": "user", "content": 提示词}],
                    "你是工作流设计专家，只输出JSON。",
                    跳过缓存=True
                )
                if not 结果.get("成功"):
                    self._返回JSON({"成功": False, "错误": "AI生成失败: " + 结果.get("错误", "")})
                    return

                回复 = 结果.get("回复内容", "").strip()
                # 提取JSON
                import re as _re
                json匹配 = _re.search(r'\{[\s\S]*\}', 回复)
                if not json匹配:
                    self._返回JSON({"成功": False, "错误": "AI输出格式错误，未找到JSON"})
                    return
                try:
                    图 = json.loads(json匹配.group())
                except Exception:
                    self._返回JSON({"成功": False, "错误": "AI输出的JSON解析失败"})
                    return

                # 自动创建不存在的员工
                已有员工 = set(e.get("姓名", "") for e in 员工列表)
                新建员工 = []
                for node in 图.get("nodes", []):
                    cfg = node.get("config", {})
                    员工名 = cfg.get("员工名", "")
                    if 员工名 and 员工名 not in 已有员工:
                        角色 = cfg.get("角色", 员工名)
                        指令 = cfg.get("指令", "")
                        # 用AI生成人设
                        人设提示 = f"请为数字员工「{员工名}」生成系统提示词。角色：{角色}。要求：1.开头写'你的名字叫{员工名}' 2.描述角色专长和限制 3.当用户问你是谁时回答你是{员工名} 4.控制在100字以内 5.只输出提示词内容，不要多余解释"
                        人设追加 = f"你的名字叫{员工名}。{角色}。"
                        try:
                            人设结果 = self.模型直连器.发送消息(
                                [{"role": "user", "content": 人设提示}],
                                "你是一个人设生成助手，简洁输出。",
                                跳过缓存=True
                            )
                            if 人设结果.get("成功"):
                                人设追加 = 人设结果.get("回复内容", "").strip()
                        except Exception:
                            pass
                        # 随机选头像
                        import random as _rand
                        头像列表 = ['👨‍💻','👩‍💻','📝','🎨','🔬','📊','📷','🎵','🛠️','⚙️','📋','🧮','💡','🔧','🚀','🏆','🤖','🧑‍💼','👩‍💼','👨‍🔧','🧑‍🎨','🧑‍🔬','🧑‍🏫','🧑‍🌾']
                        头像 = _rand.choice(头像列表)
                        # 创建员工
                        if "员工管理" in self.模块注册:
                            self.模块注册["员工管理"].创建员工({
                                "姓名": 员工名,
                                "头像": 头像,
                                "角色": 角色,
                                "目标": 角色,
                                "人设追加": 人设追加,
                                "独立记忆": True,
                                "状态": "在岗"
                            })
                            已有员工.add(员工名)
                            新建员工.append(员工名)

                self._返回JSON({"成功": True, "图": 图, "新建员工": 新建员工})
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/wf-tasks":
            """工作流定时任务管理（POST=创建, PUT=更新, DELETE=删除）"""
            调度器 = getattr(网页请求处理器, '_定时任务调度器', None)
            if not 调度器:
                self._返回JSON({"成功": False, "错误": "定时任务调度器未就绪"})
                return
            方法 = getattr(self, '_http方法', 'POST')
            if 方法 == "POST":
                try:
                    结果 = 调度器.添加工作流任务(
                        名称=数据.get("名称", ""),
                        工作流文件=数据.get("工作流文件", ""),
                        类型=数据.get("类型", "每日"),
                        时间=数据.get("时间", "08:00"),
                        星期=数据.get("星期", []),
                        间隔分钟=数据.get("间隔分钟", 0),
                        通知=数据.get("通知", True),
                        日期=数据.get("日期", "")
                    )
                    self._返回JSON(结果)
                except Exception as _e:
                    self._返回JSON({"成功": False, "错误": str(_e)})
            elif 方法 == "PUT":
                结果 = 调度器.更新工作流任务(数据.get("id", ""), 数据.get("更新", {}))
                self._返回JSON(结果)
            elif 方法 == "DELETE":
                结果 = 调度器.删除工作流任务(数据.get("id", ""))
                self._返回JSON(结果)
            else:
                self._返回JSON({"成功": False, "错误": "不支持的方法"})
        elif 路径 == "/api/wf-task-notify":
            """获取工作流任务执行通知"""
            调度器 = getattr(网页请求处理器, '_定时任务调度器', None)
            if 调度器:
                self._返回JSON({"成功": True, "通知": 调度器.获取工作流通知()})
            else:
                self._返回JSON({"成功": True, "通知": []})
        elif 路径 == "/api/resume-checkpoint":
            """从检查点续跑"""
            if self.模块注册 and "对话" in self.模块注册:
                结果 = self.模块注册["对话"].续跑检查点()
                self._返回JSON(结果)
            else:
                self._返回JSON({"成功": False, "错误": "对话模块未就绪"})
        elif 路径 == "/api/clear-checkpoint":
            """清除检查点"""
            if self.模块注册 and "对话" in self.模块注册:
                self.模块注册["对话"]._清除检查点()
                self._返回JSON({"成功": True})
            else:
                self._返回JSON({"成功": False, "错误": "对话模块未就绪"})
        elif 路径 == "/api/file-read":
            结果 = self.文件管理器.读取文件(数据.get("路径", ""))
            self._返回JSON(结果)
        elif 路径 == "/api/file-write":
            结果 = self.文件管理器.写入文件(数据.get("路径", ""), 数据.get("内容", ""))
            self._返回JSON(结果)
        elif 路径 == "/api/file-mkdir":
            结果 = self.文件管理器.创建目录(数据.get("路径", ""))
            self._返回JSON(结果)
        elif 路径 == "/api/file-create":
            结果 = self.文件管理器.新建文件(数据.get("路径", ""))
            self._返回JSON(结果)
        elif 路径 == "/api/file-delete":
            结果 = self.文件管理器.删除(数据.get("路径", ""))
            self._返回JSON(结果)
        elif 路径 == "/api/open-in-explorer":
            目标路径 = os.path.abspath(数据.get("路径", ""))
            try:
                if sys.platform == "win32":
                    os.startfile(目标路径)
                elif sys.platform == "darwin":
                    subprocess.Popen(["open", 目标路径])
                else:
                    subprocess.Popen(["xdg-open", 目标路径])
                self._返回JSON({"成功": True})
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/fs-mkdir":
            """创建文件夹"""
            import os as _os
            目标 = 数据.get("path", "")
            if not 目标:
                self._返回JSON({"成功": False, "错误": "缺少路径"})
            else:
                try:
                    _os.makedirs(目标, exist_ok=True)
                    self._返回JSON({"成功": True})
                except Exception as e:
                    self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/file-rename":
            结果 = self.文件管理器.重命名(数据.get("路径", ""), 数据.get("新名称", ""))
            self._返回JSON(结果)
        elif 路径 == "/api/file-move":
            结果 = self.文件管理器.移动(数据.get("源路径", ""), 数据.get("目标目录", ""))
            self._返回JSON(结果)
        elif 路径 == "/api/file-copy":
            结果 = self.文件管理器.复制(数据.get("源路径", ""), 数据.get("目标目录", ""), 数据.get("新名称", None))
            self._返回JSON(结果)
        elif 路径 == "/api/file-replace":
            结果 = self.文件管理器.替换文本(
                数据.get("路径", ""),
                数据.get("旧文本", ""),
                数据.get("新文本", "")
            )
            self._返回JSON(结果)
        elif 路径 == "/api/save-image":
            """保存图片文件（直接二进制body）"""
            保存路径 = 数据.get("路径", "")
            if not 保存路径:
                self._返回JSON({"成功": False, "错误": "缺少路径"})
                return
            # 数据字段是 base64 编码的图片
            图片数据 = 数据.get("数据", "")
            if not 图片数据:
                self._返回JSON({"成功": False, "错误": "缺少数据"})
                return
            # 去掉 data:image/png;base64, 前缀
            if "," in 图片数据:
                图片数据 = 图片数据.split(",", 1)[1]
            try:
                import base64
                字节 = base64.b64decode(图片数据)
                # 确保目录存在
                目录 = os.path.dirname(保存路径)
                if 目录 and not os.path.exists(目录):
                    os.makedirs(目录, exist_ok=True)
                with open(保存路径, "wb") as f:
                    f.write(字节)
                print(f"  ✅ 图片已保存: {保存路径} ({len(字节)} 字节)")
                self._返回JSON({"成功": True, "路径": 保存路径})
            except Exception as e:
                print(f"  ❌ 图片保存失败: {e}")
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/image-inpaint":
            """图片加工 - OpenCV去水印/去杂物 (multipart/form-data: image, mask, algorithm, radius)"""
            try:
                import cv2
                import base64
                import numpy as np

                ctype = self.headers.get("Content-Type", "")
                if "multipart/form-data" not in ctype:
                    self._返回JSON({"成功": False, "错误": "需要multipart/form-data"})
                    return

                # 解析multipart
                boundary = ctype.split("boundary=")[1].encode()
                body = getattr(self, '_multipart_body', b'')
                parts = body.split(b"--" + boundary)

                image_data = None
                mask_data = None
                algorithm = "TELEA"
                radius = 3

                for part in parts:
                    if b"Content-Disposition" not in part:
                        continue
                    header_end = part.find(b"\r\n\r\n")
                    if header_end < 0:
                        continue
                    header = part[:header_end].decode("utf-8", errors="replace")
                    content = part[header_end+4:]
                    if content.endswith(b"\r\n"):
                        content = content[:-2]

                    if 'name="image"' in header:
                        image_data = content
                    elif 'name="mask"' in header:
                        mask_data = content
                    elif 'name="algorithm"' in header:
                        algorithm = content.decode("utf-8", errors="replace")
                    elif 'name="radius"' in header:
                        radius = int(content.decode("utf-8", errors="replace"))

                if not image_data or not mask_data:
                    self._返回JSON({"成功": False, "错误": "缺少图片或遮罩"})
                    return

                # 解码图片和遮罩
                img_arr = np.frombuffer(image_data, np.uint8)
                img = cv2.imdecode(img_arr, cv2.IMREAD_COLOR)
                mask_arr = np.frombuffer(mask_data, np.uint8)
                mask = cv2.imdecode(mask_arr, cv2.IMREAD_GRAYSCALE)

                if img is None or mask is None:
                    self._返回JSON({"成功": False, "错误": "无法解码图片"})
                    return

                # 尺寸对齐
                if img.shape[:2] != mask.shape[:2]:
                    mask = cv2.resize(mask, (img.shape[1], img.shape[0]))

                # 二值化遮罩
                _, mask = cv2.threshold(mask, 128, 255, cv2.THRESH_BINARY)

                # 膨胀遮罩边缘，扩大修复区域3px，减少接缝
                kernel = np.ones((3, 3), np.uint8)
                mask = cv2.dilate(mask, kernel, iterations=1)

                alg = cv2.INPAINT_NS if algorithm == "NS" else cv2.INPAINT_TELEA
                result = cv2.inpaint(img, mask, radius, alg)

                # 对修复区域边缘做羽化混合，消除接缝色差
                feather_mask = cv2.GaussianBlur(mask.astype(np.float32), (21, 21), 0)
                feather_mask = np.clip(feather_mask / 255.0, 0, 1)
                feather_3ch = cv2.merge([feather_mask, feather_mask, feather_mask])
                blended = (img.astype(np.float32) * (1 - feather_3ch) +
                          result.astype(np.float32) * feather_3ch)
                result = blended.astype(np.uint8)

                # 编码返回
                _, buf = cv2.imencode(".png", result)
                b64 = base64.b64encode(buf).decode("utf-8")
                self._返回JSON({"成功": True, "图片": b64})
            except ImportError:
                self._返回JSON({"成功": False, "错误": "opencv-python未安装"})
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/permission":
            self.文件管理器.用户确认权限(
                数据.get("路径", ""),
                数据.get("操作", "读"),
                数据.get("选择", "拒绝")
            )
            self._返回JSON({"成功": True})
        elif 路径 == "/api/ask-user-response":
            """用户在前端提交询问回答"""
            from 操作.询问用户 import 询问用户
            结果 = 询问用户.提交回答(数据.get("id", ""), 数据.get("回答", {}))
            self._返回JSON(结果)
        elif 路径 == "/api/ask-user-pending":
            """获取待答询问列表（SSE失败时轮询兼容）"""
            from 操作.询问用户 import 询问用户
            self._返回JSON({"待答": 询问用户.获取待答()})
        elif 路径 == "/api/shutdown":
            self._返回JSON({"成功": True})
            def _延迟退出():
                time.sleep(0.5)
                if self._启动器实例:
                    self._启动器实例.停止()
                os._exit(0)
            threading.Thread(target=_延迟退出, daemon=True).start()
        elif 路径 == "/api/save-config":
            self.配置加载器.保存配置(
                数据.get("名称", ""),
                数据.get("数据", {}),
                数据.get("区域", "公共区")
            )
            self._返回JSON({"成功": True})
        elif 路径 == "/api/wheel-config":
            """POST：保存快速呼出配置"""
            配置路径 = self.配置加载器.项目根目录 / "公共区" / "配置" / "快速呼出配置.json"
            try:
                with open(配置路径, "w", encoding="utf-8") as f:
                    json.dump(数据, f, ensure_ascii=False, indent=2)
                # 热更新：如果有快速浮窗实例，更新其配置
                启动器 = getattr(self, '启动器实例', None)
                if 启动器 and hasattr(启动器, '快速浮窗'):
                    启动器.快速浮窗.配置 = 数据
                    启动器.快速浮窗.半径 = 数据.get("轮盘半径", 72)
                    启动器.快速浮窗.中心圆半径 = 数据.get("中心圆半径", 26)
                    启动器.快速浮窗.透明度 = 数据.get("透明度", 0.88)
                    启动器.快速浮窗.字体大小 = max(4, 数据.get("字体大小", 12))
                    启动器.快速浮窗.扇区默认色 = 数据.get("扇区默认色", "#1c1c28")
                    启动器.快速浮窗.扇区hover色 = 数据.get("扇区hover色", "#3a3a52")
                    启动器.快速浮窗.边框色 = 数据.get("边框色", "#444466")
                    启动器.快速浮窗.中心圆色 = 数据.get("中心圆色", "#15151c")
                    启动器.快速浮窗.中心圆hover色 = 数据.get("中心圆hover色", "#2a2a3a")
                    启动器.快速浮窗.文字色 = 数据.get("文字色", "#aaaacc")
                    启动器.快速浮窗.文字hover色 = 数据.get("文字hover色", "#ffffff")
                self._返回JSON({"成功": True})
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/reload-config":
            self.配置加载器.重载配置()
            self._返回JSON({"成功": True})
        elif 路径 == "/api/model-config":
            """获取或保存模型配置"""
            if not self.模型直连器:
                self._返回JSON({"错误": "模型直连器未初始化"})
            elif not 数据:
                模型详情列表 = []
                for m in self.模型直连器.模型配置列表:
                    名称 = m.get("名称", "")
                    详情 = self.模型直连器.获取模型配置详情(名称)
                    模型详情列表.append(详情)
                self._返回JSON({"成功": True, "模型列表": 模型详情列表, "当前模型": self.模型直连器.当前模型名})
            else:
                模型名 = 数据.get("模型", "")
                密钥 = 数据.get("密钥", {})
                self.模型直连器.保存模型密钥(模型名, 密钥)
                密钥路径 = self.配置加载器.项目根目录 / "隐私区" / "我的配置" / "密钥.json"
                try:
                    from 模型直连器 import 加密密钥配置
                    加密后配置 = 加密密钥配置(self.模型直连器.密钥配置)
                    with open(密钥路径, "w", encoding="utf-8") as f:
                        json.dump(加密后配置, f, ensure_ascii=False, indent=2)
                    self._返回JSON({"成功": True, "消息": "密钥已保存（加密存储）"})
                except Exception as e:
                    self._返回JSON({"错误": f"保存失败: {e}"})
        elif 路径 == "/api/tool-keys":
            """工具密钥管理（Tavily等非LLM工具的API Key）"""
            if not self.模型直连器:
                self._返回JSON({"错误": "模型直连器未初始化"})
            elif not 数据:
                # GET：返回工具密钥状态（掩码）
                密钥列表 = self.模型直连器.密钥配置.get("密钥列表", {})
                tavily配置 = 密钥列表.get("TAVILY", {})
                tavily密钥 = tavily配置.get("API密钥", "") if isinstance(tavily配置, dict) else ""
                掩码密钥 = (tavily密钥[:6] + "****" + tavily密钥[-4:]) if len(tavily密钥) > 12 else ("已配置" if tavily密钥 else "")
                self._返回JSON({"成功": True, "工具列表": [
                    {"名称": "Tavily", "描述": "AI搜索引擎，网络搜索+网页抓取", "密钥字段": "API密钥", "已配置": bool(tavily密钥), "掩码值": 掩码密钥}
                ]})
            else:
                # POST：保存工具密钥
                工具名 = 数据.get("工具", "")
                密钥值 = 数据.get("密钥", "")
                if 工具名 == "Tavily" and 密钥值:
                    self.模型直连器.保存模型密钥("TAVILY", {"API密钥": 密钥值})
                    密钥路径 = self.配置加载器.项目根目录 / "隐私区" / "我的配置" / "密钥.json"
                    try:
                        from 模型直连器 import 加密密钥配置
                        加密后配置 = 加密密钥配置(self.模型直连器.密钥配置)
                        with open(密钥路径, "w", encoding="utf-8") as f:
                            json.dump(加密后配置, f, ensure_ascii=False, indent=2)
                        self._返回JSON({"成功": True, "消息": "Tavily密钥已保存（加密存储）"})
                    except Exception as e:
                        self._返回JSON({"错误": f"保存失败: {e}"})
                else:
                    self._返回JSON({"错误": "不支持的工具或密钥为空"})
        elif 路径 == "/api/run-action":
            if self.操作注册中心:
                操作名 = 数据.get("操作", "")
                参数 = 数据.get("参数", {})
                结果 = self.操作注册中心.执行(操作名, 参数)
                self._返回JSON(结果)
            else:
                self._返回JSON({"错误": "操作注册中心未初始化"})
        elif 路径 == "/api/work-mode":
            if self.模块注册 and "对话" in self.模块注册:
                模式 = 数据.get("模式", "商量")
                成功 = self.模块注册["对话"].设置工作模式(模式)
                self._返回JSON({"成功": 成功, "当前模式": self.模块注册["对话"].工作模式})
            else:
                self._返回JSON({"错误": "对话模块未加载"})
        elif 路径 == "/api/clear-chat":
            if self.模块注册 and "对话" in self.模块注册:
                self.模块注册["对话"].清空历史()
                self._返回JSON({"成功": True})
            else:
                self._返回JSON({"错误": "对话模块未加载"})
        elif 路径 == "/api/cancel":
            """用户取消当前正在执行的对话"""
            if self.模块注册 and "对话" in self.模块注册:
                self.模块注册["对话"].取消()
                self._返回JSON({"成功": True})
            else:
                self._返回JSON({"错误": "对话模块未加载"})
        elif 路径 == "/api/tts" or 路径 == "/api/wheel-tts":
            """TTS语音合成 - 三引擎：本地Kokoro(优先) → edge-tts → SAPI回退"""
            是轮盘 = (路径 == "/api/wheel-tts")
            状态dict = 网页请求处理器._tts轮盘状态 if 是轮盘 else 网页请求处理器._tts主界面状态
            通道号 = 1 if 是轮盘 else 0
            文件后缀 = "wheel" if 是轮盘 else "main"
            文本 = 数据.get("文本", "")
            if not 文本:
                self._返回JSON({"错误": "文本为空"})
                return
            文本 = 文本[:500]
            tts音量 = 数据.get("音量", 100)
            tts音量 = max(0, min(100, int(tts音量)))
            员工名 = 数据.get("员工名", "")
            语音配置 = _获取当前员工语音配置(员工名)
            # 试听时直接用前端传的语音配置覆盖
            if 数据.get("语音配置"):
                语音配置.update(数据["语音配置"])
            状态dict["代次"] += 1
            本次代次 = 状态dict["代次"]
            状态dict["播放中"] = True

            def _tts播放(待播文本, 播放音量, 状态, 代次, 通道, 后缀, 语音cfg):
                try:
                    import time
                    os.environ.setdefault("PYGAME_HIDE_SUPPORT_PROMPT", "1")
                    import pygame
                    if not pygame.mixer.get_init():
                        pygame.mixer.init(frequency=24000, size=-16, channels=1)
                    try:
                        pygame.mixer.Channel(通道).stop()
                    except Exception:
                        pass

                    引擎 = 语音cfg.get("引擎", "本地")
                    已播放 = False

                    # ① 本地Kokoro TTS（离线，高质量）
                    if 引擎 == "本地" and not 已播放:
                        try:
                            tts = _获取KokoroTTS引擎()
                            if tts is not None:
                                import sherpa_onnx, numpy as np
                                gen_config = sherpa_onnx.GenerationConfig()
                                gen_config.sid = int(语音cfg.get("说话人ID", 47))
                                gen_config.speed = float(语音cfg.get("语速", 1.0))
                                audio = tts.generate(待播文本, gen_config)
                                if len(audio.samples) > 0:
                                    # 峰值归一化到0.9（Kokoro输出振幅不一，统一放大到可用范围）
                                    samples_arr = np.asarray(audio.samples, dtype=np.float32)
                                    峰值 = np.max(np.abs(samples_arr))
                                    if 峰值 > 1e-8:
                                        samples_arr = samples_arr * (0.9 / 峰值)
                                    增益后 = np.clip(samples_arr, -1.0, 1.0)
                                    wav文件 = _float32转WAV(增益后, audio.sample_rate)
                                    if 状态["代次"] != 代次:
                                        try: os.remove(wav文件)
                                        except: pass
                                        return
                                    # pygame播放（支持音量滑块控制）
                                    try:
                                        pygame.mixer.quit()
                                        pygame.mixer.init(frequency=audio.sample_rate, size=-16, channels=1)
                                    except Exception:
                                        pass
                                    音频obj = pygame.mixer.Sound(wav文件)
                                    ch = pygame.mixer.Channel(通道)
                                    ch.set_volume(播放音量 / 100.0)
                                    ch.play(音频obj)
                                    while ch.get_busy():
                                        if 状态["代次"] != 代次:
                                            ch.stop()
                                            break
                                        time.sleep(0.05)
                                    try: os.remove(wav文件)
                                    except: pass
                                    已播放 = True
                            else:
                                print("  ⚠️ Kokoro引擎未加载，回退到edge-tts")
                        except Exception as e:
                            print(f"  ⚠️ Kokoro TTS异常: {e}")

                    # ② edge-tts（在线，高质量）
                    if not 已播放:
                        try:
                            import asyncio, edge_tts, tempfile
                            async def _生成():
                                edge音色 = 语音cfg.get("edge音色", "zh-CN-XiaoxiaoNeural")
                                communicate = edge_tts.Communicate(
                                    待播文本, edge音色,
                                    rate='+30%', volume='+100%'
                                )
                                mp3 = os.path.join(tempfile.gettempdir(), f'zf3d_tts_{后缀}.mp3')
                                await asyncio.wait_for(communicate.save(mp3), timeout=30.0)
                                return mp3
                            mp3文件 = asyncio.run(_生成())
                            if 状态["代次"] != 代次:
                                try: os.remove(mp3文件)
                                except: pass
                                return
                            音频 = pygame.mixer.Sound(mp3文件)
                            ch = pygame.mixer.Channel(通道)
                            ch.set_volume(播放音量 / 100.0)
                            ch.play(音频)
                            while ch.get_busy():
                                if 状态["代次"] != 代次:
                                    ch.stop()
                                    break
                                time.sleep(0.1)
                            try: os.remove(mp3文件)
                            except: pass
                            已播放 = True
                        except Exception:
                            pass

                    # ③ SAPI SpVoice（离线回退）
                    if not 已播放:
                        try:
                            import pythoncom, win32com.client
                            pythoncom.CoInitialize()
                            try:
                                speaker = win32com.client.Dispatch("SAPI.SpVoice")
                                speaker.Rate = 3
                                speaker.Volume = 播放音量
                                if 状态["代次"] == 代次:
                                    speaker.Speak(待播文本, 0)
                            finally:
                                pythoncom.CoUninitialize()
                            已播放 = True
                        except Exception:
                            pass

                    # ④ PowerShell最终回退
                    if not 已播放:
                        import subprocess
                        干净文本 = 待播文本.replace("'", "''").replace('"', '')
                        cmd = f'powershell -Command "Add-Type -AssemblyName System.Speech; (New-Object System.Speech.Synthesis.SpeechSynthesizer).Speak(\'{干净文本}\')"'
                        proc = subprocess.Popen(cmd, shell=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                                               creationflags=subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0)
                        if 状态["代次"] == 代次:
                            proc.wait()
                        else:
                            proc.kill()
                finally:
                    if 状态["代次"] == 代次:
                        状态["播放中"] = False
            t = threading.Thread(target=_tts播放, args=(文本, tts音量, 状态dict, 本次代次, 通道号, 文件后缀, 语音配置), daemon=True)
            t.start()
            self._返回JSON({"成功": True})
        elif 路径 == "/api/tts-stop" or 路径 == "/api/wheel-tts-stop":
            """停止TTS播放（各自独立停止）"""
            是轮盘 = (路径 == "/api/wheel-tts-stop")
            状态dict = 网页请求处理器._tts轮盘状态 if 是轮盘 else 网页请求处理器._tts主界面状态
            通道号 = 1 if 是轮盘 else 0
            状态dict["代次"] += 1
            状态dict["播放中"] = False
            try:
                os.environ.setdefault("PYGAME_HIDE_SUPPORT_PROMPT", "1")
                import pygame
                if pygame.mixer.get_init():
                    pygame.mixer.Channel(通道号).stop()
            except Exception:
                pass
            self._返回JSON({"成功": True})
        elif 路径 == "/api/conversation-new":
            if self.模块注册 and "对话" in self.模块注册:
                结果 = self.模块注册["对话"].新建对话()
                self._返回JSON({"成功": True, "对话": 结果})
            else:
                self._返回JSON({"错误": "对话模块未加载"})
        elif 路径 == "/api/conversation-switch":
            if self.模块注册 and "对话" in self.模块注册:
                结果 = self.模块注册["对话"].切换对话(数据.get("id", ""))
                self._返回JSON(结果)
            else:
                self._返回JSON({"错误": "对话模块未加载"})
        elif 路径 == "/api/conversation-messages":
            if self.模块注册 and "对话" in self.模块注册:
                结果 = self.模块注册["对话"].获取对话消息(数据.get("id", ""))
                self._返回JSON(结果)
            else:
                self._返回JSON({"错误": "对话模块未加载"})
        elif 路径 == "/api/conversation-delete":
            if self.模块注册 and "对话" in self.模块注册:
                结果 = self.模块注册["对话"].删除对话(数据.get("id", ""))
                self._返回JSON(结果)
            else:
                self._返回JSON({"错误": "对话模块未加载"})
        elif 路径 == "/api/conversation-rename":
            if self.模块注册 and "对话" in self.模块注册:
                结果 = self.模块注册["对话"].重命名对话(数据.get("id", ""), 数据.get("标题", ""))
                self._返回JSON(结果)
            else:
                self._返回JSON({"错误": "对话模块未加载"})
        elif 路径 == "/api/conversation-save":
            if self.模块注册 and "对话" in self.模块注册:
                self.模块注册["对话"]._保存当前对话()
            self._返回JSON({"成功": True})
        elif 路径 == "/api/check-update":
            """检查GitHub是否有新版本"""
            try:
                from 更新检查器 import 更新检查器类
                # 强制从文件读取最新配置（不依赖内存缓存，避免版本号过期）
                import json as _json_upd
                _cfg路径 = str(self.配置加载器.项目根目录 / "公共区" / "配置" / "系统配置.json")
                with open(_cfg路径, "r", encoding="utf-8-sig") as _f_upd:
                    系统配置 = _json_upd.load(_f_upd)
                系统配置["项目根目录"] = str(self.配置加载器.项目根目录)
                检查器 = 更新检查器类(系统配置)
                结果 = 检查器.检查更新(强制=数据.get("强制", False))
                self._返回JSON(结果)
            except Exception as e:
                self._返回JSON({"有更新": False, "错误": str(e)})
        elif 路径 == "/api/do-update":
            """执行更新：下载并覆盖公共区"""
            try:
                from 更新检查器 import 更新检查器类
                # 强制从文件读取最新配置
                import json as _json_upd2
                _cfg路径2 = str(self.配置加载器.项目根目录 / "公共区" / "配置" / "系统配置.json")
                with open(_cfg路径2, "r", encoding="utf-8-sig") as _f_upd2:
                    系统配置 = _json_upd2.load(_f_upd2)
                系统配置["项目根目录"] = str(self.配置加载器.项目根目录)
                检查器 = 更新检查器类(系统配置)
                下载地址 = 数据.get("下载地址", "")
                if not 下载地址:
                    结果 = 检查器.检查更新(强制=True)
                    下载地址 = 结果.get("下载地址", "")
                if not 下载地址:
                    self._返回JSON({"成功": False, "错误": "无法获取下载地址"})
                    return
                结果 = 检查器.执行更新(下载地址)
                self._返回JSON(结果)
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/download-cancel":
            """取消指定下载任务"""
            try:
                from 操作.多线程下载 import 多线程下载
                下载ID = int(数据.get("下载ID", 0))
                多线程下载.取消下载(下载ID)
                self._返回JSON({"成功": True, "消息": "下载取消中..."})
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/memory-add":
            if self.模块注册 and "对话" in self.模块注册:
                self.模块注册["对话"].添加永久记忆(数据.get("内容", ""))
                self._返回JSON({"成功": True})
            else:
                self._返回JSON({"错误": "对话模块未加载"})
        elif 路径 == "/api/plan-approve":
            if self.模块注册 and "对话" in self.模块注册:
                批准 = 数据.get("批准", False)
                if 批准:
                    self.模块注册["对话"].设置工作模式("执行")
                    self._返回JSON({"成功": True, "消息": "计划已批准，切换到执行模式"})
                else:
                    self.模块注册["对话"].当前计划 = None
                    self._返回JSON({"成功": True, "消息": "计划已拒绝"})
            else:
                self._返回JSON({"错误": "对话模块未加载"})
        elif 路径 == "/api/plan-execute":
            if self.模块注册 and "对话" in self.模块注册:
                结果 = self.模块注册["对话"].执行已批准计划()
                self._返回JSON(结果)
            else:
                self._返回JSON({"错误": "对话模块未加载"})
        elif 路径 == "/api/folder-dialog":
            选中路径 = self._打开文件夹选择对话框()
            self._返回JSON({"路径": 选中路径})
        elif 路径 == "/api/switch-model":
            # 切换当前对话模型
            模型名 = 数据.get("模型", "")
            if self.模型直连器:
                结果 = self.模型直连器.切换模型(模型名)
                if 结果.get("成功"):
                    self.当前模型名 = 模型名
                    if self.模块注册 and "对话" in self.模块注册:
                        self.模块注册["对话"].模型直连器 = self.模型直连器
                    # 持久化当前模型选择到 模型规则.json
                    try:
                        模型规则路径 = self.配置加载器.项目根目录 / "公共区" / "配置" / "模型规则.json"
                        with open(模型规则路径, "r", encoding="utf-8") as f:
                            模型规则 = json.load(f)
                        模型规则["当前模型"] = 模型名
                        with open(模型规则路径, "w", encoding="utf-8") as f:
                            json.dump(模型规则, f, ensure_ascii=False, indent=2)
                    except Exception:
                        pass
                    self._返回JSON({"成功": True, "当前模型": 模型名})
                else:
                    self._返回JSON({"错误": 结果.get("错误", "切换失败")})
            else:
                self._返回JSON({"错误": "模型直连器未初始化"})
        elif 路径 == "/api/model-ranking":
            """模型排行榜：GET排行数据/POST保存顺序/POST同步排名"""
            if not self.模型直连器:
                self._返回JSON({"错误": "模型直连器未初始化"})
            elif not 数据:
                # GET: 返回排行数据
                模型列表 = self.模型直连器.获取模型列表()
                self._返回JSON({"成功": True, "模型列表": 模型列表, "当前模型": self.模型直连器.当前模型名})
            elif 数据.get("顺序"):
                # POST: 保存自定义顺序
                结果 = self.模型直连器.保存排行顺序(数据["顺序"])
                self._返回JSON(结果)
            elif 数据.get("同步"):
                # POST: 同步全球排名（多源搜索最新benchmark）
                import urllib.request as _urq
                import urllib.parse as _ups
                import re as _re
                import json as _json
                更新分数 = {}
                更新价格 = {}
                错误列表 = []

                # ---- 源1: Artificial Analysis API (JSON) ----
                try:
                    req1 = _urq.Request("https://artificialanalysis.ai/api/models", headers={"User-Agent": "Mozilla/5.0", "Accept": "application/json"})
                    resp1 = _urq.urlopen(req1, timeout=10)
                    data1 = _json.loads(resp1.read().decode("utf-8", errors="replace"))
                    if isinstance(data1, list):
                        for item in data1:
                            name = (item.get("name") or item.get("model") or "").lower()
                            score = item.get("arena_score") or item.get("elo_score") or item.get("score")
                            if score and 50 <= int(score) <= 100:
                                if "fable" in name or "mythos" in name:
                                    更新分数["Claude(Anthropic)"] = int(score)
                                elif "gpt-5" in name or "gpt5" in name:
                                    更新分数["OpenAI(ChatGPT)"] = int(score)
                                elif "gemini" in name and "3" in name:
                                    更新分数["Gemini(Google)"] = int(score)
                                elif "grok" in name and "4" in name:
                                    更新分数["Grok(xAI)"] = int(score)
                                elif "deepseek" in name:
                                    更新分数["DeepSeek(深度求索)"] = int(score)
                                elif "mistral" in name and "large" in name:
                                    更新分数["Mistral AI"] = int(score)
                                elif "qwen" in name and "3" in name:
                                    更新分数["通义千问(阿里云)"] = int(score)
                                elif "glm" in name and "5" in name:
                                    更新分数["智谱大模型(GLM)"] = int(score)
                except Exception as e1:
                    错误列表.append("源1(ArtificialAnalysis): %s" % str(e1)[:80])

                # ---- 源2: HuggingFace LMSYS leaderboard 页面 ----
                try:
                    req2 = _urq.Request("https://huggingface.co/spaces/lmsys/chatbot-arena-leaderboard", headers={"User-Agent": "Mozilla/5.0"})
                    resp2 = _urq.urlopen(req2, timeout=10)
                    html2 = resp2.read().decode("utf-8", errors="replace")[:80000]
                    # 匹配 "Claude Fable 5" 附近的 ELO 分数
                    匹配规则 = [
                        ("Claude(Anthropic)", r"[Cc]laude\s*[Ff]able\s*5?\D{0,20}?(\d{3,4})\b", 1200, 1600),
                        ("OpenAI(ChatGPT)", r"GPT.?5\.?\d?\D{0,20}?(\d{3,4})\b", 1200, 1600),
                        ("Gemini(Google)", r"[Gg]emini\s*3\.?\d?\D{0,20}?(\d{3,4})\b", 1200, 1600),
                        ("Grok(xAI)", r"[Gg]rok\s*4\.?\d*\D{0,20}?(\d{3,4})\b", 1200, 1600),
                        ("DeepSeek(深度求索)", r"[Dd]eepseek\s*V?\d?\D{0,20}?(\d{3,4})\b", 1200, 1600),
                        ("Mistral AI", r"[Mm]istral\s*[Ll]arge\D{0,20}?(\d{3,4})\b", 1200, 1600),
                    ]
                    for 模型名, 正则, 最小, 最大 in 匹配规则:
                        m = _re.search(正则, html2)
                        if m:
                            原始分 = int(m.group(1))
                            if 最小 <= 原始分 <= 最大:
                                # ELO分数转0-100: (elo-1200)/4, 钳到50-100
                                转换分 = max(50, min(100, int((原始分 - 1200) / 4)))
                                if 模型名 not in 更新分数:
                                    更新分数[模型名] = 转换分
                except Exception as e2:
                    错误列表.append("源2(HuggingFace): %s" % str(e2)[:80])

                # ---- 源3: Bing搜索（Google在国内常被墙）----
                try:
                    查询 = "AI model benchmark ranking 2026 chatbot arena score Claude GPT Gemini DeepSeek"
                    url3 = "https://www.bing.com/search?q=" + _ups.quote(查询) + "&setlang=en"
                    req3 = _urq.Request(url3, headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"})
                    resp3 = _urq.urlopen(req3, timeout=10)
                    html3 = resp3.read().decode("utf-8", errors="replace")[:80000]
                    # 从搜索摘要中匹配模型+分数
                    摘要规则 = [
                        ("Claude(Anthropic)", r"[Cc]laude\s*[Ff]able\s*5?\D{0,30}?(\d{2,3})\b", 50, 100),
                        ("OpenAI(ChatGPT)", r"GPT.?5\.?\d?\D{0,30}?(\d{2,3})\b", 50, 100),
                        ("Gemini(Google)", r"[Gg]emini\s*3\.?\d?\D{0,30}?(\d{2,3})\b", 50, 100),
                        ("Grok(xAI)", r"[Gg]rok\s*4\.?\d*\D{0,30}?(\d{2,3})\b", 50, 100),
                        ("DeepSeek(深度求索)", r"[Dd]eepseek\D{0,30}?(\d{2,3})\b", 50, 100),
                        ("Mistral AI", r"[Mm]istral\s*[Ll]arge\D{0,30}?(\d{2,3})\b", 50, 100),
                    ]
                    for 模型名, 正则, 最小, 最大 in 摘要规则:
                        m = _re.search(正则, html3)
                        if m:
                            分数 = int(m.group(1))
                            if 最小 <= 分数 <= 最大:
                                if 模型名 not in 更新分数:
                                    更新分数[模型名] = 分数
                except Exception as e3:
                    错误列表.append("源3(Bing): %s" % str(e3)[:80])

                # ---- 源4: 各模型官方定价页（获取最新价格）----
                try:
                    # OpenAI pricing API (公开JSON)
                    req4 = _urq.Request("https://openai.com/api/pricing/", headers={"User-Agent": "Mozilla/5.0", "Accept": "application/json"})
                    resp4 = _urq.urlopen(req4, timeout=8)
                    html4 = resp4.read().decode("utf-8", errors="replace")[:30000]
                    # 尝试匹配 GPT-5 价格
                    gpt_match = _re.search(r'gpt.?5[^"]*?"?input"?\s*:\s*\$?([\d.]+).*?"?output"?\s*:\s*\$?([\d.]+)', html4, _re.I)
                    if gpt_match:
                        更新价格["OpenAI(ChatGPT)"] = {"价格输入": float(gpt_match.group(1)), "价格输出": float(gpt_match.group(2))}
                except Exception as e4:
                    错误列表.append("源4(OpenAI定价): %s" % str(e4)[:80])

                # ---- 汇总结果 ----
                消息 = ""
                if 更新分数:
                    self.模型直连器.更新实力分(更新分数)
                    消息 += "实力分更新%d个: %s. " % (len(更新分数), str(更新分数))
                else:
                    消息 += "实力分未获取到新数据，保持默认值. "
                if 更新价格:
                    for 模型名, 价格 in 更新价格.items():
                        for m in self.模型直连器.模型配置列表:
                            if m.get("名称") == 模型名:
                                m["价格输入"] = 价格["价格输入"]
                                m["价格输出"] = 价格["价格输出"]
                                break
                    # 保存到json
                    模型规则路径 = self.配置加载器.项目根目录 / "公共区" / "配置" / "模型规则.json"
                    with open(模型规则路径, "r", encoding="utf-8") as f:
                        规则 = _json.load(f)
                    for m in 规则.get("模型配置列表", []):
                        if m.get("名称") in 更新价格:
                            m["价格输入"] = 更新价格[m["名称"]]["价格输入"]
                            m["价格输出"] = 更新价格[m["名称"]]["价格输出"]
                    with open(模型规则路径, "w", encoding="utf-8") as f:
                        _json.dump(规则, f, ensure_ascii=False, indent=2)
                    消息 += "价格更新%d个. " % len(更新价格)
                else:
                    消息 += "价格未获取到新数据，保持默认值. "
                if 错误列表:
                    消息 += "(部分源失败: %d/%d)" % (len(错误列表), 4)
                self._返回JSON({"成功": True, "消息": 消息, "更新分数": 更新分数, "更新价格": 更新价格, "源错误": 错误列表})
            else:
                self._返回JSON({"错误": "未知操作"})
        elif 路径 == "/api/task-run":
            if self.模块注册 and "任务" in self.模块注册:
                结果 = self.模块注册["任务"].运行(数据)
                self._返回JSON(结果)
            else:
                self._返回JSON({"错误": "任务模块未加载"})
        elif 路径 == "/api/engine-merge":
            """合并工作引擎文件到主引擎"""
            try:
                结果 = self._执行引擎合并(数据)
                self._返回JSON(结果)
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/engine-rollback":
            """回滚主引擎到指定备份"""
            try:
                备份名 = 数据.get("备份", "")
                结果 = self._执行引擎回滚(备份名)
                self._返回JSON(结果)
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/evolution-control":
            """控制进化引擎：启动/停止/暂停/恢复/设置目标"""
            启动器 = getattr(self, '_启动器实例', None)
            if not 启动器:
                self._返回JSON({"成功": False, "错误": "启动器实例未绑定"})
                return
            动作 = 数据.get("动作", "")
            if 动作 == "启动":
                if hasattr(启动器, '进化引擎'):
                    self._返回JSON({"成功": False, "错误": "进化引擎已在运行"})
                    return
                if not hasattr(启动器, '_进化引擎类'):
                    self._返回JSON({"成功": False, "错误": "进化引擎类未预加载，请检查模型规则.json中自我进化配置"})
                    return
                try:
                    进化配置 = getattr(启动器, '_进化配置', {})
                    启动器.进化引擎 = 启动器._进化引擎类(启动器.模型直连器, 启动器.项目根目录, 进化配置)
                    启动器.进化引擎.启动()
                    self._返回JSON({"成功": True, "消息": "进化引擎已启动（测试员+开发者+审查员）"})
                except Exception as e:
                    self._返回JSON({"成功": False, "错误": f"启动失败: {e}"})
                return
            if 动作 == "停止":
                if hasattr(启动器, '进化引擎'):
                    启动器.进化引擎.停止()
                    del 启动器.进化引擎
                    self._返回JSON({"成功": True, "消息": "进化引擎已停止"})
                else:
                    self._返回JSON({"成功": False, "错误": "进化引擎未运行"})
                return
            if 动作 == "重置工作引擎":
                if hasattr(启动器, '进化引擎'):
                    启动器.进化引擎.重置工作引擎()
                    del 启动器.进化引擎
                    self._返回JSON({"成功": True, "消息": "工作引擎已重置，从主引擎重新同步完成"})
                else:
                    # 引擎未运行，直接创建临时实例执行重置
                    if hasattr(启动器, '_进化引擎类'):
                        进化配置 = getattr(启动器, '_进化配置', {})
                        临时引擎 = 启动器._进化引擎类(启动器.模型直连器, 启动器.项目根目录, 进化配置)
                        临时引擎.重置工作引擎()
                        self._返回JSON({"成功": True, "消息": "工作引擎已重置，从主引擎重新同步完成"})
                    else:
                        self._返回JSON({"成功": False, "错误": "进化引擎类未预加载"})
                return
            if not hasattr(启动器, '进化引擎'):
                self._返回JSON({"成功": False, "错误": "进化引擎未启动"})
                return
            if 动作 == "暂停":
                启动器.进化引擎.暂停()
                self._返回JSON({"成功": True, "消息": "已暂停"})
            elif 动作 == "恢复":
                启动器.进化引擎.恢复()
                self._返回JSON({"成功": True, "消息": "已恢复"})
            elif 动作 == "设置目标":
                目标 = 数据.get("目标", "")
                启动器.进化引擎.设置目标(目标)
                self._返回JSON({"成功": True, "消息": f"目标已设置: {目标}"})
            elif 动作 == "对话测试模式":
                启用 = 数据.get("启用", False)
                启动器.进化引擎.设置对话测试模式(启用)
                self._返回JSON({"成功": True, "消息": f"对话测试模式: {'开启' if 启用 else '关闭'}"})
            else:
                self._返回JSON({"成功": False, "错误": f"未知动作: {动作}"})
        elif 路径 == "/api/record-devices":
            """列出可用录音设备"""
            try:
                from 录音器 import 录音器
                self._返回JSON(录音器.列出设备())
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/record-start":
            """开始录制系统音频"""
            保存目录 = 数据.get("保存目录", "")
            if not 保存目录:
                保存目录 = 网页请求处理器._最后打开的文件夹 or str(Path.home() / "Desktop")
            设备索引 = 数据.get("设备索引", None)
            try:
                from 录音器 import 录音器
                结果 = 录音器.开始录制(保存目录, 设备索引)
                self._返回JSON(结果)
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/record-stop":
            """停止录制并保存"""
            try:
                from 录音器 import 录音器
                音量倍数 = 数据.get("音量倍数", 1.0)
                结果 = 录音器.停止录制(音量倍数)
                self._返回JSON(结果)
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/record-status":
            """查询录音状态"""
            try:
                from 录音器 import 录音器
                self._返回JSON(录音器.查询状态())
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/screenrecord-devices":
            """列出录屏可用的 dshow 音频设备"""
            try:
                from 录屏器 import 录屏器
                self._返回JSON(录屏器.列出dshow设备())
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/screenrecord-select-area":
            """弹出区域选择遮罩，返回选区坐标"""
            try:
                import tkinter as tk
                from 区域选择 import 区域选择
                root = tk.Tk()
                root.withdraw()
                选择器 = 区域选择(root)
                结果 = 选择器.弹出()
                root.destroy()
                if 结果:
                    self._返回JSON({"成功": True, "区域": 结果})
                else:
                    self._返回JSON({"成功": False, "错误": "用户取消了区域选择"})
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/screenrecord-start":
            """开始录制屏幕"""
            保存目录 = 数据.get("保存目录", "")
            if not 保存目录:
                保存目录 = 网页请求处理器._最后打开的文件夹 or str(Path.home() / "Desktop")
            x = 数据.get("x", 0)
            y = 数据.get("y", 0)
            w = 数据.get("w", 0)
            h = 数据.get("h", 0)
            帧率 = 数据.get("帧率", 30)
            音频模式 = 数据.get("音频模式", "mic")
            dshow设备名 = 数据.get("dshow设备名", "")
            麦克风音量 = 数据.get("麦克风音量", 1.0)
            麦克风静音 = 数据.get("麦克风静音", False)
            系统音量 = 数据.get("系统音量", 1.0)
            系统静音 = 数据.get("系统静音", False)
            点击效果 = 数据.get("点击效果", False)
            点击音效 = 数据.get("点击音效", False)
            音效音量 = 数据.get("音效音量", 50)
            # 保存设置到服务器（供轮盘快捷录屏复用）
            网页请求处理器._录屏设置.update({
                "帧率": 帧率, "音频模式": 音频模式, "dshow设备名": dshow设备名,
                "麦克风音量": 麦克风音量, "麦克风静音": 麦克风静音,
                "系统音量": 系统音量, "系统静音": 系统静音,
                "点击效果": 点击效果, "点击音效": 点击音效, "音效音量": 音效音量,
            })
            try:
                from 录屏器 import 录屏器
                结果 = 录屏器.开始录制(保存目录, x, y, w, h, 帧率, 音频模式, dshow设备名,
                                     麦克风音量, 麦克风静音, 系统音量, 系统静音,
                                     点击效果, 点击音效, 音效音量)
                self._返回JSON(结果)
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/screenrecord-stop":
            """停止录屏（同步等待转码完成，直接返回结果）"""
            try:
                from 录屏器 import 录屏器, _写日志
                _写日志("API:screenrecord-stop", "信息", "前端调用停止录屏API")
                结果 = 录屏器.停止并等待完成(超时秒=300)
                _写日志("API:screenrecord-stop响应", "信息",
                        f"成功={结果.get('成功')} 保存路径={结果.get('保存路径', '')}")
                self._返回JSON(结果)
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/screenrecord-status":
            """查询录屏状态"""
            try:
                from 录屏器 import 录屏器, _写日志
                状态 = 录屏器.查询状态()
                _写日志("API:screenrecord-status", "信息",
                        f"转码中={状态.get('转码中')} 转码完成={状态.get('转码完成')} "
                        f"录制中={状态.get('录制中')}")
                self._返回JSON(状态)
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/screenrecord-logs":
            """查询录屏日志"""
            try:
                from 录屏器 import 录屏器
                会话ID = 数据.get("会话ID", "")
                结果 = 录屏器.查询录屏日志(会话ID if 会话ID else None)
                self._返回JSON(结果)
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/screenrecord-log":
            """前端发来的日志"""
            try:
                from 录屏器 import _写日志
                消息 = 数据.get("消息", "")
                _写日志("前端日志", "信息", 消息)
                self._返回JSON({"成功": True})
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/screenrecord-settings":
            """返回后端保存的录屏设置（供轮盘快捷录屏复用）"""
            self._返回JSON({"成功": True, "设置": 网页请求处理器._录屏设置})
        elif 路径 == "/api/screenrecord-test-volume":
            """试听音量：用系统声音文件按当前音量倍数直接播放"""
            try:
                import tempfile
                音量倍数 = 数据.get("音量倍数", 1.0)
                _ffmpeg_path = shutil.which("ffmpeg") or r"C:\ffmpeg\bin\ffmpeg.exe"
                源音频 = r"C:\Windows\Media\chord.wav"
                if not os.path.exists(源音频):
                    源音频 = r"C:\Windows\Media\ding.wav"
                临时wav = os.path.join(tempfile.gettempdir(), f"_sr_test_{int(time.time()*1000)}.wav")
                vol = max(0, 音量倍数)
                cmd = [_ffmpeg_path, "-y", "-i", 源音频,
                       "-af", f"volume={vol}",
                       "-ar", "44100", "-ac", "2", 临时wav]
                subprocess.run(cmd, capture_output=True, timeout=10,
                    creationflags=subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0)
                if os.path.exists(临时wav) and os.path.getsize(临时wav) > 100:
                    # 直接用系统命令播放wav
                    def _播放():
                        try:
                            if sys.platform == "win32":
                                import winsound
                                winsound.PlaySound(临时wav, winsound.SND_FILENAME)
                                os.remove(临时wav)
                        except Exception:
                            pass
                    threading.Thread(target=_播放, daemon=True).start()
                    self._返回JSON({"成功": True})
                else:
                    self._返回JSON({"成功": False, "错误": "音频生成失败"})
            except Exception as e:
                self._返回JSON({"成功": False, "错误": str(e)})
        elif 路径 == "/api/voice-install":
            """自动安装 sherpa-onnx + 下载模型（复用项目多线程下载模块）"""
            import shutil, tarfile
            from 操作.多线程下载 import 多线程下载
            网页请求处理器._语音安装状态 = {"步骤": "正在安装 sherpa-onnx 库...", "进度": 0, "完成": False, "错误": ""}
            def _后台安装():
                try:
                    # ① pip install sherpa-onnx
                    result = subprocess.run(
                        [sys.executable, "-m", "pip", "install", "sherpa-onnx"],
                        capture_output=True, text=True, timeout=300,
                        encoding="utf-8", errors="replace"
                    )
                    if result.returncode != 0:
                        网页请求处理器._语音安装状态["错误"] = f"sherpa-onnx安装失败: {result.stderr[:500]}"
                        return
                    网页请求处理器._语音安装状态["进度"] = 10
                    # ② 下载模型（复用项目多线程下载模块，进度自动显示在下载面板）
                    # 注意：模型必须放在纯英文路径，sherpa-onnx的C++底层不支持中文路径
                    模型目录 = _获取英文模型目录()
                    目标目录 = 模型目录 / "paraformer-streaming"
                    模型目录.mkdir(parents=True, exist_ok=True)
                    if not 目标目录.exists() or not (目标目录 / "encoder.int8.onnx").exists():
                        目标目录.mkdir(parents=True, exist_ok=True)
                        # 下载流式模型（旧离线模型不兼容流式，必须重新下载）
                        下载源 = [
                            "https://github.com/k2-fsa/sherpa-onnx/releases/download/asr-models/sherpa-onnx-streaming-paraformer-bilingual-zh-en.tar.bz2",
                        ]
                        下载器 = 多线程下载()
                        下载器.进度回调 = None
                        下载器.取消检查 = None
                        下载器._aria2c路径 = None
                        with 多线程下载._下载进度锁:
                            多线程下载._下载进度表[999] = {
                                "文件名": "语音模型.tar.bz2", "百分比": 0, "已下载MB": 0,
                                "总大小MB": 0, "速度MB每秒": 0, "ETA": "", "状态": "启动中"
                            }
                        临时文件 = str(模型目录 / "语音模型.tar.bz2")
                        成功 = False
                        for 源idx, 地址 in enumerate(下载源):
                            网页请求处理器._语音安装状态["步骤"] = f"下载模型（源{源idx+1}，多线程加速）..."
                            多线程下载._保存任务(999, {
                                "下载地址": 地址, "保存路径": 临时文件,
                                "线程数": 16, "重试次数": 5, "已取消": False,
                                "启动时间": time.strftime("%Y-%m-%d %H:%M:%S"),
                            })
                            结果 = 下载器._执行下载(地址, Path(临时文件), 16, 5, 下载ID=999, 文件名="语音模型.tar.bz2")
                            if 结果.成功:
                                成功 = True
                                break
                            网页请求处理器._语音安装状态["步骤"] = f"源{源idx+1}失败: {结果.错误[:100]}，尝试下一个源..."
                        with 多线程下载._下载进度锁:
                            多线程下载._下载进度表.pop(999, None)
                        多线程下载._移除任务(999)
                        if not 成功:
                            网页请求处理器._语音安装状态["错误"] = "所有下载源均失败，请检查网络后重试"
                            return
                        网页请求处理器._语音安装状态["步骤"] = "正在解压模型..."
                        网页请求处理器._语音安装状态["进度"] = 90
                        with tarfile.open(临时文件, "r:bz2") as tar:
                            for member in tar.getmembers():
                                基名 = os.path.basename(member.name)
                                if 基名 in ("encoder.int8.onnx", "decoder.int8.onnx", "tokens.txt"):
                                    member.name = 基名
                                    tar.extract(member, 目标目录)
                        try: os.remove(临时文件)
                        except: pass
                        for d in 模型目录.glob("sherpa-onnx-streaming-*"):
                            if d.is_dir():
                                shutil.rmtree(d, ignore_errors=True)
                    网页请求处理器._语音安装状态["进度"] = 100
                    网页请求处理器._语音安装状态["步骤"] = "安装完成"
                    网页请求处理器._语音安装状态["完成"] = True
                except Exception as e:
                    网页请求处理器._语音安装状态["错误"] = str(e)
            threading.Thread(target=_后台安装, daemon=True).start()
            self._返回JSON({"成功": True, "消息": "安装已启动"})
        elif 路径 == "/api/voice-install-status":
            """查询安装进度"""
            self._返回JSON({"成功": True, "状态": 网页请求处理器._语音安装状态})
        elif 路径 == "/api/tts-install":
            """安装Kokoro TTS模型（复用多线程下载模块）"""
            import shutil, tarfile
            from 操作.多线程下载 import 多线程下载
            网页请求处理器._tts安装状态 = {"步骤": "正在下载Kokoro TTS模型...", "进度": 0, "完成": False, "错误": ""}
            def _后台安装TTS():
                try:
                    # 检查sherpa-onnx是否已安装
                    try:
                        import sherpa_onnx
                    except ImportError:
                        网页请求处理器._tts安装状态["步骤"] = "正在安装 sherpa-onnx 库..."
                        result = subprocess.run(
                            [sys.executable, "-m", "pip", "install", "sherpa-onnx"],
                            capture_output=True, text=True, timeout=300,
                            encoding="utf-8", errors="replace"
                        )
                        if result.returncode != 0:
                            网页请求处理器._tts安装状态["错误"] = f"sherpa-onnx安装失败: {result.stderr[:500]}"
                            return
                    网页请求处理器._tts安装状态["进度"] = 10

                    模型目录 = _获取英文模型目录() / "kokoro-tts"
                    模型目录.mkdir(parents=True, exist_ok=True)

                    # 先检查模型是否已存在（避免重复下载）
                    已有模型 = (模型目录 / "model.onnx").exists() or (模型目录 / "model.int8.onnx").exists()
                    已有voices = (模型目录 / "voices.bin").exists()
                    已有tokens = (模型目录 / "tokens.txt").exists()
                    if 已有模型 and 已有voices and 已有tokens:
                        网页请求处理器._kokoroTTS引擎 = None
                        网页请求处理器._tts安装状态["进度"] = 100
                        网页请求处理器._tts安装状态["步骤"] = "模型已存在，跳过下载"
                        网页请求处理器._tts安装状态["完成"] = True
                        return

                    # 清理残留的tar.bz2
                    残留 = 模型目录 / "kokoro-tts.tar.bz2"
                    if 残留.exists():
                        try: 残留.unlink()
                        except: pass

                    # 下载模型（int8版本，126MB）
                    下载源 = [
                        "https://github.com/k2-fsa/sherpa-onnx/releases/download/tts-models/kokoro-int8-multi-lang-v1_1.tar.bz2",
                        "https://github.com/k2-fsa/sherpa-onnx/releases/download/tts-models/kokoro-multi-lang-v1_1.tar.bz2",
                    ]
                    下载器 = 多线程下载()
                    下载器.进度回调 = None
                    下载器.取消检查 = None
                    下载器._aria2c路径 = None
                    with 多线程下载._下载进度锁:
                        多线程下载._下载进度表[998] = {
                            "文件名": "kokoro-tts.tar.bz2", "百分比": 0, "已下载MB": 0,
                            "总大小MB": 0, "速度MB每秒": 0, "ETA": "", "状态": "启动中"
                        }
                    临时文件 = str(模型目录 / "kokoro-tts.tar.bz2")
                    成功 = False
                    最后错误 = ""
                    for 源idx, 地址 in enumerate(下载源):
                        网页请求处理器._tts安装状态["步骤"] = f"下载Kokoro模型（源{源idx+1}，多线程加速）..."
                        多线程下载._保存任务(998, {
                            "下载地址": 地址, "保存路径": 临时文件,
                            "线程数": 16, "重试次数": 5, "已取消": False,
                            "启动时间": time.strftime("%Y-%m-%d %H:%M:%S"),
                        })
                        结果 = 下载器._执行下载(地址, Path(临时文件), 16, 5, 下载ID=998, 文件名="kokoro-tts.tar.bz2")
                        if 结果.成功:
                            成功 = True
                            break
                        最后错误 = 结果.错误[:200] if 结果.错误 else "未知错误"
                        网页请求处理器._tts安装状态["步骤"] = f"源{源idx+1}失败，尝试下一个源..."
                    with 多线程下载._下载进度锁:
                        多线程下载._下载进度表.pop(998, None)
                    多线程下载._移除任务(998)
                    if not 成功:
                        网页请求处理器._tts安装状态["错误"] = f"下载失败: {最后错误}"
                        return
                    # 解压模型
                    网页请求处理器._tts安装状态["步骤"] = "正在解压模型..."
                    网页请求处理器._tts安装状态["进度"] = 90
                    with tarfile.open(临时文件, "r:bz2") as tar:
                        tar.extractall(模型目录)
                    # 删除tar.bz2压缩包
                    try: os.remove(临时文件)
                    except: pass
                    # 检查解压后的目录结构，可能多一层
                    子目录 = list(模型目录.glob("kokoro-*"))
                    if 子目录:
                        for src in 子目录:
                            if src.is_dir():
                                for f in src.iterdir():
                                    目标 = 模型目录 / f.name
                                    if 目标.exists():
                                        if 目标.is_dir():
                                            shutil.rmtree(目标, ignore_errors=True)
                                        else:
                                            目标.unlink()
                                    shutil.move(str(f), str(目标))
                                shutil.rmtree(src, ignore_errors=True)
                    # 验证关键文件
                    有模型 = (模型目录 / "model.onnx").exists() or (模型目录 / "model.int8.onnx").exists()
                    有voices = (模型目录 / "voices.bin").exists()
                    有tokens = (模型目录 / "tokens.txt").exists()
                    if not (有模型 and 有voices and 有tokens):
                        缺失 = []
                        if not 有模型: 缺失.append("model.onnx")
                        if not 有voices: 缺失.append("voices.bin")
                        if not 有tokens: 缺失.append("tokens.txt")
                        网页请求处理器._tts安装状态["错误"] = f"解压后缺少文件: {', '.join(缺失)}"
                        return
                    # 重置引擎缓存（下次调用时重新加载）
                    网页请求处理器._kokoroTTS引擎 = None
                    网页请求处理器._tts安装状态["进度"] = 100
                    网页请求处理器._tts安装状态["步骤"] = "安装完成"
                    网页请求处理器._tts安装状态["完成"] = True
                except Exception as e:
                    网页请求处理器._tts安装状态["错误"] = str(e)
            threading.Thread(target=_后台安装TTS, daemon=True).start()
            self._返回JSON({"成功": True, "消息": "安装已启动"})
        elif 路径 == "/api/voice-stt":
            """本地语音识别：接收base64 WAV音频，返回文字"""
            if not hasattr(网页请求处理器, '_sherpa识别器') or 网页请求处理器._sherpa识别器 is None:
                try:
                    import sherpa_onnx
                except ImportError:
                    self._返回JSON({"成功": False, "错误": "sherpa-onnx未安装，请在设置→语音中安装"})
                    return
                模型目录 = _获取英文模型目录() / "paraformer-zh-int8"
                if not 模型目录.exists():
                    self._返回JSON({"成功": False, "错误": "模型未下载，请在设置→语音中安装"})
                    return
                try:
                    网页请求处理器._sherpa识别器 = sherpa_onnx.OfflineRecognizer.from_paraformer(
                        paraformer=str(模型目录 / "model.int8.onnx"),
                        tokens=str(模型目录 / "tokens.txt"),
                        num_threads=2,
                    )
                except Exception as e:
                    self._返回JSON({"成功": False, "错误": f"模型加载失败: {e}"})
                    return
            音频base64 = 数据.get("音频", "")
            if not 音频base64:
                self._返回JSON({"成功": False, "错误": "音频数据为空"})
                return
            try:
                import base64, wave, tempfile, numpy as np
                # base64 → WAV文件（用系统临时目录，避免隐私区目录不存在）
                音频字节 = base64.b64decode(音频base64.split(",")[-1] if "," in 音频base64 else 音频base64)
                临时wav = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)
                临时wav.write(音频字节)
                临时wav.close()
                # 读取WAV
                with wave.open(临时wav.name, "rb") as wf:
                    声道数 = wf.getnchannels()
                    采样率 = wf.getframerate()
                    帧数 = wf.getnframes()
                    原始数据 = wf.readframes(帧数)
                samples = np.frombuffer(原始数据, dtype=np.int16).astype(np.float32) / 32768.0
                if 声道数 > 1:
                    samples = samples[::声道数]
                # sherpa-onnx 识别
                流 = 网页请求处理器._sherpa识别器.create_stream()
                流.accept_waveform(采样率, samples)
                网页请求处理器._sherpa识别器.decode_stream(流)
                文字 = 流.result.text.strip()
                try:
                    os.remove(临时wav.name)
                except:
                    pass
                self._返回JSON({"成功": True, "文字": 文字})
            except Exception as e:
                self._返回JSON({"成功": False, "错误": f"识别失败: {e}"})
        elif 路径.startswith("/api/employee-"):
            self._处理员工管理API(路径, 数据)
        else:
            print(f"  ❌ 未知POST API: {路径}")
            self._返回JSON({"错误": "未知API: " + 路径}, 404)

    # 员工独立对话状态：{员工名: [{role, content}, ...]}
    _员工对话历史 = {}
    # 员工上线时间：{员工名: "ISO时间"}
    _员工上线时间 = {}
    # 员工提醒队列：{员工名: {"头像": "...", "间隔分钟": 30, "待发送": "消息", "上次提醒": "ISO时间"}}
    _员工提醒队列 = {}
    # ComfyUI异步生成图片队列：[{节点id, prompt_id, 图片路径, 状态, 时间}]
    _comfyui图片队列 = []
    # 员工提醒定时器
    _员工提醒线程 = None

    def _处理员工对话(self, 数据: dict):
        """员工独立对话 — 完全隔离于主对话模块"""
        try:
            消息 = 数据.get("消息", "")
            员工名 = 数据.get("姓名", "")
            if not 消息 or not 员工名:
                self._返回JSON({"成功": False, "错误": "缺少消息或员工名"})
                return
            if not self.模块注册 or "员工管理" not in self.模块注册:
                self._返回JSON({"成功": False, "错误": "员工管理模块未加载"})
                return
            if not self.模型直连器:
                self._返回JSON({"成功": False, "错误": "模型直连器未就绪"})
                return

            模块 = self.模块注册["员工管理"]
            配置结果 = 模块.获取运行时配置(员工名)
            if not 配置结果.get("成功"):
                self._返回JSON({"成功": False, "错误": f"员工「{员工名}」不存在"})
                return
            运行时 = 配置结果["数据"]
            系统提示词 = 运行时.get("系统提示词", "")

            # 记录上线时间并注入提示词
            from datetime import datetime as _dt
            if 员工名 not in 网页请求处理器._员工上线时间:
                网页请求处理器._员工上线时间[员工名] = _dt.now().isoformat()
                # 注册定时提醒（检查人设里是否提到提醒间隔）
                间隔 = 30  # 默认30分钟
                追加词 = 运行时.get("系统提示词", "")
                if "20分钟" in 追加词:
                    间隔 = 20
                elif "15分钟" in 追加词:
                    间隔 = 15
                elif "60分钟" in 追加词 or "1小时" in 追加词:
                    间隔 = 60
                头像 = 运行时.get("头像", "🙂")
                网页请求处理器._员工提醒队列[员工名] = {
                    "头像": 头像, "间隔分钟": 间隔, "弹窗时长秒": 30, "待发送": None,
                    "上次提醒": _dt.now().isoformat(), "角色": 运行时.get("角色", "")
                }
                # 启动提醒线程
                if not 网页请求处理器._员工提醒线程:
                    网页请求处理器._员工提醒线程 = True
                    import threading as _th
                    def _提醒循环():
                        while True:
                            import time as _time
                            _time.sleep(60)
                            now = _dt.now()
                            for 姓名, 信息 in list(网页请求处理器._员工提醒队列.items()):
                                if 信息["待发送"]:
                                    continue
                                上次 = _dt.fromisoformat(信息["上次提醒"])
                                经过 = (now - 上次).total_seconds() / 60
                                if 经过 >= 信息["间隔分钟"]:
                                    信息["待发送"] = f"⏰ 已工作{int(经过)}分钟，该休息了！起来走走，喝口水～"
                                    信息["上次提醒"] = now.isoformat()
                    _th.Thread(target=_提醒循环, daemon=True).start()
            上线时间 = 网页请求处理器._员工上线时间[员工名]
            现在时间 = _dt.now()
            工作时长分钟 = int((现在时间 - _dt.fromisoformat(上线时间)).total_seconds() / 60)
            间隔信息 = 网页请求处理器._员工提醒队列.get(员工名, {})
            系统提示词 += (
                f"\n\n## 时间信息\n你于{上线时间[:19]}上线，已工作{工作时长分钟}分钟。当前时间：{现在时间.strftime('%Y-%m-%d %H:%M:%S')}。\n"
                f"你有定时提醒功能，当前设置：每{间隔信息.get('间隔分钟',30)}分钟提醒，弹窗{间隔信息.get('弹窗时长秒',30)}秒。\n"
                f"你可以根据用户的工作状态动态调整：用户忙时拉长间隔，用户闲时缩短间隔。\n"
                f"调整方法：在回复末尾加上 [提醒设置:间隔=X分钟,时长=Y秒] 即可，系统会自动解析并应用。\n"
                f"例如：[提醒设置:间隔=15分钟,时长=10秒] 表示15分钟后提醒，弹窗10秒。\n"
                f"也可以只改一项：[提醒设置:间隔=45分钟] 或 [提醒设置:时长=60秒]"
            )
            # 注入当前工作目录
            当前文件夹 = 数据.get("当前文件夹", "")
            if 当前文件夹:
                系统提示词 += f"\n\n## 工作目录\n用户当前打开的文件夹：{当前文件夹}"

            # 老板注入团队信息
            树结果 = 模块.获取员工树()
            if 树结果.get("成功"):
                for 节点 in 树结果["数据"]:
                    if 节点.get("姓名") == 员工名:
                        下属 = 节点.get("下属", [])
                        if 下属:
                            团队信息 = "\n\n## 你的团队\n你是一名管理者，以下是你的下属：\n"
                            for s in 下属:
                                团队信息 += f"- {s.get('姓名','')}（{s.get('角色','')}）[{s.get('状态','')}]\n"
                            团队信息 += "\n你可以调配下属干活，也可以自己处理。"
                            系统提示词 += 团队信息
                        break

            # 获取或创建该员工的独立对话历史
            if 员工名 not in 网页请求处理器._员工对话历史:
                网页请求处理器._员工对话历史[员工名] = []
            历史 = 网页请求处理器._员工对话历史[员工名]

            # 添加用户消息
            历史.append({"role": "user", "content": 消息})

            # SSE流式响应
            self.send_response(200)
            self.send_header("Content-Type", "text/event-stream; charset=utf-8")
            self.send_header("Cache-Control", "no-cache")
            self.send_header("Connection", "close")
            self.end_headers()

            def _SSE写入(事件数据):
                try:
                    行 = f"data: {json.dumps(事件数据, ensure_ascii=False)}\n\n"
                    self.wfile.write(行.encode("utf-8"))
                    self.wfile.flush()
                except Exception:
                    pass

            # 调用LLM（直接用模型直连器，不走主对话模块）
            结果 = self.模型直连器.发送消息(历史, 系统提示词)

            if 结果.get("成功"):
                回复 = 结果.get("回复内容", "")
                # 解析员工动态提醒设置 [提醒设置:间隔=X分钟,时长=Y秒]
                import re as _re
                设置匹配 = _re.search(r'\[提醒设置[：:](.+?)\]', 回复)
                if 设置匹配 and 员工名 in 网页请求处理器._员工提醒队列:
                    设置文本 = 设置匹配.group(1)
                    信息 = 网页请求处理器._员工提醒队列[员工名]
                    间隔匹配 = _re.search(r'间隔[=:]?\s*(\d+)', 设置文本)
                    时长匹配 = _re.search(r'时长[=:]?\s*(\d+)', 设置文本)
                    if 间隔匹配:
                        信息["间隔分钟"] = int(间隔匹配.group(1))
                    if 时长匹配:
                        信息["弹窗时长秒"] = int(时长匹配.group(1))
                    # 从回复中移除设置标记
                    回复 = _re.sub(r'\s*\[提醒设置[：:].+?\]', '', 回复).strip()
                # 流式推送
                _SSE写入({"类型": "推理流", "记录": [{"类型": "流式回复", "内容": {"内容": 回复}}]})
                # 保存到历史
                历史.append({"role": "assistant", "content": 回复})
                # 保存到SQLite（员工对话记录表）
                try:
                    from 存储引擎 import 存储引擎类
                    存储实例 = 存储引擎类._实例引用
                    if 存储实例:
                        时间戳 = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        存储实例._执行("INSERT INTO 员工对话记录 (员工名, 角色, 内容, 时间, 来源) VALUES (?,?,?,?,?)",
                            (员工名, "user", 消息, 时间戳, "员工对话"))
                        存储实例._执行("INSERT INTO 员工对话记录 (员工名, 角色, 内容, 时间, 来源) VALUES (?,?,?,?,?)",
                            (员工名, "assistant", 回复, 时间戳, "员工对话"))
                except Exception:
                    pass
                # 保存到员工记忆文件（JSON，兼容旧逻辑）
                记忆路径 = 运行时.get("记忆路径")
                if 记忆路径:
                    try:
                        import json as _json
                        from pathlib import Path as _P
                        绝对路径 = self.配置加载器.项目根目录 / 记忆路径.lstrip("./")
                        绝对路径.parent.mkdir(parents=True, exist_ok=True)
                        记忆数据 = []
                        if 绝对路径.exists():
                            记忆数据 = _json.loads(绝对路径.read_text(encoding="utf-8"))
                        记忆数据.append({"用户": 消息, "助手": 回复, "时间": datetime.now().isoformat()})
                        绝对路径.write_text(_json.dumps(记忆数据, ensure_ascii=False, indent=2), encoding="utf-8")
                    except Exception:
                        pass
                _SSE写入({"类型": "完成", "结果": {"成功": True, "回复": 回复}})
            else:
                _SSE写入({"类型": "完成", "结果": {"成功": False, "错误": 结果.get("错误", "LLM调用失败")}})
        except Exception as e:
            if isinstance(e, (ConnectionAbortedError, ConnectionResetError, BrokenPipeError)):
                return
            try:
                self._返回JSON({"成功": False, "错误": f"员工对话异常: {str(e)}"})
            except Exception:
                return

    def _处理员工任务(self, 数据: dict):
        """员工任务执行 — 负责人出题→执行人逐题解答，SSE流式推送进度"""
        try:
            任务目标 = 数据.get("任务目标", "")
            数量 = int(数据.get("数量", 10))
            负责人 = 数据.get("负责人", "")
            执行人列表 = 数据.get("执行人列表", [])
            当前文件夹 = 数据.get("当前文件夹", "")

            if not 任务目标:
                self._返回JSON({"成功": False, "错误": "缺少任务目标"})
                return
            if not 负责人:
                self._返回JSON({"成功": False, "错误": "缺少负责人"})
                return
            if not 执行人列表:
                self._返回JSON({"成功": False, "错误": "缺少执行人"})
                return
            if not self.模块注册 or "员工管理" not in self.模块注册:
                self._返回JSON({"成功": False, "错误": "员工管理模块未加载"})
                return
            if not self.模型直连器:
                self._返回JSON({"成功": False, "错误": "模型直连器未就绪"})
                return

            模块 = self.模块注册["员工管理"]

            def _获取员工提示词(员工名):
                """获取员工运行时配置，返回(系统提示词, 头像, 角色)"""
                配置结果 = 模块.获取运行时配置(员工名)
                if not 配置结果.get("成功"):
                    return None, None, None
                运行时 = 配置结果["数据"]
                提示词 = 运行时.get("系统提示词", "")
                if 当前文件夹:
                    提示词 += f"\n\n## 工作目录\n{当前文件夹}"
                return 提示词, 运行时.get("头像", "🙂"), 运行时.get("角色", "")

            # SSE响应头
            self.send_response(200)
            self.send_header("Content-Type", "text/event-stream; charset=utf-8")
            self.send_header("Cache-Control", "no-cache")
            self.send_header("Connection", "close")
            self.end_headers()

            def _SSE写入(事件数据):
                try:
                    行 = f"data: {json.dumps(事件数据, ensure_ascii=False)}\n\n"
                    self.wfile.write(行.encode("utf-8"))
                    self.wfile.flush()
                except Exception:
                    pass

            import re as _re
            from datetime import datetime as _dt

            # ========== 阶段1: 负责人出题 ==========
            _SSE写入({"类型": "阶段", "阶段": "出题", "状态": "开始", "负责人": 负责人})

            负责人提示词, 负责人头像, 负责人角色 = _获取员工提示词(负责人)
            if not 负责人提示词:
                _SSE写入({"类型": "完成", "结果": {"成功": False, "错误": f"负责人「{负责人}」不存在"}})
                return

            出题指令 = (
                f"请根据任务目标「{任务目标}」，出{数量}道题。\n"
                f"格式要求：每道题独占一行，用数字编号开头（如 1. xxx），只输出题目，不要输出答案和解释。\n"
                f"请确保恰好输出{数量}道题。"
            )

            出题结果 = self.模型直连器.发送消息(
                [{"role": "user", "content": 出题指令}], 负责人提示词
            )

            if not 出题结果.get("成功"):
                _SSE写入({"类型": "完成", "结果": {"成功": False, "错误": "负责人出题失败: " + 出题结果.get("错误", "")}})
                return

            出题回复 = 出题结果.get("回复内容", "")

            # 解析题目列表
            题目列表 = []
            for 行 in 出题回复.strip().split("\n"):
                行 = 行.strip()
                匹配 = _re.match(r'^\d+[.、)）]\s*(.+)', 行)
                if 匹配:
                    题目列表.append(匹配.group(1).strip())

            if not 题目列表:
                # 按行分割作为后备
                题目列表 = [行.strip() for 行 in 出题回复.strip().split("\n") if 行.strip() and not 行.startswith("#")]

            实际数量 = len(题目列表)
            if 实际数量 == 0:
                _SSE写入({"类型": "完成", "结果": {"成功": False, "错误": "负责人未生成有效题目"}})
                return

            _SSE写入({"类型": "阶段", "阶段": "出题", "状态": "完成", "题目": 题目列表, "数量": 实际数量})

            # ========== 阶段2: 执行人逐题解答 ==========
            汇总结果 = []
            执行人数 = len(执行人列表)

            for i, 题目 in enumerate(题目列表):
                执行人 = 执行人列表[i % 执行人数]
                执行人提示词, 执行人头像, 执行人角色 = _获取员工提示词(执行人)

                if not 执行人提示词:
                    _SSE写入({"类型": "进度", "当前": i + 1, "总数": 实际数量, "题目": 题目, "结果": f"执行人「{执行人}」不存在", "执行人": 执行人, "执行人头像": "❓", "成功": False})
                    汇总结果.append({"题目": 题目, "结果": f"执行人不存在", "执行人": 执行人, "成功": False})
                    continue

                _SSE写入({"类型": "进度", "阶段": "解答中", "当前": i + 1, "总数": 实际数量, "题目": 题目, "执行人": 执行人, "执行人头像": 执行人头像})

                解答指令 = f"请解答以下问题，直接给出答案：\n{题目}"
                解答结果 = self.模型直连器.发送消息(
                    [{"role": "user", "content": 解答指令}], 执行人提示词
                )

                if 解答结果.get("成功"):
                    回复 = 解答结果.get("回复内容", "").strip()
                else:
                    回复 = "解答失败: " + 解答结果.get("错误", "")

                _SSE写入({"类型": "进度", "阶段": "已完成", "当前": i + 1, "总数": 实际数量, "题目": 题目, "结果": 回复, "执行人": 执行人, "执行人头像": 执行人头像, "成功": 解答结果.get("成功", False)})
                汇总结果.append({"题目": 题目, "结果": 回复, "执行人": 执行人, "成功": 解答结果.get("成功", False)})

            # ========== 阶段3: 完成 ==========
            正确数 = sum(1 for r in 汇总结果 if r.get("成功"))
            _SSE写入({"类型": "完成", "结果": {"成功": True, "汇总": 汇总结果, "总数": 实际数量, "完成数": len(汇总结果)}})

        except Exception as e:
            if isinstance(e, (ConnectionAbortedError, ConnectionResetError, BrokenPipeError)):
                return
            try:
                self._返回JSON({"成功": False, "错误": f"任务执行异常: {str(e)}"})
            except Exception:
                return

    def _处理员工工作流(self, 数据: dict):
        """节点工作流执行引擎 — 拓扑排序+并行分支，SSE推送每个节点状态"""
        try:
            节点列表 = 数据.get("节点", [])
            连接列表 = 数据.get("连接", [])
            当前文件夹 = 数据.get("当前文件夹", "")
            单节点上游 = 数据.get("单节点上游输入", "")  # 单节点执行时前端传的上游输入

            if not 节点列表:
                self._返回JSON({"成功": False, "错误": "工作流无节点"})
                return
            if not self.模块注册 or "员工管理" not in self.模块注册:
                self._返回JSON({"成功": False, "错误": "员工管理模块未加载"})
                return
            if not self.模型直连器:
                self._返回JSON({"成功": False, "错误": "模型直连器未就绪"})
                return

            模块 = self.模块注册["员工管理"]
            import threading as _th
            import concurrent.futures as _cf
            from datetime import datetime as _dt2
            import time as _time2

            # 生成会话ID
            wf会话ID = _dt2.now().strftime("%Y%m%d_%H%M%S")

            # 构建邻接表（保留loop属性）
            节点映射 = {n["id"]: n for n in 节点列表}
            出边 = {n["id"]: [] for n in 节点列表}
            入边 = {n["id"]: [] for n in 节点列表}
            连线映射 = {}  # (from,to) → conn对象
            for conn in 连接列表:
                f, t = conn.get("from", ""), conn.get("to", "")
                if f in 出边 and t in 入边:
                    出边[f].append(t)
                    入边[t].append(f)
                    连线映射[(f, t)] = conn

            # 检测是否有循环连线
            循环连线 = {k: v for k, v in 连线映射.items() if v.get("loop")}
            有循环 = len(循环连线) > 0

            # 拓扑排序（Kahn算法）— 无循环时用拓扑，有循环时用分层迭代
            入度 = {nid: len(入边[nid]) for nid in 节点映射}
            就绪队列 = [nid for nid, d in 入度.items() if d == 0]
            执行顺序 = []
            临时入度 = dict(入度)
            while 就绪队列:
                批次 = list(就绪队列)
                就绪队列 = []
                for nid in 批次:
                    执行顺序.append(nid)
                    for 下游 in 出边[nid]:
                        临时入度[下游] -= 1
                        if 临时入度[下游] == 0:
                            就绪队列.append(下游)

            if not 有循环 and len(执行顺序) < len(节点列表):
                self._返回JSON({"成功": False, "错误": "工作流存在循环依赖，请设置循环连线属性"})
                return

            # 按拓扑层分组（同层可并行）
            层级 = {}
            for nid in 执行顺序:
                if not 入边[nid]:
                    层级[nid] = 0
                else:
                    上游层级 = [层级.get(p, 0) for p in 入边[nid] if p in 层级]
                    层级[nid] = max(上游层级) + 1 if 上游层级 else 0
            最大层 = max(层级.values()) if 层级 else 0
            层分组 = [[] for _ in range(最大层 + 1)]
            for nid, lv in 层级.items():
                层分组[lv].append(nid)

            # SSE响应头
            self.send_response(200)
            self.send_header("Content-Type", "text/event-stream; charset=utf-8")
            self.send_header("Cache-Control", "no-cache")
            self.send_header("Connection", "close")
            self.end_headers()

            def _SSE写入(事件数据):
                try:
                    行 = f"data: {json.dumps(事件数据, ensure_ascii=False)}\n\n"
                    self.wfile.write(行.encode("utf-8"))
                    self.wfile.flush()
                except Exception as _sse_err:
                    print(f"  [WF SSE错误] {_sse_err} | 数据类型={事件数据.get('类型','?')}")

            # 存储引擎初始化（在SSE建立后，可以推送调试信息）
            try:
                from 存储引擎 import 获取存储引擎
                _wf存储 = 获取存储引擎()
            except Exception as _wf_err:
                _wf存储 = None
                _SSE写入({"类型": "调试", "消息": f"存储引擎初始化失败: {str(_wf_err)}"})

            # 节点输出存储
            节点输出 = {}

            def _获取员工提示词(员工名):
                配置结果 = 模块.获取运行时配置(员工名)
                if not 配置结果.get("成功"):
                    return None
                运行时 = 配置结果["数据"]
                提示词 = 运行时.get("系统提示词", "")
                if 当前文件夹:
                    提示词 += f"\n\n## 工作目录\n{当前文件夹}"
                return 提示词

            def _获取员工运行时(员工名):
                """获取员工运行时配置（含工具调用等属性）"""
                配置结果 = 模块.获取运行时配置(员工名)
                if not 配置结果.get("成功"):
                    return None
                return 配置结果["数据"]

            def _收集上游输入(nid):
                """收集所有上游节点的输出，合并为文本"""
                # 单节点模式：直接返回前端传的上游输入
                if 单节点上游:
                    return 单节点上游
                上游 = 入边.get(nid, [])
                if not 上游:
                    return ""
                # 按拓扑层级排序上游，保证输入顺序=执行顺序
                上游排序 = sorted(上游, key=lambda uid: 层级.get(uid, 0))
                parts = []
                for uid in 上游排序:
                    out = 节点输出.get(uid, "")
                    if out:
                        node = 节点映射.get(uid, {})
                        parts.append(f"【来自{node.get('name', uid)}】\n{out}")
                return "\n\n".join(parts)

            def _wf写日志(nid, nname, ntype, 状态, 输入, 输出, 错误, 耗时):
                """安全写入工作流日志"""
                if not _wf存储: return
                try:
                    _wf存储.写工作流日志(wf会话ID, nid, nname, ntype, 状态, 输入, 输出, 错误, 耗时)
                except Exception:
                    pass

            def _执行工具员工(nid, nname, 员工名, 初始消息, 提示词, 上游输入):
                """工具员工ReAct循环：多轮LLM+操作调用，思考过程不输出，只返回最终结果"""
                对话历史 = [{"role": "user", "content": 初始消息}]
                工具定义 = self.操作注册中心.获取工具定义() if hasattr(self.操作注册中心, '获取工具定义') else []
                工具提示 = "\n\n## 工具使用规则\n你拥有可调用的工具（function calling）。必须通过工具调用实际执行操作，不要只用文字描述你做了什么。如果你需要生成图片，必须调用生图工具；如果你需要读取文件，必须调用文件读取工具。绝不允许在文字中声称已执行操作而实际未调用工具。"
                最大步数 = 15
                最终回复 = ""
                for 步 in range(最大步数):
                    结果 = self.模型直连器.发送消息(对话历史, 提示词 + 工具提示, 工具列表=工具定义, 跳过缓存=True)
                    if not 结果.get("成功"):
                        return 结果.get("错误", "LLM调用失败")
                    回复 = 结果.get("回复内容", "").strip()
                    工具调用列表 = 结果.get("工具调用")
                    if not 工具调用列表:
                        # LLM没调用工具但声称执行了操作 → 强制再试一次
                        if 步 == 0 and ("已使用" in 回复 or "已调用" in 回复 or "已提交" in 回复 or "已生成" in 回复 or "已完成" in 回复):
                            对话历史.append({"role": "assistant", "content": 回复})
                            对话历史.append({"role": "user", "content": "你刚才声称执行了操作，但实际并未通过工具调用执行。请现在使用function calling调用相应工具来真正执行操作。不要只描述，要实际调用。"})
                            continue
                        最终回复 = 回复
                        break
                    # 工具调用是列表，取第一个
                    调用 = 工具调用列表[0]
                    英文名 = 调用.get("名称", 调用.get("name", ""))
                    英文参数 = 调用.get("参数", 调用.get("arguments", {}))
                    if not 英文名:
                        最终回复 = 回复
                        break
                    # 英文参数名→中文参数名映射
                    if hasattr(self.操作注册中心, '解析工具调用'):
                        工具名, 工具参数 = self.操作注册中心.解析工具调用(英文名, 英文参数)
                    else:
                        工具名, 工具参数 = 英文名, 英文参数
                    执行结果 = self.操作注册中心.执行(工具名, 工具参数)
                    操作输出 = 执行结果.get("数据") or 执行结果.get("结果") or str(执行结果)
                    对话历史.append({"role": "assistant", "content": 回复})
                    对话历史.append({"role": "user", "content": f"操作结果：\n{操作输出}\n\n请继续，或如果任务已完成请给出最终回复。"})
                else:
                    最终回复 = 回复

                # 最终总结：只输出结果，不输出思考过程
                if len(对话历史) > 1:
                    总结结果 = self.模型直连器.发送消息(
                        对话历史 + [{"role": "user", "content": "请根据以上工作过程，给出简洁的最终结果。只输出结果本身，不要重复思考过程和操作细节。"}],
                        提示词, 跳过缓存=True
                    )
                    if 总结结果.get("成功"):
                        总结 = 总结结果.get("回复内容", "").strip()
                        if 总结:
                            return 总结
                return 最终回复

            def _执行节点(nid):
                node = 节点映射[nid]
                ntype = node.get("type", "")
                nname = node.get("name", "")
                config = node.get("config", {})
                _t0 = _time2.time()

                # 禁用节点跳过执行
                if node.get("disabled"):
                    _SSE写入({"类型": "节点开始", "id": nid, "name": nname, "type": ntype})
                    节点输出[nid] = ""
                    _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": "🚫 节点已禁用", "成功": True})
                    _wf写日志(nid, nname, ntype, "跳过", "", "节点已禁用", "", int((_time2.time()-_t0)*1000))
                    return

                print(f"  [WF] 开始执行节点: {nname} ({ntype})")

                _SSE写入({"类型": "节点开始", "id": nid, "name": nname, "type": ntype})

                try:
                    if ntype == "target" or ntype == "目标":
                        目标 = config.get("目标", "")
                        数量 = int(config.get("数量", 1))
                        输出 = f"任务目标：{目标}\n数量：{数量}"
                        节点输出[nid] = 输出
                        print(f"  [WF] 目标节点输出: {输出[:60]}")
                        _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": 输出, "成功": True})
                        print(f"  [WF] 节点完成已推送: {nname}")
                        _wf写日志(nid, nname, ntype, "成功", "", 输出, "", int((_time2.time()-_t0)*1000))

                    elif ntype == "employee" or ntype == "员工":
                        员工名 = node.get("员工名", nname)
                        运行时 = _获取员工运行时(员工名)
                        if not 运行时:
                            节点输出[nid] = f"员工「{员工名}」不存在"
                            _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": "员工不存在", "成功": False})
                            _wf写日志(nid, nname, ntype, "失败", "", "员工不存在", f"员工「{员工名}」不存在", int((_time2.time()-_t0)*1000))
                            return

                        提示词 = 运行时.get("系统提示词", "")
                        if 当前文件夹:
                            提示词 += f"\n\n## 工作目录\n{当前文件夹}"
                        提示词 += "\n\n## 工作流上下文\n你是提示词增强流水线中的一个加工节点。上游给你一段画面提示词，你的任务：在上游提示词基础上加工（追加你负责的元素或做最终审核），然后直接输出完整的提示词。规则：1.去掉上游的【来自xxx】标签，只保留纯提示词内容 2.不要用方括号[]包裹提示词 3.直接输出纯文本提示词，不要任何格式标记 4.不要解释工作过程 5.不要输出设计文档"
                        启用工具 = 运行时.get("工具调用", False)

                        上游输入 = _收集上游输入(nid)
                        指令 = node.get("config", {}).get("指令", "")
                        if 指令 and 上游输入:
                            消息 = f"【用户指令】{指令}\n\n【上游输入】\n{上游输入}"
                        elif 指令:
                            消息 = f"【用户指令】{指令}"
                        else:
                            消息 = 上游输入 if 上游输入 else "请根据任务目标开始工作"

                        if 启用工具 and self.操作注册中心:
                            # 工具员工：ReAct循环
                            回复 = _执行工具员工(nid, nname, 员工名, 消息, 提示词, 上游输入)
                        else:
                            # 普通员工：单轮对话
                            print(f"  [WF] 调用LLM: {员工名}, 消息长度={len(消息)}")
                            结果 = self.模型直连器.发送消息(
                                [{"role": "user", "content": 消息}], 提示词, 跳过缓存=True
                            )
                            if 结果.get("成功"):
                                回复 = 结果.get("回复内容", "").strip()
                            else:
                                回复 = None
                                错误 = 结果.get("错误", "LLM调用失败")

                        if 回复 is not None:
                            节点输出[nid] = 回复
                            _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": 上游输入, "输出": 回复, "成功": True})
                            print(f"  [WF] 节点完成已推送: {nname}, 输出长度={len(回复)}")
                            _wf写日志(nid, nname, ntype, "成功", 上游输入, 回复, "", int((_time2.time()-_t0)*1000))
                        else:
                            节点输出[nid] = 错误
                            _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": 上游输入, "输出": 错误, "成功": False})
                            print(f"  [WF] 节点失败已推送: {nname}, 错误={错误[:60]}")
                            _wf写日志(nid, nname, ntype, "失败", 上游输入, 错误, 错误, int((_time2.time()-_t0)*1000))

                    elif ntype == "input" or ntype == "输入":
                        # 输入节点：直接输出config中的内容
                        config = node.get("config", {})
                        输出 = config.get("内容", "")
                        节点输出[nid] = 输出
                        _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": 输出[:200] if 输出 else "(空)", "成功": True})
                        _wf写日志(nid, nname, ntype, "成功", "", 输出[:200], "", int((_time2.time()-_t0)*1000))

                    elif ntype == "print" or ntype == "打印":
                        上游输入 = _收集上游输入(nid)
                        节点输出[nid] = 上游输入
                        _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": 上游输入, "输出": 上游输入, "成功": True})
                        _wf写日志(nid, nname, ntype, "成功", 上游输入, 上游输入, "", int((_time2.time()-_t0)*1000))

                    elif ntype == "prompt" or ntype == "文本输入":
                        # 文本输入节点：零token，直接输出用户填的提示词
                        输出 = config.get("提示词") or config.get("prompt") or ""
                        节点输出[nid] = 输出
                        _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": 输出[:500], "成功": True})
                        _wf写日志(nid, nname, ntype, "成功", "", 输出[:500], "", int((_time2.time()-_t0)*1000))

                    elif ntype == "text" or ntype == "文本":
                        # 文本拼接节点：零token，直接拼接指令+上游输入（不带名片）
                        指令 = config.get("指令") or config.get("text") or config.get("content") or ""
                        # 收集上游原始输出（不带【来自xxx】标签）
                        上游输入 = ""
                        if 单节点上游:
                            上游输入 = 单节点上游
                        else:
                            上游 = 入边.get(nid, [])
                            上游排序 = sorted(上游, key=lambda uid: 层级.get(uid, 0))
                            parts = []
                            for uid in 上游排序:
                                out = 节点输出.get(uid, "")
                                if out:
                                    parts.append(out)
                            上游输入 = "\n".join(parts)
                        if 上游输入 and 指令:
                            输出 = 上游输入 + 指令
                        elif 指令:
                            输出 = 指令
                        elif 上游输入:
                            输出 = 上游输入
                        else:
                            输出 = ""
                        节点输出[nid] = 输出
                        _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": 上游输入[:200], "输出": 输出[:500], "成功": True})
                        _wf写日志(nid, nname, ntype, "成功", 上游输入[:200], 输出[:500], "", int((_time2.time()-_t0)*1000))

                    elif ntype == "alert":
                        # 弹窗提醒节点：零token，弹窗显示内容+播放提示音
                        提醒内容 = config.get("提醒内容") or ""
                        # 收集上游输入
                        上游输入 = ""
                        if 单节点上游:
                            上游输入 = 单节点上游
                        else:
                            上游 = 入边.get(nid, [])
                            上游排序 = sorted(上游, key=lambda uid: 层级.get(uid, 0))
                            parts = []
                            for uid in 上游排序:
                                out = 节点输出.get(uid, "")
                                if out:
                                    parts.append(out)
                            上游输入 = "\n".join(parts)
                        # 组合最终提醒内容
                        if 提醒内容 and 上游输入:
                            最终内容 = 提醒内容 + "\n\n" + 上游输入
                        elif 提醒内容:
                            最终内容 = 提醒内容
                        elif 上游输入:
                            最终内容 = 上游输入
                        else:
                            最终内容 = "提醒时间到了！"
                        节点输出[nid] = 最终内容
                        # 推送弹窗通知到前端
                        声音开关 = config.get("声音", "default")
                        _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": 上游输入[:200], "输出": 最终内容[:500], "成功": True, "alert": True, "alert内容": 最终内容, "alert声音": 声音开关})
                        _wf写日志(nid, nname, ntype, "成功", 上游输入[:200], 最终内容[:500], "", int((_time2.time()-_t0)*1000))

                    elif ntype in ("file_read", "file_write", "file_mkdir", "file_search",
                                   "web_search", "web_fetch", "file_download",
                                   "run_command", "system_info",
                                   "image_watermark", "image_crop", "image_resize", "image_rotate",
                                   "code_search", "code_glob",
                                   "play_music", "play_video", "tts_speak", "video_convert"):
                        # 通用操作节点：通过操作注册中心执行，零token
                        上游输入 = _收集上游输入(nid)
                        # 参数：config中的值
                        参数 = dict(config)
                        # 检查上游节点是否有文件路径信息（拖入的文件节点）
                        上游节点列表 = 入边.get(nid, [])
                        for uid in 上游节点列表:
                            src_node = 节点映射.get(uid, {})
                            src_config = src_node.get("config", {})
                            if src_config.get("路径") and not 参数.get("路径"):
                                参数["路径"] = src_config["路径"]
                            if src_config.get("图片路径") and not 参数.get("图片路径"):
                                参数["图片路径"] = src_config["图片路径"]
                            if src_config.get("输入路径") and not 参数.get("输入路径"):
                                参数["输入路径"] = src_config["输入路径"]
                        # 如果路径还是空，但上游有内容（浏览器拖入的文件内容已预读），直接用内容作为输出
                        if ntype == "file_read" and not 参数.get("路径") and 上游输入:
                            节点输出[nid] = 上游输入
                            _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "使用上游文件内容", "输出": 上游输入[:500], "成功": True})
                            _wf写日志(nid, nname, ntype, "成功", "上游内容", 上游输入[:500], "", int((_time2.time()-_t0)*1000))
                            return
                        # 如果有上游文本输入，填充到常见字段
                        if 上游输入:
                            if "路径" in 参数 and not 参数["路径"]:
                                参数["路径"] = 上游输入
                            if "关键词" in 参数 and not 参数["关键词"]:
                                参数["关键词"] = 上游输入
                            if "内容" in 参数 and not 参数["内容"]:
                                参数["内容"] = 上游输入
                            if "文本" in 参数 and not 参数["文本"]:
                                参数["文本"] = 上游输入
                            if "网址" in 参数 and not 参数["网址"]:
                                参数["网址"] = 上游输入
                            if "歌曲名" in 参数 and not 参数["歌曲名"]:
                                参数["歌曲名"] = 上游输入
                            if "视频名" in 参数 and not 参数["视频名"]:
                                参数["视频名"] = 上游输入
                            if "命令" in 参数 and not 参数["命令"]:
                                参数["命令"] = 上游输入
                        try:
                            # 操作名映射
                            操作映射 = {
                                "file_read": "读取文件", "file_write": "写入文件", "file_mkdir": "创建文件夹", "file_search": "搜索文件",
                                "web_search": "网络搜索", "web_fetch": "网页抓取", "file_download": "多线程下载",
                                "run_command": "运行命令", "system_info": "系统信息",
                                "image_watermark": "去水印", "image_crop": "裁剪图片", "image_resize": "缩放图片", "image_rotate": "旋转图片",
                                "code_search": "搜索代码", "code_glob": "Glob搜索",
                                "play_music": "搜索音乐", "play_video": "搜索视频", "tts_speak": "朗读文本", "video_convert": "视频转码"
                            }
                            操作名 = 操作映射.get(ntype, ntype)
                            if self.操作注册中心:
                                结果 = self.操作注册中心.执行(操作名, 参数)
                                输出 = 结果.get("数据") or 结果.get("结果") or str(结果)
                                if isinstance(输出, dict):
                                    输出 = json.dumps(输出, ensure_ascii=False, indent=2)
                                节点输出[nid] = 输出
                                _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": str(参数)[:200], "输出": str(输出)[:500], "成功": True})
                            else:
                                节点输出[nid] = "操作注册中心未就绪"
                                _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": "操作注册中心未就绪", "成功": False})
                        except Exception as _op_err:
                            节点输出[nid] = str(_op_err)
                            _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": f"执行失败: {_op_err}", "成功": False})
                        _wf写日志(nid, nname, ntype, "成功" if 节点输出.get(nid) else "失败", str(参数)[:200], str(节点输出.get(nid, ""))[:500], "", int((_time2.time()-_t0)*1000))

                    elif ntype == "comfyui" or ntype == "生图":
                        # ComfyUI直出节点：不走LLM，直接调用ComfyUI出图，零token消耗
                        # 全局锁：同一时间只允许一个ComfyUI任务
                        if not hasattr(网页请求处理器, '_comfyui执行中'):
                            网页请求处理器._comfyui执行中 = {}
                        if not hasattr(网页请求处理器, '_comfyui全局锁'):
                            import threading as _th_lock
                            网页请求处理器._comfyui全局锁 = _th_lock.Lock()
                        # 检查是否有任何ComfyUI任务正在执行
                        if 网页请求处理器._comfyui执行中:
                            节点输出[nid] = "⏳ 有其他ComfyUI任务在执行中，等待完成后再试"
                            _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": "⏳ 等待其他ComfyUI任务完成", "成功": False})
                            _wf写日志(nid, nname, ntype, "跳过", "", "等待ComfyUI全局锁", "其他任务执行中", int((_time2.time()-_t0)*1000))
                        else:
                            上游输入 = _收集上游输入(nid)
                            提示词 = config.get("提示词") or config.get("prompt") or 上游输入 or ""
                            # 清理提示词：去掉上游标签、多余格式，只保留画面描述
                            import re as _re_clean
                            提示词 = _re_clean.sub(r'【来自[^】]*】\s*', '', 提示词)
                            提示词 = _re_clean.sub(r'```[a-z]*\n?', '', 提示词)
                            提示词 = _re_clean.sub(r'\*\*[^*]*\*\*:', '', 提示词)
                            提示词 = _re_clean.sub(r'^\s*[-•]\s*', '', 提示词, flags=_re_clean.MULTILINE)
                            提示词 = _re_clean.sub(r'\n{3,}', '\n\n', 提示词)
                            提示词 = 提示词.strip()
                            工作流名 = config.get("工作流") or config.get("workflow") or ""
                            宽度 = int(config.get("宽度") or config.get("width") or 0)
                            高度 = int(config.get("高度") or config.get("height") or 0)
                            if not 提示词:
                                节点输出[nid] = "无提示词输入"
                                _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": "无提示词", "成功": False})
                                _wf写日志(nid, nname, ntype, "失败", "", "无提示词", "缺少提示词", int((_time2.time()-_t0)*1000))
                            else:
                                try:
                                    from 操作.ComfyUI操作 import ComfyUI一键生图, _获取ComfyUI地址, _加载工作流文件, _注入种子, _注入宽高, _注入提示词通用, _API请求, _获取默认保存目录
                                    地址 = _获取ComfyUI地址()
                                    print(f"  [WF-ComfyUI] 探测地址: {地址}")
                                    保存目录 = 当前文件夹 if 当前文件夹 else _获取默认保存目录()
                                    print(f"  [WF-ComfyUI] 保存目录: {保存目录}")
                                    import os as _os2
                                    _os2.makedirs(保存目录, exist_ok=True)
                                    import json as _json2, uuid as _uuid2, random as _random2
                                    if 工作流名:
                                        成功2, 结果2, _ = _加载工作流文件(工作流名)
                                        if not 成功2:
                                            raise Exception(结果2)
                                        工作流dict = 结果2
                                    else:
                                        工作流dict = {
                                            "3": {"class_type": "KSampler", "inputs": {"seed": _random2.randint(0, 2**32-1), "steps": 20, "cfg": 7.0, "sampler_name": "euler", "scheduler": "normal", "denoise": 1.0, "model": ["4", 0], "positive": ["6", 0], "negative": ["7", 0], "latent_image": ["5", 0]}},
                                            "4": {"class_type": "CheckpointLoaderSimple", "inputs": {"ckpt_name": ""}},
                                            "5": {"class_type": "EmptyLatentImage", "inputs": {"width": 宽度 or 512, "height": 高度 or 512, "batch_size": 1}},
                                            "6": {"class_type": "CLIPTextEncode", "inputs": {"text": 提示词, "clip": ["4", 1]}},
                                            "7": {"class_type": "CLIPTextEncode", "inputs": {"text": "low quality, bad anatomy", "clip": ["4", 1]}},
                                            "8": {"class_type": "VAEDecode", "inputs": {"samples": ["3", 0], "vae": ["4", 2]}},
                                            "9": {"class_type": "SaveImage", "inputs": {"filename_prefix": "zf3d", "images": ["8", 0]}}
                                        }
                                    种子 = _random2.randint(0, 2**32-1)
                                    _注入种子(工作流dict, 种子)
                                    _注入宽高(工作流dict, 宽度, 高度)
                                    _注入提示词通用(工作流dict, 提示词)
                                    client_id = str(_uuid2.uuid4())
                                    payload = {"prompt": 工作流dict, "client_id": client_id}
                                    成功3, 结果3 = _API请求(地址, "/prompt", "POST", payload)
                                    if not 成功3:
                                        raise Exception(f"提交失败: {结果3}")
                                    prompt_id = 结果3.get("prompt_id", "")
                                    网页请求处理器._comfyui执行中[nid] = prompt_id
                                    输出文本 = f"✅ 已提交生成任务(不等待)\n提示词: {提示词[:60]}{'...' if len(提示词)>60 else ''}\nprompt_id: {prompt_id}"
                                    节点输出[nid] = 输出文本
                                    _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": 提示词[:200], "输出": 输出文本, "成功": True, "图片": ""})

                                    def _轮询图片(nid, prompt_id, 地址, 保存目录):
                                        import time as _time3
                                        for _ in range(120):
                                            _time3.sleep(2)
                                            try:
                                                ok, hist = _API请求(地址, f"/history/{prompt_id}")
                                                if ok and prompt_id in hist:
                                                    条目 = hist[prompt_id]
                                                    状态 = 条目.get("status", {})
                                                    if isinstance(状态, dict) and 状态.get("status_str") == "error":
                                                        网页请求处理器._comfyui图片队列.append({"节点id": nid, "prompt_id": prompt_id, "图片路径": "", "状态": "失败", "时间": _time3.time()})
                                                        网页请求处理器._comfyui执行中.pop(nid, None)
                                                        return
                                                    import os as _os3
                                                    for 节点id, 输出 in 条目.get("outputs", {}).items():
                                                        for img in 输出.get("images", []):
                                                            文件名 = img.get("filename", "")
                                                            if not 文件名:
                                                                continue
                                                            import urllib.request as _ur2, urllib.parse as _up2
                                                            参数 = {"filename": 文件名, "type": img.get("type", "output")}
                                                            if img.get("subfolder"):
                                                                参数["subfolder"] = img["subfolder"]
                                                            url = f"http://{地址}/view?{_up2.urlencode(参数)}"
                                                            保存路径 = _os3.path.join(保存目录, 文件名)
                                                            try:
                                                                req = _ur2.Request(url, headers={"User-Agent": "ZF3D-Agent"})
                                                                with _ur2.urlopen(req, timeout=30) as resp2:
                                                                    with open(保存路径, 'wb') as f2:
                                                                        f2.write(resp2.read())
                                                                print(f"  [WF-ComfyUI] 图片已保存: {保存路径}")
                                                            except Exception as dl_err:
                                                                print(f"  [WF-ComfyUI] 图片下载失败: {dl_err}, url={url}")
                                                            网页请求处理器._comfyui图片队列.append({"节点id": nid, "prompt_id": prompt_id, "图片路径": 保存路径, "文件名": 文件名, "状态": "完成", "时间": _time3.time()})
                                                            网页请求处理器._comfyui执行中.pop(nid, None)
                                                            return
                                            except Exception as poll_err:
                                                print(f"  [WF-ComfyUI] 轮询异常: {poll_err}")
                                        网页请求处理器._comfyui图片队列.append({"节点id": nid, "prompt_id": prompt_id, "图片路径": "", "状态": "超时", "时间": _time3.time()})
                                        网页请求处理器._comfyui执行中.pop(nid, None)

                                    import threading as _th2
                                    _th2.Thread(target=_轮询图片, args=(nid, prompt_id, 地址, 保存目录), daemon=True).start()
                                    _wf写日志(nid, nname, ntype, "成功", 提示词[:200], 输出文本, "", int((_time2.time()-_t0)*1000))
                                except Exception as e:
                                    错误 = f"ComfyUI出图失败: {e}"
                                    节点输出[nid] = 错误
                                    _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": 提示词[:200], "输出": 错误, "成功": False})
                                    _wf写日志(nid, nname, ntype, "失败", 提示词[:200], 错误, 错误, int((_time2.time()-_t0)*1000))
                    elif ntype == "comfyui-edit" or ntype == "图片修改":
                        # ComfyUI图片修改节点：上传图片+注入提示词，不走LLM
                        if not hasattr(网页请求处理器, '_comfyui执行中'):
                            网页请求处理器._comfyui执行中 = {}
                        if 网页请求处理器._comfyui执行中:
                            节点输出[nid] = "⏳ 等待其他ComfyUI任务完成"
                            _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": "⏳ 等待ComfyUI", "成功": False})
                            _wf写日志(nid, nname, ntype, "跳过", "", "等待全局锁", "", int((_time2.time()-_t0)*1000))
                        else:
                            上游输入 = _收集上游输入(nid)
                            提示词 = config.get("提示词") or config.get("prompt") or 上游输入 or ""
                            工作流名 = config.get("工作流") or config.get("workflow") or ""
                            图片路径 = config.get("图片路径") or config.get("路径") or ""
                            if not 工作流名:
                                节点输出[nid] = "未指定工作流"
                                _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": "未指定工作流", "成功": False})
                            elif not 图片路径:
                                节点输出[nid] = "未提供图片路径"
                                _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": "未提供图片路径", "成功": False})
                            else:
                                try:
                                    from 操作.ComfyUI操作 import ComfyUI图片修改, _获取ComfyUI地址, _获取默认保存目录
                                    操作实例 = ComfyUI图片修改()
                                    操作参数 = {"工作流": 工作流名, "图片路径": 图片路径}
                                    if 提示词:
                                        操作参数["提示词"] = 提示词
                                    保存目录 = 当前文件夹 if 当前文件夹 else _获取默认保存目录()
                                    操作参数["保存目录"] = 保存目录
                                    网页请求处理器._comfyui执行中[nid] = "edit"
                                    # 后台执行
                                    def _执行图片修改(nid, 操作实例, 操作参数, 保存目录):
                                        import time as _time_edit
                                        结果 = 操作实例.执行(操作参数)
                                        输出文本 = 结果.结果 if hasattr(结果, '结果') else str(结果)
                                        图片路径_out = ""
                                        if hasattr(结果, '元数据') and 结果.元数据:
                                            sd = 结果.元数据.get("保存目录", "")
                                            if sd:
                                                import os as _os_edit
                                                图片文件 = [f for f in _os_edit.listdir(sd) if f.endswith(('.png','.jpg','.jpeg','.webp'))]
                                                if 图片文件:
                                                    图片文件.sort(key=lambda f: _os_edit.getmtime(_os_edit.join(sd, f)), reverse=True)
                                                    图片路径_out = _os_edit.join(sd, 图片文件[0])
                                        网页请求处理器._comfyui图片队列.append({"节点id": nid, "图片路径": 图片路径_out, "文件名": "", "状态": "完成" if 结果.成功 else "失败", "时间": _time_edit.time()})
                                        网页请求处理器._comfyui执行中.pop(nid, None)
                                    import threading as _th_edit
                                    _th_edit.Thread(target=_执行图片修改, args=(nid, 操作实例, 操作参数, 保存目录), daemon=True).start()
                                    节点输出[nid] = f"✅ 已提交图片修改(不等待)\n工作流: {工作流名}\n图片: {图片路径}"
                                    _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": 提示词[:200], "输出": 节点输出[nid], "成功": True, "图片": ""})
                                    _wf写日志(nid, nname, ntype, "成功", 提示词[:200], 节点输出[nid], "", int((_time2.time()-_t0)*1000))
                                except Exception as e:
                                    网页请求处理器._comfyui执行中.pop(nid, None)
                                    错误 = f"图片修改失败: {e}"
                                    节点输出[nid] = 错误
                                    _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": 错误, "成功": False})
                                    _wf写日志(nid, nname, ntype, "失败", "", 错误, 错误, int((_time2.time()-_t0)*1000))

                    elif ntype == "comfyui-video" or ntype == "视频生成":
                        # ComfyUI视频生成节点：文生视频/图生视频，不走LLM
                        if not hasattr(网页请求处理器, '_comfyui执行中'):
                            网页请求处理器._comfyui执行中 = {}
                        if 网页请求处理器._comfyui执行中:
                            节点输出[nid] = "⏳ 等待其他ComfyUI任务完成"
                            _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": "⏳ 等待ComfyUI", "成功": False})
                            _wf写日志(nid, nname, ntype, "跳过", "", "等待全局锁", "", int((_time2.time()-_t0)*1000))
                        else:
                            上游输入 = _收集上游输入(nid)
                            提示词 = config.get("提示词") or config.get("prompt") or 上游输入 or ""
                            工作流名 = config.get("工作流") or config.get("workflow") or ""
                            图片路径 = config.get("图片路径") or config.get("路径") or ""
                            if not 工作流名:
                                节点输出[nid] = "未指定工作流"
                                _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": "未指定工作流", "成功": False})
                            else:
                                try:
                                    from 操作.ComfyUI操作 import ComfyUI视频生成, _获取ComfyUI地址, _获取默认保存目录
                                    操作实例 = ComfyUI视频生成()
                                    操作参数 = {"提示词": 提示词, "工作流": 工作流名}
                                    if 图片路径:
                                        操作参数["图片路径"] = 图片路径
                                    保存目录 = 当前文件夹 if 当前文件夹 else _获取默认保存目录()
                                    操作参数["保存目录"] = 保存目录
                                    网页请求处理器._comfyui执行中[nid] = "video"
                                    def _执行视频生成(nid, 操作实例, 操作参数, 保存目录):
                                        import time as _time_vid
                                        结果 = 操作实例.执行(操作参数)
                                        输出文本 = 结果.结果 if hasattr(结果, '结果') else str(结果)
                                        视频路径 = ""
                                        if hasattr(结果, '元数据') and 结果.元数据:
                                            sd = 结果.元数据.get("保存目录", "")
                                            if sd:
                                                import os as _os_vid
                                                视频文件 = [f for f in _os_vid.listdir(sd) if f.endswith(('.mp4','.webm','.gif','.avi'))]
                                                if 视频文件:
                                                    视频文件.sort(key=lambda f: _os_vid.getmtime(_os_vid.join(sd, f)), reverse=True)
                                                    视频路径 = _os_vid.join(sd, 视频文件[0])
                                        网页请求处理器._comfyui图片队列.append({"节点id": nid, "图片路径": 视频路径, "文件名": "", "状态": "完成" if 结果.成功 else "失败", "时间": _time_vid.time()})
                                        网页请求处理器._comfyui执行中.pop(nid, None)
                                    import threading as _th_vid
                                    _th_vid.Thread(target=_执行视频生成, args=(nid, 操作实例, 操作参数, 保存目录), daemon=True).start()
                                    节点输出[nid] = f"✅ 已提交视频生成(不等待)\n工作流: {工作流名}\n提示词: {提示词[:60]}"
                                    _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": 提示词[:200], "输出": 节点输出[nid], "成功": True, "图片": ""})
                                    _wf写日志(nid, nname, ntype, "成功", 提示词[:200], 节点输出[nid], "", int((_time2.time()-_t0)*1000))
                                except Exception as e:
                                    网页请求处理器._comfyui执行中.pop(nid, None)
                                    错误 = f"视频生成失败: {e}"
                                    节点输出[nid] = 错误
                                    _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": 错误, "成功": False})
                                    _wf写日志(nid, nname, ntype, "失败", "", 错误, 错误, int((_time2.time()-_t0)*1000))

                    elif ntype == "cloud_image" or ntype == "云端出图":
                        上游输入 = _收集上游输入(nid)
                        提示词 = config.get("提示词") or config.get("prompt") or 上游输入 or ""
                        提示词 = _re_clean.sub(r'【来自[^】]*】\s*', '', 提示词)
                        提示词 = _re_clean.sub(r'```[a-z]*\n?', '', 提示词)
                        提示词 = 提示词.strip()
                        服务商 = config.get("服务商") or config.get("provider") or "agnes"
                        宽度 = int(config.get("宽度") or config.get("width") or 1024)
                        高度 = int(config.get("高度") or config.get("height") or 1024)
                        if not 提示词:
                            节点输出[nid] = "无提示词输入"
                            _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": "无提示词", "成功": False})
                            _wf写日志(nid, nname, ntype, "失败", "", "无提示词", "缺少提示词", int((_time2.time()-_t0)*1000))
                        else:
                            try:
                                import os as _os_ci, json as _json_ci, urllib.request as _ur_ci, time as _time_ci, uuid as _uuid_ci, base64 as _b64_ci
                                _cfg路径 = _os_ci.path.join(_os_ci.path.dirname(_os_ci.path.dirname(_os_ci.path.dirname(_os_ci.path.abspath(__file__)))), "公共区", "配置", "系统配置.json")
                                with open(_cfg路径, "r", encoding="utf-8-sig") as _f:
                                    _系统配置 = _json_ci.load(_f)
                                _保存目录 = 当前文件夹 if 当前文件夹 else "."
                                _os_ci.makedirs(_保存目录, exist_ok=True)
                                _img_bytes = None
                                _文件名 = ""
                                if 服务商.lower() in ("agnes", "agnesai"):
                                    _api_key = _系统配置.get("Agnes_API密钥", "")
                                    if not _api_key:
                                        try:
                                            _密钥路径 = _os_ci.path.join(_os_ci.path.dirname(_os_ci.path.dirname(_os_ci.path.dirname(_os_ci.path.abspath(__file__)))), "隐私区", "我的配置", "密钥.json")
                                            if _os_ci.exists(_密钥路径):
                                                with open(_密钥路径, "r", encoding="utf-8-sig") as _kf:
                                                    _kd = _json_ci.load(_kf)
                                                _api_key = _kd.get("密钥列表", {}).get("AgnesAI(全模态免费)", {}).get("API密钥", "")
                                        except Exception: pass
                                    if not _api_key: raise Exception("未配置AgnesAI密钥，请在设置→模型中配置")
                                    _url = "https://apihub.agnes-ai.com/v1/images/generations"
                                    _payload = _json_ci.dumps({"model": "agnes-image-2.1-flash", "prompt": 提示词, "n": 1, "size": f"{宽度}x{高度}"}).encode("utf-8")
                                    _req = _ur_ci.Request(_url, data=_payload, headers={"Content-Type": "application/json", "Authorization": f"Bearer {_api_key}"}, method="POST")
                                    _SSE写入({"类型": "节点进度", "id": nid, "name": nname, "信息": "正在调用AgnesAI免费生图API..."})
                                    _resp = _ur_ci.urlopen(_req, timeout=120)
                                    _result = _json_ci.loads(_resp.read().decode("utf-8"))
                                    for _item in _result.get("data", []):
                                        if "b64_json" in _item: _img_bytes = _b64_ci.b64decode(_item["b64_json"]); break
                                        if "url" in _item: _img_bytes = _ur_ci.urlopen(_ur_ci.Request(_item["url"], headers={"User-Agent": "ZF3D-Agent"}), timeout=60).read(); break
                                    _文件名 = f"agnes_{_uuid_ci.uuid4().hex[:8]}.png"
                                elif 服务商.lower() in ("seedream", "字节跳动"):
                                    _api_key = _系统配置.get("Seedream_API密钥", "")
                                    if not _api_key: raise Exception("未配置Seedream_API密钥")
                                    _url = "https://ark.cn-beijing.volces.com/api/v3/images/generations"
                                    _payload = _json_ci.dumps({"model": "bytedance/seedream-v5.0-lite", "prompt": 提示词, "size": f"{宽度}x{高度}", "n": 1}).encode("utf-8")
                                    _req = _ur_ci.Request(_url, data=_payload, headers={"Content-Type": "application/json", "Authorization": f"Bearer {_api_key}"}, method="POST")
                                    _SSE写入({"类型": "节点进度", "id": nid, "name": nname, "信息": "正在调用Seedream API..."})
                                    _resp = _ur_ci.urlopen(_req, timeout=120)
                                    _result = _json_ci.loads(_resp.read().decode("utf-8"))
                                    for _item in _result.get("data", []):
                                        if "b64_json" in _item: _img_bytes = _b64_ci.b64decode(_item["b64_json"]); break
                                        if "url" in _item: _img_bytes = _ur_ci.urlopen(_ur_ci.Request(_item["url"], headers={"User-Agent": "ZF3D-Agent"}), timeout=60).read(); break
                                    _文件名 = f"seedream_{_uuid_ci.uuid4().hex[:8]}.png"
                                elif 服务商.lower() in ("nano_banana", "google", "banana"):
                                    _api_key = _系统配置.get("Google_API密钥", "")
                                    if not _api_key: raise Exception("未配置Google_API密钥")
                                    _ar = "9:16" if 高度 > 宽度 else ("16:9" if 宽度 > 高度 else "1:1")
                                    _url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-image:generateContent?key={_api_key}"
                                    _payload = _json_ci.dumps({"contents": [{"role": "user", "parts": [{"text": 提示词}]}], "generationConfig": {"responseModalities": ["IMAGE"], "imageConfig": {"aspectRatio": _ar}}}).encode("utf-8")
                                    _req = _ur_ci.Request(_url, data=_payload, headers={"Content-Type": "application/json"}, method="POST")
                                    _SSE写入({"类型": "节点进度", "id": nid, "name": nname, "信息": "正在调用Nano Banana API..."})
                                    _resp = _ur_ci.urlopen(_req, timeout=180)
                                    _result = _json_ci.loads(_resp.read().decode("utf-8"))
                                    for _c in _result.get("candidates", []):
                                        for _p in _c.get("content", {}).get("parts", []):
                                            if "inlineData" in _p: _img_bytes = _b64_ci.b64decode(_p["inlineData"]["data"]); break
                                        if _img_bytes: break
                                    _文件名 = f"nano_banana_{_uuid_ci.uuid4().hex[:8]}.png"
                                elif 服务商.lower() in ("grok", "xai"):
                                    _api_key = _系统配置.get("Grok_API密钥", "")
                                    if not _api_key: raise Exception("未配置Grok_API密钥")
                                    _size = "1024x1024"
                                    if 宽度 > 高度: _size = "1280x720"
                                    elif 高度 > 宽度: _size = "720x1280"
                                    _url = "https://api.x.ai/v1/images/generations"
                                    _payload = _json_ci.dumps({"model": "grok-2-image-1212", "prompt": 提示词, "n": 1, "size": _size, "response_format": "b64_json"}).encode("utf-8")
                                    _req = _ur_ci.Request(_url, data=_payload, headers={"Content-Type": "application/json", "Authorization": f"Bearer {_api_key}"}, method="POST")
                                    _SSE写入({"类型": "节点进度", "id": nid, "name": nname, "信息": "正在调用Grok API..."})
                                    _resp = _ur_ci.urlopen(_req, timeout=120)
                                    _result = _json_ci.loads(_resp.read().decode("utf-8"))
                                    for _item in _result.get("data", []):
                                        if "b64_json" in _item: _img_bytes = _b64_ci.b64decode(_item["b64_json"]); break
                                    _文件名 = f"grok_{_uuid_ci.uuid4().hex[:8]}.png"
                                elif 服务商.lower() in ("gpt_image", "openai", "gpt-image"):
                                    _api_key = _系统配置.get("OpenAI_Image_API密钥", "")
                                    if not _api_key: raise Exception("未配置OpenAI_Image_API密钥")
                                    _size = "1024x1024"
                                    if 宽度 > 高度: _size = "1536x1024"
                                    elif 高度 > 宽度: _size = "1024x1536"
                                    _url = "https://api.openai.com/v1/images/generations"
                                    _payload = _json_ci.dumps({"model": "gpt-image-1", "prompt": 提示词, "size": _size, "n": 1}).encode("utf-8")
                                    _req = _ur_ci.Request(_url, data=_payload, headers={"Content-Type": "application/json", "Authorization": f"Bearer {_api_key}"}, method="POST")
                                    _SSE写入({"类型": "节点进度", "id": nid, "name": nname, "信息": "正在调用GPT Image API..."})
                                    _resp = _ur_ci.urlopen(_req, timeout=180)
                                    _result = _json_ci.loads(_resp.read().decode("utf-8"))
                                    for _item in _result.get("data", []):
                                        if "b64_json" in _item: _img_bytes = _b64_ci.b64decode(_item["b64_json"]); break
                                    _文件名 = f"gpt_image_{_uuid_ci.uuid4().hex[:8]}.png"
                                else:
                                    raise Exception(f"未知服务商: {服务商}")
                                if not _img_bytes: raise Exception(f"API返回无图片: {str(_result)[:300]}")
                                _保存路径 = _os_ci.path.join(_保存目录, _文件名)
                                with open(_保存路径, "wb") as _f2: _f2.write(_img_bytes)
                                _输出文本 = f"✅ {服务商}出图成功\n提示词: {提示词[:60]}{'...' if len(提示词)>60 else ''}\n图片: {_保存路径}"
                                节点输出[nid] = _输出文本
                                _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": 提示词[:200], "输出": _输出文本, "成功": True, "图片": _保存路径})
                                if not hasattr(网页请求处理器, '_comfyui图片队列'): 网页请求处理器._comfyui图片队列 = []
                                网页请求处理器._comfyui图片队列.append({"节点id": nid, "图片路径": _保存路径, "文件名": _文件名, "状态": "完成", "时间": _time_ci.time()})
                                _wf写日志(nid, nname, ntype, "成功", 提示词[:200], _输出文本, "", int((_time2.time()-_t0)*1000))
                            except Exception as e:
                                错误 = f"云端出图失败: {e}"
                                节点输出[nid] = 错误
                                _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": 提示词[:200], "输出": 错误, "成功": False})
                                _wf写日志(nid, nname, ntype, "失败", 提示词[:200], 错误, 错误, int((_time2.time()-_t0)*1000))

                except Exception as _node_err:
                    print(f"  [WF] 节点执行异常: {nname}: {_node_err}")
                    节点输出[nid] = str(_node_err)
                    _SSE写入({"类型": "节点完成", "id": nid, "name": nname, "输入": "", "输出": f"执行异常: {_node_err}", "成功": False})

            # 逐层执行（同层并行）— 支持循环连线
            总节点数 = len(节点列表)
            已完成 = 0

            if 有循环:
                # ===== 循环执行模式 =====
                # 找到循环连线涉及的节点对
                循环节点对 = []
                for (f, t), conn in 循环连线.items():
                    loop = conn.get("loop", {})
                    循环节点对.append({"from": f, "to": t, "loop": loop})

                # 找到循环回路中的所有节点（from到to路径上的节点）
                循环节点集 = set()
                for pair in 循环节点对:
                    循环节点集.add(pair["from"])
                    循环节点集.add(pair["to"])

                # 先执行非循环节点（不在循环节点集中的节点）
                非循环层 = [lv for lv in range(len(层分组)) if not any(nid in 循环节点集 for nid in 层分组[lv])]
                for lv in 非循环层:
                    层 = 层分组[lv]
                    if not 层: continue
                    if len(层) == 1:
                        _执行节点(层[0])
                    else:
                        with _cf.ThreadPoolExecutor(max_workers=len(层)) as executor:
                            futures = {executor.submit(_执行节点, nid): nid for nid in 层}
                            for f in _cf.as_completed(futures):
                                f.result()
                    已完成 += len(层)
                    _SSE写入({"类型": "进度", "已完成": 已完成, "总数": 总节点数})

                # 执行循环
                for pair in 循环节点对:
                    loop = pair["loop"]
                    f, t = pair["from"], pair["to"]
                    fnode = 节点映射.get(f, {})
                    tnode = 节点映射.get(t, {})
                    fname = fnode.get("name", f)
                    tname = tnode.get("name", t)

                    if loop.get("type") == "for":
                        count = int(loop.get("count", 3))
                        _SSE写入({"类型": "循环轮次", "轮次": 0, "总数": count, "信息": f"固定循环 {count} 次"})
                        累积输出 = []
                        for i in range(count):
                            轮次 = i + 1
                            _SSE写入({"类型": "循环轮次", "轮次": 轮次, "总数": count, "信息": f"第{轮次}/{count}轮", "重置": [f, t]})
                            # 执行from节点
                            _执行节点(f)
                            累积输出.append(节点输出.get(f, ""))
                            # 执行to节点（带累积输入）
                            上游 = 累积输出[-1]
                            if len(累积输出) > 1:
                                上游 = "\n\n".join([f"第{j+1}轮：{v}" for j, v in enumerate(累积输出)])
                            # 临时设置from的输出为累积
                            原始输出 = 节点输出.get(f, "")
                            节点输出[f] = 上游
                            _执行节点(t)
                            节点输出[f] = 原始输出  # 恢复
                            已完成 += 2
                            _SSE写入({"类型": "进度", "已完成": min(已完成, 总节点数), "总数": 总节点数})
                        _SSE写入({"类型": "循环结束", "循环类型": "for", "总轮次": count, "信息": f"For循环完成，共执行{count}轮"})

                    elif loop.get("type") == "while":
                        condition = loop.get("condition", "")
                        maxLoop = int(loop.get("maxLoop", 10))
                        _SSE写入({"类型": "循环轮次", "轮次": 0, "总数": maxLoop, "信息": f"条件循环：{condition}（最多{maxLoop}次）"})
                        for i in range(maxLoop):
                            轮次 = i + 1
                            _SSE写入({"类型": "循环轮次", "轮次": 轮次, "总数": maxLoop, "信息": f"第{轮次}轮", "重置": [f, t]})
                            # 执行from节点
                            _执行节点(f)
                            # 执行to节点
                            _执行节点(t)
                            已完成 += 2
                            _SSE写入({"类型": "进度", "已完成": min(已完成, 总节点数), "总数": 总节点数})
                            # AI判断是否满足条件
                            判断输出 = 节点输出.get(t, "")
                            判断提示 = f"当前节点输出：\n{判断输出}\n\n判断条件：{condition}\n\n请回答：是或否（只回答一个字）"
                            判断结果 = self.模型直连器.发送消息(
                                [{"role": "user", "content": 判断提示}],
                                "你是条件判断助手，只回答'是'或'否'。",
                                跳过缓存=True
                            )
                            判断回复 = ""
                            if 判断结果.get("成功"):
                                判断回复 = 判断结果.get("回复内容", "").strip()
                            _SSE写入({"类型": "循环轮次", "轮次": 轮次, "信息": f"条件判断：{判断回复}"})
                            if "是" in 判断回复:
                                _SSE写入({"类型": "循环轮次", "轮次": 轮次, "信息": f"条件满足，跳出循环（第{轮次}轮）"})
                                break
                        else:
                            _SSE写入({"类型": "循环轮次", "轮次": maxLoop, "信息": f"已达最大循环次数{maxLoop}，强制停止"})
                        # While循环结束提示
                        if "是" in 判断回复:
                            _SSE写入({"类型": "循环结束", "循环类型": "while", "总轮次": 轮次, "信息": f"条件循环完成，第{轮次}轮满足条件「{condition}」"})
                        else:
                            _SSE写入({"类型": "循环结束", "循环类型": "while", "总轮次": maxLoop, "信息": f"条件循环结束，已达最大{maxLoop}次未满足「{condition}」"})

                # 执行剩余非循环节点
                循环后层 = [lv for lv in range(len(层分组)) if lv not in 非循环层]
                # 重新执行所有层（简化：从拓扑层中找到循环节点之后的层）
                已执行节点 = set()
                for lv in 非循环层:
                    已执行节点.update(层分组[lv])
                已执行节点.update(循环节点集)
                剩余节点 = [nid for nid in 执行顺序 if nid not in 已执行节点]
                if 剩余节点:
                    # 按层级排序剩余节点
                    剩余排序 = sorted(剩余节点, key=lambda nid: 层级.get(nid, 0))
                    # 重新分组
                    剩余层级 = {}
                    for nid in 剩余排序:
                        上游在剩余 = [p for p in 入边[nid] if p in set(剩余节点)]
                        if not 上游在剩余:
                            剩余层级[nid] = 0
                        else:
                            剩余层级[nid] = max(剩余层级.get(p, 0) for p in 上游在剩余) + 1
                    剩余最大层 = max(剩余层级.values()) if 剩余层级 else 0
                    剩余层分组 = [[] for _ in range(剩余最大层 + 1)]
                    for nid, lv in 剩余层级.items():
                        剩余层分组[lv].append(nid)
                    for 层 in 剩余层分组:
                        if not 层: continue
                        if len(层) == 1:
                            _执行节点(层[0])
                        else:
                            with _cf.ThreadPoolExecutor(max_workers=len(层)) as executor:
                                futures = {executor.submit(_执行节点, nid): nid for nid in 层}
                                for f2 in _cf.as_completed(futures):
                                    f2.result()
                        已完成 += len(层)
                        _SSE写入({"类型": "进度", "已完成": min(已完成, 总节点数), "总数": 总节点数})

            else:
                # ===== 普通执行模式（原逻辑） =====
                for 层 in 层分组:
                    if len(层) == 1:
                        _执行节点(层[0])
                    else:
                        with _cf.ThreadPoolExecutor(max_workers=len(层)) as executor:
                            futures = {executor.submit(_执行节点, nid): nid for nid in 层}
                            for f in _cf.as_completed(futures):
                                f.result()
                    已完成 += len(层)
                    _SSE写入({"类型": "进度", "已完成": 已完成, "总数": 总节点数})

            _SSE写入({"类型": "完成", "结果": {"成功": True, "总数": 总节点数, "会话ID": wf会话ID}})

        except Exception as e:
            if isinstance(e, (ConnectionAbortedError, ConnectionResetError, BrokenPipeError)):
                return
            try:
                self._返回JSON({"成功": False, "错误": f"工作流执行异常: {str(e)}"})
            except Exception:
                return

    def _处理员工管理API(self, 路径: str, 数据: dict):
        """处理员工管理相关API"""
        if not self.模块注册 or "员工管理" not in self.模块注册:
            self._返回JSON({"成功": False, "错误": "员工管理模块未加载"})
            return
        模块 = self.模块注册["员工管理"]
        if 路径 == "/api/employee-switch":
            结果 = 模块.切换员工(数据.get("姓名", ""))
            # 只记录切换，不修改主对话模块
            self._返回JSON(结果)
        elif 路径 == "/api/employee-create":
            self._返回JSON(模块.创建员工(数据.get("配置", {})))
        elif 路径 == "/api/employee-update":
            self._返回JSON(模块.更新员工(数据.get("姓名", ""), 数据.get("配置", {})))
        elif 路径 == "/api/employee-delete":
            self._返回JSON(模块.删除员工(数据.get("姓名", "")))
        elif 路径 == "/api/employee-assign":
            self._返回JSON(模块.分配员工(数据.get("员工名", ""), 数据.get("老板名", "")))
        elif 路径 == "/api/employee-unassign":
            self._返回JSON(模块.移除分配(数据.get("员工名", ""), 数据.get("老板名", "")))
        elif 路径 == "/api/employee-status":
            self._返回JSON(模块.设置状态(数据.get("姓名", ""), 数据.get("状态", "")))
        elif 路径 == "/api/employee-notify-config":
            姓名 = 数据.get("姓名", "")
            if 姓名 in 网页请求处理器._员工提醒队列:
                信息 = 网页请求处理器._员工提醒队列[姓名]
                if "间隔分钟" in 数据:
                    信息["间隔分钟"] = int(数据["间隔分钟"])
                if "弹窗时长秒" in 数据:
                    信息["弹窗时长秒"] = int(数据["弹窗时长秒"])
                self._返回JSON({"成功": True, "数据": f"提醒间隔{信息['间隔分钟']}分钟，弹窗{信息['弹窗时长秒']}秒"})
            else:
                self._返回JSON({"成功": False, "错误": "该员工未开启提醒"})
        elif 路径 == "/api/employee-generate-persona":
            """AI生成员工人设（非流式，直接返回JSON）"""
            姓名 = 数据.get("姓名", "")
            角色 = 数据.get("角色", "")
            if not 姓名 or not 角色:
                self._返回JSON({"成功": False, "错误": "缺少姓名或角色"})
                return
            if not self.模型直连器:
                self._返回JSON({"成功": False, "错误": "模型未就绪"})
                return
            # AI润色用户描述（补充细节，不改原意）
            润色提示 = f"用户想创建一个AI图片提示词加工流水线的员工「{姓名}」，职责描述：{角色}\n请对这段描述进行润色和补充，使其更清晰、更专业，但不要改变用户的原始意图。只输出润色后的描述（50字以内），不要其他内容。"
            润色结果 = self.模型直连器.发送消息([{"role": "user", "content": 润色提示}], "简洁专业地润色，保留原意。")
            润色描述 = ""
            if 润色结果.get("成功"):
                润色描述 = 润色结果.get("回复内容", "").strip()
            # 人设 = 名字 + 用户原始描述 + AI润色 + 工作规则
            人设 = f"你的名字叫{姓名}。你在AI提示词加工流水线中工作。\n\n## 用户描述\n{角色}\n\n## 角色润色\n{润色描述}\n\n## 工作规则\n收到上游提示词后，按照上述职责进行加工，输出完整提示词。不要解释工作过程，禁止输出设计文档。当用户问你是谁时，回答你是{姓名}。"
            self._返回JSON({"成功": True, "数据": 人设})
        elif 路径 == "/api/employee-clear-history":
            """清除员工聊天记录（不影响技能和经验）"""
            姓名 = 数据.get("姓名", "")
            # 清空内存
            网页请求处理器._员工对话历史[姓名] = []
            # 清空记忆文件
            if self.模块注册 and "员工管理" in self.模块注册:
                运行时结果 = self.模块注册["员工管理"].获取运行时配置(姓名)
                if 运行时结果.get("成功"):
                    记忆路径 = 运行时结果["数据"].get("记忆路径")
                    if 记忆路径:
                        try:
                            绝对路径 = self.配置加载器.项目根目录 / 记忆路径.lstrip("./")
                            if 绝对路径.exists():
                                绝对路径.write_text("[]", encoding="utf-8")
                        except Exception:
                            pass
            self._返回JSON({"成功": True, "数据": "聊天记录已清除"})
        elif 路径 == "/api/employee-save-history":
            """保存员工聊天历史（单条删除后同步）"""
            姓名 = 数据.get("姓名", "")
            历史 = 数据.get("历史", [])
            网页请求处理器._员工对话历史[姓名] = 历史
            # 同步到记忆文件
            if self.模块注册 and "员工管理" in self.模块注册:
                运行时结果 = self.模块注册["员工管理"].获取运行时配置(姓名)
                if 运行时结果.get("成功"):
                    记忆路径 = 运行时结果["数据"].get("记忆路径")
                    if 记忆路径:
                        try:
                            import json as _json
                            绝对路径 = self.配置加载器.项目根目录 / 记忆路径.lstrip("./")
                            绝对路径.parent.mkdir(parents=True, exist_ok=True)
                            记忆数据 = []
                            i = 0
                            while i < len(历史):
                                role = 历史[i].get("role", "")
                                content = 历史[i].get("content", "")
                                if role == "user" and i + 1 < len(历史) and 历史[i + 1].get("role") == "assistant":
                                    记忆数据.append({"用户": content, "助手": 历史[i + 1].get("content", "")})
                                    i += 2
                                elif role == "user":
                                    记忆数据.append({"用户": content, "助手": ""})
                                    i += 1
                                elif role == "assistant":
                                    记忆数据.append({"用户": "", "助手": content})
                                    i += 1
                                else:
                                    i += 1
                            绝对路径.write_text(_json.dumps(记忆数据, ensure_ascii=False, indent=2), encoding="utf-8")
                        except Exception:
                            pass
            self._返回JSON({"成功": True})
        else:
            self._返回JSON({"成功": False, "错误": "未知员工管理API: " + 路径})

    def _引擎差异分析(self) -> dict:
        """对比工作引擎与主引擎的文件差异"""
        import hashlib
        项目根 = self.配置加载器.项目根目录
        主引擎路径 = 项目根 / "公共区"
        工作引擎路径 = 项目根 / "隐私区" / "我的工作引擎" / "公共区"
        忽略模式 = ["__pycache__", ".pyc", ".git", ".log", ".bak", "__pycache__"]
        最大差异 = 50

        if not 工作引擎路径.exists():
            return {"成功": True, "新增": [], "修改": [], "删除": [], "未变": 0, "提示": "工作引擎目录为空或不存在"}

        def _忽略(路径: str) -> bool:
            return any(p in 路径 for p in 忽略模式)

        def _哈希(文件路径) -> str:
            h = hashlib.md5()
            with open(文件路径, "rb") as f:
                for chunk in iter(lambda: f.read(8192), b""):
                    h.update(chunk)
            return h.hexdigest()

        # 收集主引擎文件
        主文件 = {}
        for f in 主引擎路径.rglob("*"):
            if f.is_file() and not _忽略(str(f)):
                相对 = f.relative_to(主引擎路径).as_posix()
                主文件[相对] = _哈希(f)

        # 收集工作引擎文件
        工作文件 = {}
        for f in 工作引擎路径.rglob("*"):
            if f.is_file() and not _忽略(str(f)):
                相对 = f.relative_to(工作引擎路径).as_posix()
                工作文件[相对] = _哈希(f)

        新增 = [k for k in 工作文件 if k not in 主文件]
        删除 = [k for k in 主文件 if k not in 工作文件]
        修改 = [k for k in 工作文件 if k in 主文件 and 工作文件[k] != 主文件[k]]
        未变 = sum(1 for k in 工作文件 if k in 主文件 and 工作文件[k] == 主文件[k])

        # 截断
        新增 = sorted(新增)[:最大差异]
        删除 = sorted(删除)[:最大差异]
        修改 = sorted(修改)[:最大差异]

        return {"成功": True, "新增": 新增, "修改": 修改, "删除": 删除, "未变": 未变}

    def _执行引擎合并(self, 数据: dict) -> dict:
        """执行文件合并：备份→检测→复制→记日志"""
        import shutil
        import py_compile
        项目根 = self.配置加载器.项目根目录
        主引擎路径 = 项目根 / "公共区"
        工作引擎路径 = 项目根 / "隐私区" / "我的工作引擎" / "公共区"
        文件列表 = 数据.get("文件列表", [])
        执行合并 = 数据.get("执行", False)

        if not 文件列表:
            return {"成功": False, "错误": "未选择文件"}

        if not 工作引擎路径.exists():
            return {"成功": False, "错误": "工作引擎目录不存在"}

        # 检测阶段
        检测结果 = []
        全部通过 = True
        for 相对路径 in 文件列表:
            源文件 = 工作引擎路径 / 相对路径
            if not 源文件.exists():
                检测结果.append({"路径": 相对路径, "状态": "源文件不存在"})
                全部通过 = False
                continue
            # 语法检查
            if 相对路径.endswith(".py"):
                try:
                    py_compile.compile(str(源文件), doraise=True)
                    检测结果.append({"路径": 相对路径, "状态": "通过"})
                except py_compile.PyCompileError as e:
                    检测结果.append({"路径": 相对路径, "状态": f"语法错误: {str(e)[:200]}"})
                    全部通过 = False
            elif 相对路径.endswith(".json"):
                try:
                    with open(源文件, "r", encoding="utf-8") as f:
                        json.load(f)
                    检测结果.append({"路径": 相对路径, "状态": "通过"})
                except Exception as e:
                    检测结果.append({"路径": 相对路径, "状态": f"JSON错误: {e}"})
                    全部通过 = False
            else:
                检测结果.append({"路径": 相对路径, "状态": "通过"})

        if not 全部通过:
            return {"成功": False, "错误": "检测未通过，已阻止合并", "检测结果": 检测结果}

        if not 执行合并:
            return {"成功": True, "已检测": True, "检测结果": 检测结果, "消息": "检测通过，可执行合并"}

        # 执行合并：先备份
        from datetime import datetime
        时间戳 = datetime.now().strftime("%Y%m%d_%H%M%S")
        备份目录 = 项目根 / "引擎管理" / "备份" / 时间戳
        备份目录.mkdir(parents=True, exist_ok=True)
        for 相对路径 in 文件列表:
            源文件 = 主引擎路径 / 相对路径
            if 源文件.exists():
                备份目标 = 备份目录 / 相对路径
                备份目标.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(源文件, 备份目标)

        # 复制工作引擎→主引擎
        合并数 = 0
        for 相对路径 in 文件列表:
            源文件 = 工作引擎路径 / 相对路径
            目标文件 = 主引擎路径 / 相对路径
            目标文件.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(源文件, 目标文件)
            合并数 += 1

        # 清理旧备份（保留10个）
        备份根 = 项目根 / "引擎管理" / "备份"
        if 备份根.exists():
            备份列表 = sorted(备份根.iterdir())
            while len(备份列表) > 10:
                最旧 = 备份列表.pop(0)
                shutil.rmtree(最旧, ignore_errors=True)

        # 记录合并日志
        合并日志路径 = 项目根 / "引擎管理" / "合并日志.json"
        try:
            with open(合并日志路径, "r", encoding="utf-8") as f:
                合并日志 = json.load(f)
        except Exception:
            合并日志 = {"记录": []}
        合并日志["记录"].append({
            "时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "方向": "工作引擎→主引擎",
            "变更摘要": f"合并 {合并数} 个文件: {', '.join(文件列表[:5])}{'...' if len(文件列表) > 5 else ''}",
            "备份": 时间戳
        })
        with open(合并日志路径, "w", encoding="utf-8") as f:
            json.dump(合并日志, f, ensure_ascii=False, indent=2)

        # 在工作引擎git中打标签，标记此commit已合并到主引擎
        启动器 = getattr(self, '_启动器实例', None)
        if 启动器 and hasattr(启动器, '进化引擎'):
            try:
                启动器.进化引擎._git打标签(f"已合并_{时间戳}")
            except Exception:
                pass  # 标签失败不影响合并结果

        return {"成功": True, "合并数": 合并数, "备份": 时间戳, "检测结果": 检测结果}

    def _执行引擎回滚(self, 备份名: str) -> dict:
        """从指定备份恢复主引擎文件"""
        import shutil
        项目根 = self.配置加载器.项目根目录
        备份目录 = 项目根 / "引擎管理" / "备份" / 备份名
        主引擎路径 = 项目根 / "公共区"

        if not 备份目录.exists():
            return {"成功": False, "错误": f"备份不存在: {备份名}"}

        恢复数 = 0
        for f in 备份目录.rglob("*"):
            if f.is_file():
                相对 = f.relative_to(备份目录)
                目标 = 主引擎路径 / 相对
                目标.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(f, 目标)
                恢复数 += 1

        # 记录日志
        from datetime import datetime
        合并日志路径 = 项目根 / "引擎管理" / "合并日志.json"
        try:
            with open(合并日志路径, "r", encoding="utf-8") as f:
                合并日志 = json.load(f)
        except Exception:
            合并日志 = {"记录": []}
        合并日志["记录"].append({
            "时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "方向": "回滚",
            "变更摘要": f"从备份 [{备份名}] 恢复 {恢复数} 个文件"
        })
        with open(合并日志路径, "w", encoding="utf-8") as f:
            json.dump(合并日志, f, ensure_ascii=False, indent=2)

        return {"成功": True, "恢复数": 恢复数, "备份": 备份名}

    def _处理对话(self, 数据: dict):
        try:
            消息 = 数据.get("消息", "")
            上下文 = 数据.get("上下文", {})

            # 构建文件上下文注入文本
            文件上下文提示 = self._构建文件上下文提示(上下文)

            # 注入当前工作目录到操作注册中心（供操作类作为默认保存路径）
            当前文件夹 = 上下文.get("当前文件夹", "") if 上下文 else ""
            if 当前文件夹 and self.操作注册中心:
                self.操作注册中心.设置当前工作目录(当前文件夹)

            if self.模块注册 and "对话" in self.模块注册:
                对话模块 = self.模块注册["对话"]
                # 重置取消标志（防止上次取消后新建对话仍处于取消状态）
                对话模块._取消标志 = False
                # 将文件上下文注入对话模块
                if 文件上下文提示:
                    对话模块.文件上下文 = 文件上下文提示
                else:
                    对话模块.文件上下文 = ""

                # SSE流式响应
                self.send_response(200)
                self.send_header("Content-Type", "text/event-stream; charset=utf-8")
                self.send_header("Cache-Control", "no-cache")
                self.send_header("Connection", "keep-alive")
                self.send_header("Access-Control-Allow-Origin", "http://localhost:8765")
                self.end_headers()

                def _SSE写入(事件数据):
                    try:
                        行 = f"data: {json.dumps(事件数据, ensure_ascii=False)}\n\n"
                        self.wfile.write(行.encode("utf-8"))
                        self.wfile.flush()
                    except Exception:
                        pass  # SSE流已关闭（后台下载线程仍在推送进度），静默忽略

                # 拦截推理流推送，实时写SSE
                原始推入 = 对话模块._推入推理流
                def _SSE推入(类型, 内容):
                    原始推入(类型, 内容)
                    _SSE写入({"类型": "推理流", "记录": [{"类型": 类型, "内容": 内容}]})
                对话模块._推入推理流 = _SSE推入
                # 同步patch操作注册中心的进度回调，使操作（如询问用户）也走SSE
                if self.操作注册中心:
                    self.操作注册中心.设置进度回调(_SSE推入)

                结果 = 对话模块.运行({"消息": 消息})

                # 恢复原始方法
                对话模块._推入推理流 = 原始推入
                if self.操作注册中心:
                    self.操作注册中心.设置进度回调(原始推入)

                # 发送最终结果
                _SSE写入({"类型": "完成", "结果": 结果})
                return
            elif self.模型直连器:
                消息列表 = [{"role": "user", "content": 消息}]
                系统提示词 = 数据.get("系统提示词", "")
                if 文件上下文提示:
                    系统提示词 = (系统提示词 + "\n\n" if 系统提示词 else "") + 文件上下文提示
                结果 = self.模型直连器.发送消息(消息列表, 系统提示词)
                if 结果["成功"]:
                    self._返回JSON({"成功": True, "回复": 结果["回复内容"]})
                else:
                    self._返回JSON({"成功": False, "错误": 结果.get("错误", "调用失败")})
            else:
                self._返回JSON({"成功": False, "错误": "无可用的模型或对话模块"})
        except Exception as e:
            if isinstance(e, (ConnectionAbortedError, ConnectionResetError, BrokenPipeError)):
                return  # 客户端已断开（用户点停止），无需处理
            import traceback
            traceback.print_exc()
            if self.运行诊断器:
                self.运行诊断器.记录错误("网页服务._处理对话", e)
            # 异常也写入对话历史，确保不丢失
            if self.模块注册 and "对话" in self.模块注册:
                try:
                    对话模块 = self.模块注册["对话"]
                    错误信息 = f"❌ 对话处理异常: {type(e).__name__}: {str(e)[:300]}"
                    对话模块.对话历史.append({"角色": "助手", "内容": 错误信息, "时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S")})
                    对话模块._保存当前对话()
                except Exception:
                    pass
            # SSE模式下通过事件发送错误（header已发送，不能再返回JSON）
            if self.模块注册 and "对话" in self.模块注册:
                try:
                    错误结果 = {"成功": False, "错误": f"对话处理遇到问题: {str(e)[:200]}", "回复": f"❌ 对话处理遇到问题，请稍后重试"}
                    错误行 = f"data: {json.dumps({'类型': '完成', '结果': 错误结果}, ensure_ascii=False)}\n\n"
                    self.wfile.write(错误行.encode("utf-8"))
                    self.wfile.flush()
                except Exception:
                    pass
            else:
                self._返回JSON({"成功": False, "错误": f"对话处理遇到问题，请稍后重试: {str(e)[:200]}"})

    def _构建文件上下文提示(self, 上下文: dict) -> str:
        """根据前端传来的上下文构建系统提示词注入"""
        if not 上下文:
            return ""
        部分 = ["\n\n## 当前工作环境\n"]
        # 当前文件夹
        当前文件夹 = 上下文.get("当前文件夹", "")
        if 当前文件夹:
            # 读取目录树摘要（浅层）
            try:
                树结果 = self.文件管理器.目录树(当前文件夹, 1)
                if 树结果.get("成功") and 树结果.get("树", {}).get("子项"):
                    子项列表 = 树结果["树"]["子项"]
                    文件列表 = [f["名称"] for f in 子项列表 if f["类型"] == "文件"]
                    目录列表 = [d["名称"] for d in 子项列表 if d["类型"] == "目录"]
                    部分.append(f"当前打开的文件夹: {当前文件夹}")
                    if 目录列表:
                        部分.append(f"  子目录: {', '.join(目录列表[:30])}")
                    if 文件列表:
                        部分.append(f"  文件: {', '.join(文件列表[:30])}")
            except:
                部分.append(f"当前打开的文件夹: {当前文件夹}")
        # 打开的文件列表
        打开的文件 = 上下文.get("打开的文件列表", [])
        if 打开的文件:
            部分.append(f"已打开的文件: {', '.join(f['名称'] for f in 打开的文件)}")
        # 当前正在编辑的文件
        当前文件 = 上下文.get("当前文件")
        if 当前文件:
            内容预览 = (当前文件.get("内容") or "")[:3000]
            部分.append(f"\n当前正在编辑的文件: {当前文件.get('名称')} ({当前文件.get('路径')})")
            部分.append(f"文件内容:\n```\n{内容预览}\n{'...(已截断)' if len(当前文件.get('内容') or '') > 3000 else ''}\n```")
        # 选中文件/文件夹
        选中文件 = 上下文.get("选中文件", [])
        if 选中文件:
            文件数 = sum(1 for f in 选中文件 if f.get("类型") != "目录")
            文件夹数 = sum(1 for f in 选中文件 if f.get("类型") == "目录")
            部分.append(f"\n## 📋 用户已选中以下文件/文件夹（共{len(选中文件)}项：{文件数}个文件，{文件夹数}个文件夹）:")
            for 项 in 选中文件:
                类型标签 = "📁" if 项.get("类型") == "目录" else "📄"
                部分.append(f"- {类型标签} {项.get('名称', '')} ({项.get('路径', '')})")
            部分.append("用户已在资源管理器中选中上述文件/文件夹，请在回复时知晓这些被选中的内容。用户可能会要求你对这些文件执行批量操作。")
        # 框选文本
        框选 = 上下文.get("框选文本")
        if 框选 and 框选.get("内容"):
            框选原文 = 框选['内容']
            框选文件路径 = 框选.get('所在文件', '')
            框选文件名 = 框选.get('所在文件名', '')
            旧文本JSON = json.dumps(框选原文, ensure_ascii=False)
            # 判断是否为Word文档
            后缀 = Path(框选文件名).suffix.lower() if 框选文件名 else ""
            is_word = 后缀 == ".docx"
            is_excel = 后缀 in (".xlsx", ".xls")
            if is_word:
                操作名 = "替换Word文本"
                部分.append(f"\n## ⚠️ 框选文本操作（最高优先级）")
                部分.append(f"用户已在Word文档「{框选文件名}」中框选文本，你【必须】使用「替换Word文本」操作修改.docx文件！")
                部分.append(f"【禁止】使用「替换文本」操作（那是用于纯文本文件的，无法操作.docx）")
                部分.append(f"【禁止】使用「读取文件」（你已能看到原文）")
                部分.append(f"【禁止】回复文字让用户确认，直接执行替换即可")
                部分.append(f"\n框选原文（必须原封不动作为旧文本参数）:")
                部分.append(f"```\n{框选原文}\n```")
                部分.append(f"文件路径: {框选文件路径}")
                部分.append(f"\n直接输出以下JSON即可（只需修改新文本字段）:")
                部分.append(f'{{"思考": "根据用户指令处理Word文档中的框选文本", "操作": "替换Word文本", "参数": {{"路径": "{框选文件路径}", "旧文本": {旧文本JSON}, "新文本": "修改后的内容"}}}}')
                部分.append(f"\n例如删除就写: \"新文本\": \"\"")
                部分.append(f"例如改写就写: \"新文本\": \"改写后的文本\"")
            elif is_excel:
                操作名 = "替换Excel文本"
                部分.append(f"\n## ⚠️ 框选文本操作（最高优先级）")
                部分.append(f"用户已在Excel文档「{框选文件名}」中框选文本，你【必须】使用「替换Excel文本」操作修改.xlsx文件！")
                部分.append(f"【禁止】使用「替换文本」操作（那是用于纯文本文件的，无法操作.xlsx）")
                部分.append(f"【禁止】使用「读取文件」（你已能看到原文）")
                部分.append(f"【禁止】回复文字让用户确认，直接执行替换即可")
                部分.append(f"\n框选原文（必须原封不动作为旧文本参数）:")
                部分.append(f"```\n{框选原文}\n```")
                部分.append(f"文件路径: {框选文件路径}")
                部分.append(f"\n直接输出以下JSON即可（只需修改新文本字段）:")
                部分.append(f'{{"思考": "根据用户指令处理Excel文档中的框选文本", "操作": "替换Excel文本", "参数": {{"路径": "{框选文件路径}", "旧文本": {旧文本JSON}, "新文本": "修改后的内容"}}}}')
                部分.append(f"\n例如删除就写: \"新文本\": \"\"")
                部分.append(f"例如改写就写: \"新文本\": \"改写后的文本\"")
            else:
                操作名 = "替换文本"
                部分.append(f"\n## ⚠️ 框选文本操作（最高优先级）")
                部分.append(f"用户已在文件「{框选文件名}」中框选文本，你【必须】使用「替换文本」操作！")
                部分.append(f"【禁止】使用读取文件（你已能看到原文）")
                部分.append(f"【禁止】使用写入文件重写整个文件")
                部分.append(f"【禁止】回复文字让用户确认，直接执行替换即可")
                部分.append(f"\n框选原文（必须原封不动作为旧文本参数）:")
                部分.append(f"```\n{框选原文}\n```")
                部分.append(f"文件路径: {框选文件路径}")
                部分.append(f"\n直接输出以下JSON即可（只需修改新文本字段）:")
                部分.append(f'{{"思考": "根据用户指令处理框选文本", "操作": "替换文本", "参数": {{"路径": "{框选文件路径}", "旧文本": {旧文本JSON}, "新文本": "修改后的内容"}}}}')
                部分.append(f"\n例如删除就写: \"新文本\": \"\"")
                部分.append(f"例如改写就写: \"新文本\": \"改写后的文本\"")
        # 文件操作提示
        部分.append("\n你可以通过以下操作来操作文件（在回复中使用JSON格式调用）:")
        if 框选 and 框选.get("内容"):
            # 有框选时不再重复列出其他操作，避免模型分心
            pass
        else:
            部分.append("""```json
{"思考": "分析需求", "操作": "读取文件", "参数": {"路径": "文件路径"}}
{"思考": "分析需求", "操作": "写入文件", "参数": {"路径": "文件路径", "内容": "文件内容"}}
{"思考": "分析需求", "操作": "创建文件", "参数": {"路径": "文件路径", "内容": "初始内容"}}
{"思考": "分析需求", "操作": "追加文件", "参数": {"路径": "文件路径", "内容": "追加内容"}}
{"思考": "分析需求", "操作": "列出目录", "参数": {"路径": "目录路径"}}
{"思考": "分析需求", "操作": "删除文件", "参数": {"路径": "文件路径"}}
{"思考": "分析需求", "操作": "替换文本", "参数": {"路径": "文件路径", "旧文本": "被替换的原文", "新文本": "替换后的新文本"}}
```""")
        return "\n".join(部分)

    # ===== 开发模式热重载 =====
    _dev文件时间戳 = {}  # {文件路径: mtime}

    def _处理开发热重载SSE(self):
        """SSE端点：监控界面目录文件变化，变化时推送reload事件"""
        self.send_response(200)
        self.send_header("Content-Type", "text/event-stream; charset=utf-8")
        self.send_header("Cache-Control", "no-cache")
        self.send_header("Access-Control-Allow-Origin", "http://localhost:8765")
        self.end_headers()
        try:
            # 初始握手
            self.wfile.write(b"data: {\"type\":\"connected\"}\n\n")
            self.wfile.flush()
            界面目录 = self.界面目录
            监控后缀 = {".js", ".css", ".html"}
            检查间隔 = 1.0
            for _ in range(36000):  # 最多10小时
                time.sleep(检查间隔)
                变更文件 = []
                try:
                    for f in 界面目录.rglob("*"):
                        if f.is_file() and f.suffix.lower() in 监控后缀:
                            try:
                                mtime = f.stat().st_mtime
                            except Exception:
                                continue
                            key = str(f)
                            旧时间 = 网页请求处理器._dev文件时间戳.get(key)
                            if 旧时间 is None:
                                网页请求处理器._dev文件时间戳[key] = mtime
                            elif mtime != 旧时间:
                                网页请求处理器._dev文件时间戳[key] = mtime
                                变更文件.append(f.name)
                except Exception:
                    pass
                if 变更文件:
                    消息 = json.dumps({"type": "reload", "files": 变更文件}, ensure_ascii=False)
                    self.wfile.write(f"data: {消息}\n\n".encode("utf-8"))
                    self.wfile.flush()
                    break  # 推送一次reload后断开，前端刷新后重新连接
        except (ConnectionAbortedError, ConnectionResetError, BrokenPipeError, BrokenPipeError):
            pass
        except Exception:
            pass

    def _返回文件(self, 路径: Path, 类型: str, 查询串: str = ""):
        try:
            if 路径.exists():
                self.send_response(200)
                self.send_header("Content-Type", f"{类型}; charset=utf-8")
                self.send_header("Access-Control-Allow-Origin", "http://localhost:8765")
                # 开发模式：所有文件no-cache，浏览器刷新即获取最新
                _开发模式 = False
                try:
                    _开发模式 = self.配置加载器.配置缓存.get("系统配置", {}).get("开发模式", False)
                except Exception:
                    pass
                if not _开发模式 and ("?v=" in 查询串 or "&v=" in 查询串):
                    self.send_header("Cache-Control", "max-age=86400")  # 缓存1天
                else:
                    self.send_header("Cache-Control", "no-cache")
                self.end_headers()
                with open(路径, "rb") as f:
                    self.wfile.write(f.read())
            else:
                self.send_response(404)
                self.send_header("Content-Type", "text/plain; charset=utf-8")
                self.end_headers()
                self.wfile.write("file not found".encode("utf-8"))
        except (ConnectionAbortedError, ConnectionResetError, BrokenPipeError):
            pass  # 客户端已断开，无需处理

    def _返回JSON(self, 数据: dict, 状态码: int = 200):
        try:
            响应体 = json.dumps(数据, ensure_ascii=False).encode("utf-8")
            self.send_response(状态码)
            self.send_header("Content-Type", "application/json; charset=utf-8")
            self.send_header("Content-Length", str(len(响应体)))
            self.send_header("Access-Control-Allow-Origin", "http://localhost:8765")
            self.end_headers()
            self.wfile.write(响应体)
        except (ConnectionAbortedError, ConnectionResetError, BrokenPipeError):
            pass  # 客户端已断开（如用户点停止），无需记录

    def _打开文件夹选择对话框(self) -> str:
        """用tkinter弹出Windows原生文件夹选择对话框"""
        try:
            import tkinter as tk
            from tkinter import filedialog
            root = tk.Tk()
            root.withdraw()
            root.attributes("-topmost", True)
            result = filedialog.askdirectory(title="选择文件夹")
            root.destroy()
            return result or ""
        except Exception as e:
            print(f"⚠ 文件夹选择对话框失败: {e}")
            return ""

    def _猜测类型(self, 路径: str) -> str:
        后缀映射 = {
            ".html": "text/html", ".css": "text/css", ".js": "application/javascript",
            ".json": "application/json", ".png": "image/png", ".jpg": "image/jpeg",
            ".svg": "image/svg+xml", ".ico": "image/x-icon"
        }
        后缀 = Path(路径).suffix.lower()
        return 后缀映射.get(后缀, "application/octet-stream")

    def log_message(self, format, *args):
        """输出简要HTTP请求日志"""
        print(f"  [{self.log_date_time_string()}] {args[0] if args else ''}")

    # ============ 股票数据接口 ============

    def _东财请求(self, url, headers=None):
        """东方财富API请求（带重试）"""
        import urllib.request
        默认headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
            "Referer": "https://quote.eastmoney.com/"
        }
        if headers:
            默认headers.update(headers)
        for 尝试 in range(3):
            try:
                req = urllib.request.Request(url, headers=默认headers)
                resp = urllib.request.urlopen(req, timeout=8)
                return json.loads(resp.read().decode("utf-8"))
            except Exception:
                if 尝试 >= 2:
                    raise
                import time; time.sleep(0.3)

    # 缓存最新交易日(避免每次查MAX)
    _最新交易日缓存 = None
    _最新交易日时间 = 0

    def _获取最新交易日(self):
        """获取最新交易日(60秒内缓存)"""
        import time as _time
        now = _time.time()
        if self._最新交易日缓存 and (now - self._最新交易日时间) < 60:
            return self._最新交易日缓存
        try:
            import sqlite3 as _sqlite3
            db路径 = str(Path(__file__).parent.parent.parent / "隐私区" / "我的数据" / "股票缓存.db")
            conn = _sqlite3.connect(db路径, timeout=5)
            c = conn.cursor()
            c.execute("SELECT MAX(日期) FROM K线数据 WHERE 周期='daily'")
            row = c.fetchone()
            conn.close()
            if row and row[0]:
                self._最新交易日缓存 = row[0]
                self._最新交易日时间 = now
                return row[0]
        except:
            pass
        return None

    def _获取股票盘面(self, 页码: int = 1, 排序字段: str = "f3", 排序方向: str = "desc") -> dict:
        """获取盘面数据 — 本地数据库SQL原生分页+排序, 一次连接完成所有查询"""
        import sqlite3 as _sqlite3
        from datetime import datetime
        try:
            db路径 = str(Path(__file__).parent.parent.parent / "隐私区" / "我的数据" / "股票缓存.db")
            conn = _sqlite3.connect(db路径, timeout=10)
            conn.row_factory = _sqlite3.Row
            c = conn.cursor()

            # 1. 最新交易日(用缓存)
            最新日 = self._获取最新交易日()
            if not 最新日:
                c.execute("SELECT MAX(日期) as d FROM K线数据 WHERE 周期='daily'")
                r = c.fetchone()
                最新日 = r["d"] if r else None
            if not 最新日:
                conn.close()
                return {"成功": False, "错误": "数据库无K线数据"}

            # 2. 排序+分页(一次SQL取20行)
            排序列 = {
                "f3": "(收-开)/开",   # 涨幅
                "f6": "量",           # 成交量(额为空,用量代替)
                "f8": "量",            # 成交量
            }.get(排序字段, "(收-开)/开")
            方向 = "DESC" if 排序方向 == "desc" else "ASC"
            每页 = 20
            offset = (页码 - 1) * 每页

            c.execute(f"""
                SELECT k.代码, k.收, k.量, k.额, a.名称,
                       (k.收-k.开)/k.开*100 as 涨幅
                FROM K线数据 k JOIN A股列表 a ON k.代码=a.代码
                WHERE k.周期='daily' AND k.日期=?
                ORDER BY {排序列} {方向}
                LIMIT ? OFFSET ?
            """, (最新日, 每页, offset))
            涨幅榜 = []
            for r in c.fetchall():
                涨幅 = r["涨幅"] if r["涨幅"] else 0
                # 成交额: 如果额为空, 用量×均价近似
                额 = r["额"] if r["额"] and r["额"] > 0 else (r["量"] or 0) * r["收"]
                涨幅榜.append({
                    "代码": r["代码"], "名称": r["名称"],
                    "最新价": round(r["收"], 2), "涨幅": round(涨幅, 2),
                    "涨速": 0,
                    "成交额": self._格式化成交额(额),
                    "量比": 0, "换手率": 0
                })

            # 3. 总数+市场总览(合并成一次查询)
            c.execute("""
                SELECT COUNT(*) as 总数,
                    SUM(CASE WHEN 收>开 THEN 1 ELSE 0 END) as 上涨,
                    SUM(CASE WHEN 收<开 THEN 1 ELSE 0 END) as 下跌,
                    SUM(CASE WHEN 收=开 THEN 1 ELSE 0 END) as 平盘,
                    SUM(CASE WHEN (收-开)/开*100>=9.8 THEN 1 ELSE 0 END) as 涨停,
                    SUM(CASE WHEN (收-开)/开*100<=-9.8 THEN 1 ELSE 0 END) as 跌停
                FROM K线数据 WHERE 周期='daily' AND 日期=?
            """, (最新日,))
            row = c.fetchone()
            总数 = row["总数"]

            # 4. 跌幅榜(最低10只)
            c.execute("""
                SELECT k.代码, k.收, k.额, a.名称,
                       (k.收-k.开)/k.开*100 as 涨幅
                FROM K线数据 k JOIN A股列表 a ON k.代码=a.代码
                WHERE k.周期='daily' AND k.日期=?
                ORDER BY 涨幅 ASC LIMIT 10
            """, (最新日,))
            跌幅榜 = []
            for r in c.fetchall():
                涨幅 = r["涨幅"] if r["涨幅"] else 0
                跌幅榜.append({
                    "代码": r["代码"], "名称": r["名称"],
                    "最新价": round(r["收"], 2), "涨幅": round(涨幅, 2),
                    "主力净流入": 0, "成交额": self._格式化成交额(r["额"] or 0)
                })

            conn.close()

            return {
                "成功": True,
                "时间": datetime.now().strftime("%H:%M:%S"),
                "指数": [],
                "涨幅榜": 涨幅榜,
                "跌幅榜": 跌幅榜,
                "市场总览": {"上涨": row["上涨"] or 0, "下跌": row["下跌"] or 0,
                            "平盘": row["平盘"] or 0, "涨停": row["涨停"] or 0, "跌停": row["跌停"] or 0},
                "涨幅榜总数": 总数,
                "当前页": 页码,
                "排序字段": 排序字段,
                "排序方向": 排序方向
            }
        except Exception as e:
            pass

        # 回退: 东财API
        from datetime import datetime
        try:
            # 1. 获取指数
            指数 = []
            指数代码 = [
                ("1.000001", "上证指数"), ("0.399001", "深证成指"),
                ("0.399006", "创业板指"), ("1.000688", "科创50")
            ]
            指数url = "https://push2.eastmoney.com/api/qt/ulist.np/get?fields=f1,f2,f3,f4,f12,f14&secids=" + ",".join([c for c, _ in 指数代码])
            try:
                数据 = self._东财请求(指数url)
                for item in (数据.get("data", {}).get("diff", {}) or {}).values():
                    指数.append({
                        "代码": item.get("f12", ""),
                        "名称": item.get("f14", ""),
                        "最新价": round(item.get("f2", 0) / 100, 2) if item.get("f2") else 0,
                        "涨跌幅": round(item.get("f3", 0) / 100, 2) if item.get("f3") else 0
                    })
            except Exception:
                pass

            # 2. 获取股票列表（按指定字段排序，支持翻页）
            涨幅榜 = []
            总数涨 = 0
            po = 1 if 排序方向 == "desc" else 0  # 1降序 0升序
            url涨 = f"https://push2.eastmoney.com/api/qt/clist/get?pn={页码}&pz=20&po={po}&np=1&fltt=2&invt=2&fields=f2,f3,f6,f7,f8,f9,f10,f12,f14,f15,f16,f17,f18,f20,f21,f62,f184,f66&fs=m:0+t:6,m:0+t:80,m:1+t:2,m:1+t:23&fid={排序字段}"
            try:
                数据 = self._东财请求(url涨)
                总数涨 = 数据.get("data", {}).get("total", 0)
                for item in (数据.get("data", {}).get("diff", []) or []):
                    涨幅榜.append({
                        "代码": item.get("f12", ""),
                        "名称": item.get("f14", ""),
                        "最新价": round(item.get("f2", 0) / 100, 2) if item.get("f2") else 0,
                        "涨幅": round(item.get("f3", 0) / 100, 2) if item.get("f3") else 0,
                        "涨速": round(item.get("f7", 0) / 100, 2) if item.get("f7") else 0,
                        "主力净流入": round((item.get("f62", 0) or 0) / 100000000, 2),
                        "成交额": self._格式化成交额(item.get("f6", 0)),
                        "量比": round(item.get("f8", 0) / 100, 2) if item.get("f8") else 0,
                        "换手率": round(item.get("f8", 0) / 100, 2) if item.get("f8") else 0,
                        "PE": round(item.get("f9", 0) / 100, 2) if item.get("f9") else 0,
                        "PB": round(item.get("f10", 0) / 100, 2) if item.get("f10") else 0
                    })
            except Exception:
                pass

            # 3. 获取跌幅榜
            跌幅榜 = []
            url跌 = "https://push2.eastmoney.com/api/qt/clist/get?pn=1&pz=10&po=0&np=1&fltt=2&invt=2&fields=f2,f3,f6,f12,f14,f62&fs=m:0+t:6,m:0+t:80,m:1+t:2,m:1+t:23&fid=f3"
            try:
                数据 = self._东财请求(url跌)
                for item in (数据.get("data", {}).get("diff", []) or []):
                    跌幅榜.append({
                        "代码": item.get("f12", ""),
                        "名称": item.get("f14", ""),
                        "最新价": round(item.get("f2", 0) / 100, 2) if item.get("f2") else 0,
                        "涨幅": round(item.get("f3", 0) / 100, 2) if item.get("f3") else 0,
                        "主力净流入": round((item.get("f62", 0) or 0) / 100000000, 2),
                        "成交额": self._格式化成交额(item.get("f6", 0))
                    })
            except Exception:
                pass

            # 4. 市场总览：涨跌家数、涨停跌停数
            市场总览 = {"上涨": 0, "下跌": 0, "平盘": 0, "涨停": 0, "跌停": 0}
            try:
                url总 = "https://push2.eastmoney.com/api/qt/clist/get?pn=1&pz=5000&po=1&np=1&fltt=2&invt=2&fields=f3,f12&fs=m:0+t:6,m:0+t:80,m:1+t:2,m:1+t:23&fid=f3"
                数据总 = self._东财请求(url总)
                for item in (数据总.get("data", {}).get("diff", []) or []):
                    pct = round(item.get("f3", 0) / 100, 2) if item.get("f3") else 0
                    code = item.get("f12", "")
                    涨停幅 = 19.9 if (code.startswith("30") or code.startswith("68")) else 9.9
                    跌停幅 = -19.9 if (code.startswith("30") or code.startswith("68")) else -9.9
                    if pct > 0:
                        市场总览["上涨"] += 1
                        if pct >= 涨停幅:
                            市场总览["涨停"] += 1
                    elif pct < 0:
                        市场总览["下跌"] += 1
                        if pct <= 跌停幅:
                            市场总览["跌停"] += 1
                    else:
                        市场总览["平盘"] += 1
            except Exception:
                pass

            return {
                "成功": True,
                "时间": datetime.now().strftime("%H:%M:%S"),
                "指数": 指数,
                "涨幅榜": 涨幅榜,
                "跌幅榜": 跌幅榜,
                "市场总览": 市场总览,
                "涨幅榜总数": 总数涨,
                "当前页": 页码,
                "排序字段": 排序字段,
                "排序方向": 排序方向
            }
        except Exception as e:
            return {"成功": False, "错误": str(e)}

    def _格式化成交额(self, 值):
        """格式化成交额（东财返回的是元）"""
        if not 值: return "-"
        亿 = 值 / 100000000
        if 亿 >= 1: return f"{亿:.1f}亿"
        万 = 值 / 10000
        return f"{万:.0f}万"

    def _获取股票K线(self, 代码: str, 周期: str = "daily") -> dict:
        """获取K线数据 — 一次连接读数据库, 回退东财API"""
        try:
            import sqlite3 as _sqlite3
            db路径 = str(Path(__file__).parent.parent.parent / "隐私区" / "我的数据" / "股票缓存.db")
            conn = _sqlite3.connect(db路径, timeout=5)
            conn.row_factory = _sqlite3.Row
            c = conn.cursor()

            # 一次查询: K线+名称(JOIN)
            c.execute("""
                SELECT k.日期, k.开, k.收, k.高, k.低, k.量, k.额, a.名称
                FROM K线数据 k
                LEFT JOIN A股列表 a ON k.代码 = a.代码
                WHERE k.代码=? AND k.周期=?
                ORDER BY k.日期 DESC LIMIT 120
            """, (代码, 周期))
            rows = c.fetchall()
            conn.close()

            if rows:
                rows.reverse()  # 正序
                结果 = [{"日期": r["日期"], "开": float(r["开"]), "收": float(r["收"]),
                         "高": float(r["高"]), "低": float(r["低"]),
                         "量": float(r["量"]), "额": float(r["额"]) if r["额"] else 0.0} for r in rows]
                名称 = rows[0]["名称"] or 代码
                最新 = 结果[-1]["收"]
                前收 = 结果[-2]["收"] if len(结果) >= 2 else 最新
                涨跌幅 = (最新 - 前收) / 前收 * 100 if 前收 > 0 else 0
                # MA计算
                收盘价 = [d["收"] for d in 结果]
                def ma(n, data):
                    return round(sum(data[-n:]) / n, 2) if len(data) >= n else 0
                return {
                    "成功": True, "数据": 结果,
                    "MA5": ma(5, 收盘价), "MA10": ma(10, 收盘价), "MA20": ma(20, 收盘价),
                    "股票信息": {"名称": 名称, "代码": 代码, "最新价": 最新, "涨跌幅": 涨跌幅}
                }
        except Exception as e:
            pass  # 数据库没有, 回退API

        # 回退: 东财API
        try:
            secid = self._代码转secid(代码)
            if not secid:
                return {"成功": False, "错误": f"无法识别股票代码: {代码}"}
            周期映射 = {"daily": 101, "weekly": 102, "monthly": 103}
            klt = 周期映射.get(周期, 101)
            天数 = 120 if klt == 101 else (200 if klt == 102 else 240)
            url = f"https://push2his.eastmoney.com/api/qt/stock/kline/get?secid={secid}&fields1=f1,f2,f3,f4,f5,f6&fields2=f51,f52,f53,f54,f55,f56,f57&klt={klt}&fqt=1&beg=0&end=20500101&lmt={天数}"
            数据 = self._东财请求(url)
            klines = 数据.get("data", {}).get("klines", []) or []
            结果 = []
            for k in klines:
                parts = k.split(",")
                if len(parts) >= 7:
                    结果.append({
                        "日期": parts[0], "开": float(parts[1]),
                        "收": float(parts[2]), "高": float(parts[3]),
                        "低": float(parts[4]), "量": float(parts[5]),
                        "额": float(parts[6])
                    })
            # 股票信息
            info = 数据.get("data", {}) or {}
            股票信息 = {
                "名称": info.get("name", ""),
                "代码": info.get("code", 代码),
                "最新价": 结果[-1]["收"] if 结果 else 0,
                "涨跌幅": ((结果[-1]["收"] - 结果[-2]["收"]) / 结果[-2]["收"] * 100) if len(结果) >= 2 else 0
            }
            # 计算MA5/MA10/MA20
            if len(结果) >= 5:
                股票信息["MA5"] = round(sum(d["收"] for d in 结果[-5:]) / 5, 2)
            if len(结果) >= 10:
                股票信息["MA10"] = round(sum(d["收"] for d in 结果[-10:]) / 10, 2)
            if len(结果) >= 20:
                股票信息["MA20"] = round(sum(d["收"] for d in 结果[-20:]) / 20, 2)
            # 计算完整MA序列（供前端画线）
            ma5_list, ma10_list, ma20_list = [], [], []
            for i in range(len(结果)):
                if i >= 4:
                    ma5_list.append(round(sum(d["收"] for d in 结果[i-4:i+1]) / 5, 2))
                else:
                    ma5_list.append(None)
                if i >= 9:
                    ma10_list.append(round(sum(d["收"] for d in 结果[i-9:i+1]) / 10, 2))
                else:
                    ma10_list.append(None)
                if i >= 19:
                    ma20_list.append(round(sum(d["收"] for d in 结果[i-19:i+1]) / 20, 2))
                else:
                    ma20_list.append(None)
            return {"成功": True, "数据": 结果, "MA5": ma5_list, "MA10": ma10_list, "MA20": ma20_list, "股票信息": 股票信息, "周期": 周期}
        except Exception as e:
            return {"成功": False, "错误": str(e)}

    def _获取股票分时(self, 代码: str) -> dict:
        """获取分时数据"""
        try:
            secid = self._代码转secid(代码)
            if not secid:
                return {"成功": False, "错误": f"无法识别股票代码: {代码}"}
            url = f"https://push2his.eastmoney.com/api/qt/stock/trends2/get?secid={secid}&fields1=f1,f2,f3,f4,f5,f6,f7,f8,f9,f10,f11,f12,f13&fields2=f51,f52,f53,f54,f55,f56,f57&iscr=0&ndays=1"
            数据 = self._东财请求(url)
            trends = 数据.get("data", {}).get("trends", []) or []
            结果 = []
            for t in trends:
                parts = t.split(",")
                if len(parts) >= 6:
                    结果.append({
                        "时间": parts[0],
                        "价格": float(parts[1]),
                        "均价": float(parts[2]) if len(parts) > 2 else None,
                        "量": float(parts[4]) if len(parts) > 4 else 0
                    })
            info = 数据.get("data", {}) or {}
            昨收 = info.get("prePreClose", 0) or info.get("preClose", 0)
            股票信息 = {
                "名称": info.get("name", ""),
                "代码": info.get("code", 代码),
                "最新价": 结果[-1]["价格"] if 结果 else 0,
                "涨跌幅": ((结果[-1]["价格"] - 昨收) / 昨收 * 100) if 结果 and 昨收 else 0
            }
            return {"成功": True, "数据": 结果, "昨收价": 昨收, "股票信息": 股票信息}
        except Exception as e:
            return {"成功": False, "错误": str(e)}

    def _代码转secid(self, 代码: str) -> str:
        """股票代码转东财secid（如 600519 → 1.600519, 000001 → 0.000001）"""
        代码 = 代码.strip()
        # 指数
        指数映射 = {"000001": "1.000001", "399001": "0.399001", "399006": "0.399006", "000688": "1.000688"}
        if 代码 in 指数映射:
            return 指数映射[代码]
        # 个股：6开头=上海(1)，0/3开头=深圳(0)
        if 代码.startswith("6"):
            return f"1.{代码}"
        elif 代码.startswith(("0", "3")):
            return f"0.{代码}"
        elif 代码.startswith("8") or 代码.startswith("4"):
            return f"0.{代码}"  # 北交所
        return ""

    def _批量查询行情(self, 代码列表: list) -> dict:
        """批量查询多只股票的实时行情（一次API请求）"""
        try:
            if not 代码列表:
                return {"成功": True, "数据": []}
            from 股票缓存 import 获取股票缓存
            缓存 = 获取股票缓存()
            # 先查缓存
            joined_codes = "-".join(sorted(代码列表))
            缓存键 = "batch_" + joined_codes
            cached = 缓存.读取缓存(缓存键, "batch")
            if cached is not None:
                return cached
            # 批量请求东财
            secids = []
            for 代码 in 代码列表:
                secid = self._代码转secid(代码)
                if secid:
                    secids.append(secid)
            if not secids:
                return {"成功": False, "错误": "无有效代码"}
            url = "https://push2.eastmoney.com/api/qt/ulist.np/get?fields=f2,f3,f4,f6,f12,f14,f15,f16,f17,f62&secids=" + ",".join(secids)
            数据 = self._东财请求(url)
            结果 = []
            for item in (数据.get("data", {}).get("diff", {}) or {}).values():
                结果.append({
                    "代码": item.get("f12", ""),
                    "名称": item.get("f14", ""),
                    "最新价": round(item.get("f2", 0) / 100, 2) if item.get("f2") else 0,
                    "涨跌幅": round(item.get("f3", 0) / 100, 2) if item.get("f3") else 0,
                    "涨跌额": round(item.get("f4", 0) / 100, 2) if item.get("f4") else 0,
                    "成交额": self._格式化成交额(item.get("f6", 0)),
                    "最高": round(item.get("f15", 0) / 100, 2) if item.get("f15") else 0,
                    "最低": round(item.get("f16", 0) / 100, 2) if item.get("f16") else 0,
                    "主力净流入": round((item.get("f62", 0) or 0) / 100000000, 2)
                })
            返回 = {"成功": True, "数据": 结果}
            缓存.写入缓存(缓存键, "batch", 返回)
            return 返回
        except Exception as e:
            return {"成功": False, "错误": str(e)}

    def _导出K线CSV(self, 代码: str, 周期: str = "daily") -> str:
        """导出K线数据为CSV字符串"""
        from 股票缓存 import 获取股票缓存
        缓存 = 获取股票缓存()
        数据 = 缓存.读取或请求(f"kline_{代码}_{周期}", "kline", lambda: self._获取股票K线(代码, 周期))
        if not 数据 or not 数据.get("成功"):
            return "错误,数据获取失败\n"
        lines = ["日期,开盘,收盘,最高,最低,成交量,成交额"]
        for d in 数据.get("数据", []):
            lines.append(f"{d['日期']},{d['开']},{d['收']},{d['高']},{d['低']},{d['量']},{d['额']}")
        return "\n".join(lines)

    def _返回CSV(self, csv内容: str, 文件名: str):
        """返回CSV文件下载响应"""
        self.send_response(200)
        self.send_header("Content-Type", "text/csv; charset=utf-8")
        self.send_header("Content-Disposition", f'attachment; filename="{文件名}"')
        body = csv内容.encode("utf-8-sig")  # BOM for Excel
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _搜索股票(self, 关键词: str) -> dict:
        """搜索股票（代码/名称模糊匹配）"""
        try:
            关键词 = 关键词.strip()
            if not 关键词:
                return {"成功": True, "结果": []}
            url = f"https://searchapi.eastmoney.com/api/suggest/get?input={关键词}&type=14&token=D43BF722C8E33BDC906FB84D85E326E8"
            数据 = self._东财请求(url)
            结果 = []
            for item in (数据.get("QuotationCodeTable", {}).get("Data", []) or []):
                code = item.get("Code", "")
                name = item.get("Name", "")
                mkt = item.get("MktNum", "")
                # 只保留A股
                if mkt in ("0", "1") or code.startswith(("0", "3", "6", "8", "4")):
                    cat = "指数" if "指数" in name or "成指" in name else "A股"
                    结果.append({"代码": code, "名称": name, "类型": cat})
            return {"成功": True, "结果": 结果[:20]}
        except Exception as e:
            return {"成功": False, "错误": str(e)}

    def _获取股票详情(self, 代码: str) -> dict:
        """获取个股详情：PE/PB/市值/换手率等"""
        try:
            secid = self._代码转secid(代码)
            if not secid:
                return {"成功": False, "错误": f"无法识别股票代码: {代码}"}
            url = f"https://push2.eastmoney.com/api/qt/stock/get?secid={secid}&fields=f57,f58,f84,f85,f86,f92,f116,f117,f162,f167,f168,f169,f170,f171,f173,f177,f183,f184,f186,f187,f188,f190,f191"
            数据 = self._东财请求(url)
            d = 数据.get("data", {}) or {}
            if not d:
                return {"成功": False, "错误": "未获取到数据"}
            def _val(key):
                v = d.get(key, 0)
                return v if v else 0
            def _div100(key):
                v = d.get(key, 0)
                return round(v / 100, 2) if v else 0
            详情 = {
                "代码": d.get("f57", 代码),
                "名称": d.get("f58", ""),
                "最新价": _div100("f84") if d.get("f84") else _val("f43"),
                "涨跌幅": _div100("f170"),
                "涨跌额": _div100("f169"),
                "成交量": self._格式化成交额(d.get("f135", 0) or d.get("f5", 0)),
                "成交额": self._格式化成交额(d.get("f6", 0)),
                "振幅": _div100("f171"),
                "换手率": _div100("f168"),
                "市盈率(动)": _div100("f162"),
                "市盈率(静)": _div100("f167"),
                "市净率": _div100("f184"),
                "总市值": self._格式化成交额(d.get("f116", 0)),
                "流通市值": self._格式化成交额(d.get("f117", 0)),
                "52周最高": _div100("f177"),
                "52周最低": _div100("f183"),
                "上市日期": d.get("f186", ""),
            }
            return {"成功": True, "详情": 详情}
        except Exception as e:
            return {"成功": False, "错误": str(e)}

    def _获取板块行情(self) -> dict:
        """获取板块行情：行业板块+概念板块"""
        try:
            结果 = {"行业": [], "概念": []}
            # 行业板块
            url行 = "https://push2.eastmoney.com/api/qt/clist/get?pn=1&pz=20&po=1&np=1&fltt=2&invt=2&fields=f2,f3,f4,f8,f12,f14,f104,f105,f128&fs=m:90+t:2&fid=f3"
            try:
                数据 = self._东财请求(url行)
                for item in (数据.get("data", {}).get("diff", []) or []):
                    结果["行业"].append({
                        "代码": item.get("f12", ""),
                        "名称": item.get("f14", ""),
                        "涨跌幅": round(item.get("f3", 0) / 100, 2) if item.get("f3") else 0,
                        "涨家数": item.get("f104", 0),
                        "跌家数": item.get("f105", 0),
                        "领涨股": item.get("f128", ""),
                        "换手率": round(item.get("f8", 0) / 100, 2) if item.get("f8") else 0
                    })
            except Exception:
                pass
            # 概念板块
            url概 = "https://push2.eastmoney.com/api/qt/clist/get?pn=1&pz=20&po=1&np=1&fltt=2&invt=2&fields=f2,f3,f4,f8,f12,f14,f104,f105,f128&fs=m:90+t:3&fid=f3"
            try:
                数据 = self._东财请求(url概)
                for item in (数据.get("data", {}).get("diff", []) or []):
                    结果["概念"].append({
                        "代码": item.get("f12", ""),
                        "名称": item.get("f14", ""),
                        "涨跌幅": round(item.get("f3", 0) / 100, 2) if item.get("f3") else 0,
                        "涨家数": item.get("f104", 0),
                        "跌家数": item.get("f105", 0),
                        "领涨股": item.get("f128", ""),
                        "换手率": round(item.get("f8", 0) / 100, 2) if item.get("f8") else 0
                    })
            except Exception:
                pass
            return {"成功": True, "板块": 结果}
        except Exception as e:
            return {"成功": False, "错误": str(e)}

    def _获取资金流向(self, 代码: str) -> dict:
        """获取个股资金流向明细 — 优先本地数据库, 回退东财API"""
        # 优先从本地数据库读取
        try:
            import sqlite3 as _sqlite3
            db路径 = str(Path(__file__).parent.parent.parent / "隐私区" / "我的数据" / "股票缓存.db")
            conn = _sqlite3.connect(db路径, timeout=10)
            c = conn.cursor()
            # 检查是否有资金流向表
            c.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='资金流向'")
            if c.fetchone():
                c.execute("SELECT 日期, 主力净流入, 超大单净流入, 大单净流入, 中单净流入, 小单净流入 FROM 资金流向 WHERE 代码=? ORDER BY 日期 DESC LIMIT 20", (代码,))
                rows = c.fetchall()
                conn.close()
                if rows:
                    rows.reverse()
                    结果 = [{"日期": r[0], "主力净流入": float(r[1] or 0)/10000, "超大单净流入": float(r[2] or 0)/10000,
                             "大单净流入": float(r[3] or 0)/10000, "中单净流入": float(r[4] or 0)/10000,
                             "小单净流入": float(r[5] or 0)/10000, "主力净流入占比": 0} for r in rows]
                    return {"成功": True, "数据": 结果, "代码": 代码}
            conn.close()
        except Exception:
            pass

        # 回退: 东财API
            secid = self._代码转secid(代码)
            if not secid:
                return {"成功": False, "错误": f"无法识别股票代码: {代码}"}
            url = f"https://push2.eastmoney.com/api/qt/stock/fflow/daykline/get?secid={secid}&fields1=f1,f2,f3,f7&fields2=f51,f52,f53,f54,f55,f56,f57,f58,f59,f60,f61,f62,f63,f64,f65&lmt=5"
            数据 = self._东财请求(url)
            klines = 数据.get("data", {}).get("klines", []) or []
            结果 = []
            for k in klines:
                parts = k.split(",")
                if len(parts) >= 10:
                    结果.append({
                        "日期": parts[0],
                        "主力净流入": round(float(parts[1]) / 10000, 2),
                        "小单净流入": round(float(parts[2]) / 10000, 2),
                        "中单净流入": round(float(parts[3]) / 10000, 2),
                        "大单净流入": round(float(parts[5]) / 10000, 2) if len(parts) > 5 else 0,
                        "超大单净流入": round(float(parts[4]) / 10000, 2) if len(parts) > 4 else 0,
                        "主力净流入占比": round(float(parts[6]) / 100, 2) if len(parts) > 6 else 0
                    })
            return {"成功": True, "数据": 结果, "代码": 代码}
        except Exception as e:
            return {"成功": False, "错误": str(e)}


class 网页服务类:
    """Web服务管理器"""

    def __init__(self, 端口: int, 界面目录: Path):
        self.端口 = 端口
        self.界面目录 = 界面目录
        self.服务器 = None

    def 启动(self, 文件管理器=None, 配置加载器=None, 模型直连器=None,
             模块注册=None, 操作注册中心=None, 启动器实例=None, 运行诊断器=None):
        网页请求处理器.界面目录 = self.界面目录
        网页请求处理器.文件管理器 = 文件管理器
        网页请求处理器.配置加载器 = 配置加载器
        网页请求处理器.模型直连器 = 模型直连器
        网页请求处理器.模块注册 = 模块注册
        网页请求处理器.操作注册中心 = 操作注册中心
        网页请求处理器.运行诊断器 = 运行诊断器
        网页请求处理器._启动器实例 = 启动器实例
        # 直接设置调度器引用，避免多级属性查找出错
        if 启动器实例 and hasattr(启动器实例, '定时任务调度器'):
            网页请求处理器._定时任务调度器 = 启动器实例.定时任务调度器
        # 设置当前模型名
        if 模型直连器:
            网页请求处理器.当前模型名 = 模型直连器.当前模型名

        # 自定义线程服务器：客户端断开不崩溃
        class _健壮HTTPServer(ThreadingHTTPServer):
            daemon_threads = True  # 守护线程，主进程退出时自动清理
            allow_reuse_address = True

            def handle_error(self, request, client_address):
                """覆盖默认错误处理，防止连接异常打印到stderr导致误判"""
                import sys
                exc = sys.exc_info()[1]
                if isinstance(exc, (ConnectionAbortedError, ConnectionResetError,
                                    BrokenPipeError, OSError)):
                    return  # 客户端断开，静默忽略
                super().handle_error(request, client_address)

        self.服务器 = _健壮HTTPServer(("0.0.0.0", self.端口), 网页请求处理器)
        print(f"   ✅ Web服务已启动: http://localhost:{self.端口}")
        self.服务器.serve_forever()

    def 停止(self):
        if self.服务器:
            self.服务器.shutdown()
            print("🌐 Web服务已停止")
